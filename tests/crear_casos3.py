"""
Genera 20 casos N2 nuevos, los guarda como Word en ~/Downloads/casos3/
y corre cada uno contra el motor de sentencias.
"""
import sys, os, json, re, time
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', 'src'))
os.environ['GEMINI_API_KEY'] = 'AIzaSyDwQp6HcIiqamO0KFeiexL2T560-zvw0_w'

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

CASOS3_DIR = os.path.join(os.path.expanduser('~'), 'Downloads', 'casos3')
os.makedirs(CASOS3_DIR, exist_ok=True)

# ============================================================
# 20 CASOS N2 — TEMAS VARIADOS
# ============================================================
CASOS = [
    {
        "titulo": "Caso 1 — Medina c/ Aguas Cordobesas (Contaminación / Prescripción ambiental)",
        "filename": "Caso01_Medina_c_AguasCordobesas.docx",
        "caratula": "Medina, Roberto Ángel c/ Aguas Cordobesas S.A. s/ daño ambiental y resarcimiento",
        "expediente": "CA-22114/2024",
        "rama": "civil_comercial",
        "tipo_proceso": "ordinario",
        "tribunal": "Juzgado Federal de Córdoba N° 2",
        "nivel": "N2",
        "hechos": (
            "El actor es vecino de la localidad de Jesús María (Córdoba) que en 2016 detectó contaminación "
            "con arsénico en el suministro de agua potable de su propiedad rural. Realizó análisis privados "
            "en 2016 que confirmaron niveles 4 veces superiores al máximo admitido por el Código Alimentario "
            "Argentino (CAA). Presentó denuncia administrativa ante ERSEP en 2017, que fue archivada sin "
            "resolución. Recién en 2022 obtuvo respuesta del ente regulador reconociendo la irregularidad. "
            "Inició demanda civil en 2023 por los daños a la salud (cuadro renal crónico diagnosticado en 2019) "
            "y daños al inmueble. Aguas Cordobesas opone prescripción: sostiene que el plazo de 3 años del "
            "art. 2561 CCCN comenzó desde 2016. El actor sostiene que la prescripción se suspendió durante "
            "el trámite administrativo y que el daño es continuado."
        ),
        "partes": [
            {
                "nombre": "Medina, Roberto Ángel",
                "rol": "actor",
                "pretension": "Resarcimiento por daño a la salud ($3.200.000), daño al inmueble ($850.000) y daño moral ($500.000). Con carácter preventivo, solicita orden de cese de contaminación.",
                "fundamentos": (
                    "Ley 25.675 (LGA) arts. 28 y 30: daño ambiental de incidencia colectiva no prescribe según "
                    "art. 29. Art. 2561 CCCN: prescripción de 3 años para daños a la persona, pero el dies a quo "
                    "se corre al momento en que el damnificado conoce el daño con certeza (2022, respuesta ERSEP). "
                    "Arts. 2539-2543 CCCN: suspensión de prescripción durante trámite administrativo previo "
                    "obligatorio. Daño continuado: la contaminación persiste, el plazo no comienza a correr."
                ),
                "jurisprudencia": (
                    "CSJN, 'Mendoza c/ Estado Nacional' (Riachuelo), Fallos 329:2316. CSJN, 'Barrick Exploraciones "
                    "Argentinas', 2019. CCA Córdoba, 'Asociación Vecinal Pilar c/ EPAS', 2021."
                ),
            },
            {
                "nombre": "Aguas Cordobesas S.A.",
                "rol": "demandado",
                "pretension": "Excepción de prescripción. Subsidiariamente, rechazo de la demanda por falta de nexo causal entre el arsénico y la enfermedad renal.",
                "fundamentos": (
                    "Art. 2561 CCCN: el plazo de 3 años corre desde que el damnificado conoció el daño (2016, "
                    "análisis propios). El trámite administrativo no es 'previo obligatorio' en materia de daño "
                    "civil. El daño ambiental colectivo (LGA art. 29) no se confunde con el daño individual "
                    "resarcible. La afección renal reconoce múltiples factores etiológicos no vinculados al "
                    "arsénico del suministro."
                ),
                "jurisprudencia": (
                    "CSJN, 'Villarreal c/ Aguas Provinciales', 2018. CNCiv. Sala H, 'Gramajo c/ AYSA', 2021."
                ),
            },
        ],
        "cuestiones": [
            "¿El dies a quo de la prescripción del daño ambiental individual corre desde el conocimiento del hecho contaminante o desde la certeza médica del daño a la salud?",
            "¿La reclamación administrativa ante el ente regulador suspende el plazo de prescripción civil aunque no sea formalmente obligatoria?",
            "¿El daño por contaminación continuada constituye un daño de tracto sucesivo que renueva el plazo de prescripción?",
        ],
    },
    {
        "titulo": "Caso 2 — Pérez Ríos s/ error de tipo (Delito informático / Deepfake)",
        "filename": "Caso02_PerezRios_DelitoInformatico.docx",
        "caratula": "Ministerio Público Fiscal c/ Pérez Ríos, Ignacio s/ distribución de material de explotación sexual infantil",
        "expediente": "CPF-4417/2024",
        "rama": "penal",
        "tipo_proceso": "juicio_oral",
        "tribunal": "Tribunal Oral en lo Criminal Federal N° 5",
        "nivel": "N2",
        "hechos": (
            "El imputado distribuyó por Telegram 14 imágenes de menores en situación sexual. Las pericias "
            "técnicas determinaron que el 100% de las imágenes eran generadas por inteligencia artificial "
            "(deepfakes): no existe ningún menor real identificable. El imputado sabía que eran imágenes "
            "generadas por IA y así lo comunicó en el grupo de Telegram. No hay víctima real identificada. "
            "La defensa plantea que la Ley 25.087 (modificatoria del CP) sanciona la distribución de "
            "material de 'menores reales' y que el art. 128 CP no abarca imágenes sintéticas. "
            "El Fiscal sostiene que la norma protege la indemnidad sexual como bien jurídico abstracto "
            "y que la IA no hace a la conducta atípica."
        ),
        "partes": [
            {
                "nombre": "Ministerio Público Fiscal",
                "rol": "actor",
                "pretension": "Condena por distribución de material de explotación sexual infantil (art. 128 CP), pena de 6 a 15 años.",
                "fundamentos": (
                    "Art. 128 CP: protege la indemnidad sexual y el orden público, no exclusivamente al menor "
                    "identificable. Interpretación teleológica: la IA normaliza la demanda de CSAM real. "
                    "Protocolo de Lanzarote (Convenio de Budapest): los Estados deben criminalizar el material "
                    "sintético. La posibilidad técnica de generar imágenes hiperrealistas hace indiferente la "
                    "distinción real/sintético desde el punto de vista del bien jurídico."
                ),
                "jurisprudencia": (
                    "CSJN, 'Irurzun c/ Fiscal', Fallos 340:669. US v. Whorley (4th Cir. 2008). "
                    "TEDH, 'K.U. c/ Finlandia', 2008."
                ),
            },
            {
                "nombre": "Pérez Ríos, Ignacio",
                "rol": "demandado",
                "pretension": "Absolución. Las imágenes sintéticas no configuran el tipo penal del art. 128 CP por ausencia de menor real.",
                "fundamentos": (
                    "Art. 128 CP: la expresión 'representaciones de menores' requiere menor real conforme "
                    "interpretación literal y sistemática. Art. 18 CN: legalidad estricta. In dubio pro reo. "
                    "El error de tipo excluye el dolo si el imputado creyera que eran reales (que no ocurre "
                    "aquí). La conducta es atípica porque el bien jurídico protegido (indemnidad del menor "
                    "real) no está en juego."
                ),
                "jurisprudencia": (
                    "CNCP Sala I, 'Forno', 2019. Zaffaroni, Manual de Derecho Penal, § 34."
                ),
            },
        ],
        "cuestiones": [
            "¿El art. 128 CP abarca imágenes de menores generadas íntegramente por inteligencia artificial cuando no existe ningún menor real involucrado?",
            "¿El principio de legalidad estricta (art. 18 CN) impide extender el tipo penal a material sintético no mencionado expresamente en la norma?",
            "¿Qué bien jurídico protege el art. 128 CP y cómo incide esa determinación en la solución del caso?",
        ],
    },
    {
        "titulo": "Caso 3 — Gómez c/ Claro (Datos personales / Algoritmo crediticio)",
        "filename": "Caso03_Gomez_c_Claro.docx",
        "caratula": "Gómez, Valentina c/ AMX Argentina S.A. (Claro) y Veraz S.A. s/ hábeas data y daños",
        "expediente": "HD-9821/2024",
        "rama": "civil_comercial",
        "tipo_proceso": "hábeas_data",
        "tribunal": "Juzgado Nacional en lo Civil y Comercial Federal N° 7",
        "nivel": "N2",
        "hechos": (
            "La actora solicitó un crédito hipotecario al Banco Nación, que fue rechazado por un 'score "
            "crediticio bajo' generado automáticamente por Veraz S.A. Al consultar sus datos en Veraz, "
            "la actora descubrió que figuraba como deudora de Claro por una deuda de $18.000 por un "
            "servicio que nunca contrató (línea dada de alta con sus datos pero con domicilio y tarjeta "
            "diferentes). Claro reconoció el error pero solo eliminó la deuda del sistema de Claro, "
            "no notificó la rectificación a Veraz. El score de Veraz siguió siendo negativo 8 meses "
            "después. La actora perdió la oportunidad de acceder al crédito hipotecario (el beneficio "
            "venció). Reclama hábeas data correctivo y daños por la oportunidad perdida."
        ),
        "partes": [
            {
                "nombre": "Gómez, Valentina",
                "rol": "actora",
                "pretension": "Rectificación inmediata del score en Veraz, eliminación del dato falso, indemnización por pérdida de chance ($8.500.000 diferencia crédito/alquiler) y daño moral.",
                "fundamentos": (
                    "Ley 25.326 (PDPA) arts. 16 y 34: derecho de rectificación y actualización en 5 días hábiles. "
                    "Art. 4 PDPA: los datos deben ser exactos y actualizados. Responsabilidad solidaria de Claro "
                    "(fuente del dato falso) y Veraz (que lo mantiene). Reglamento GDPR (aplicable por analogía): "
                    "decisiones automatizadas que afectan derechos requieren intervención humana. "
                    "Art. 1738 CCCN: pérdida de chance como daño resarcible autónomo."
                ),
                "jurisprudencia": (
                    "CSJN, 'Martínez c/ Organización Veraz', Fallos 332:1625. CNCiv. Sala F, 'Bernal c/ Nosis', "
                    "2022. TDPD España, resolución 2023/0123 (scoring crediticio automatizado)."
                ),
            },
            {
                "nombre": "AMX Argentina S.A. (Claro) y Veraz S.A.",
                "rol": "demandado",
                "pretension": "Claro: cumplió con la baja de la deuda en su sistema. Veraz: es un mero repositorio, la obligación de notificación pesa sobre la fuente (Claro).",
                "fundamentos": (
                    "Ley 25.326 art. 26: los bancos de datos de riesgo crediticio tienen régimen especial. "
                    "Veraz no tiene obligación de rectificar de oficio: debe recibir la notificación de la "
                    "fuente (Claro). La pérdida de chance del crédito hipotecario no es imputable a Claro "
                    "sino a la demora de la actora en reclamar. La deuda falsa fue originada por fraude "
                    "de terceros, no por negligencia de Claro."
                ),
                "jurisprudencia": (
                    "CNCom. Sala C, 'López c/ Equifax', 2021. CNCP, 'Defensoría del Pueblo c/ Veraz', 2019."
                ),
            },
        ],
        "cuestiones": [
            "¿El responsable de la fuente del dato falso (Claro) tiene obligación de notificar la rectificación a todas las bases de datos a las que cedió el dato, o basta con la baja en su propio sistema?",
            "¿Veraz tiene obligación autónoma de actualizar el score cuando conoce que el dato fuente era falso, aunque no haya recibido notificación formal?",
            "¿La pérdida de la oportunidad de acceder al crédito hipotecario es resarcible como pérdida de chance y cuál es el estándar de causalidad aplicable?",
        ],
    },
    {
        "titulo": "Caso 4 — López c/ ART Galeno (Silicosis / Enfermedad profesional lenta)",
        "filename": "Caso04_Lopez_c_ARTGaleno.docx",
        "caratula": "López, Héctor Ramón c/ Galeno ART S.A. y Cementos Avellaneda S.A. s/ enfermedad profesional",
        "expediente": "SRT-18834/2024",
        "rama": "laboral",
        "tipo_proceso": "laboral_ordinario",
        "tribunal": "Juzgado Nacional del Trabajo N° 5",
        "nivel": "N2",
        "hechos": (
            "El actor trabajó como operario en la planta de Cementos Avellaneda durante 22 años (2000-2022) "
            "expuesto a sílice cristalina sin los elementos de protección adecuados. En 2022 fue diagnosticado "
            "con silicosis estadio II. La ART Galeno reconoció la enfermedad profesional pero determinó el "
            "índice de incapacidad (IBM) en 38%, que el actor impugna. La SRT fijó un IBM del 45% en instancia "
            "administrativa. El actor demanda también a la empleadora directamente (acción civil) alegando dolo "
            "eventual: la empresa conocía los riesgos por inspecciones laborales de 2008 y 2014 que ordenaron "
            "mejoras nunca implementadas. La demandada alega que el art. 39 LRT limita la acción civil."
        ),
        "partes": [
            {
                "nombre": "López, Héctor Ramón",
                "rol": "actor",
                "pretension": "Fijación de incapacidad en 58% (pericia médica propia). Acción civil contra empleadora por dolo eventual. Daño moral y psicológico no cubiertos por LRT.",
                "fundamentos": (
                    "LRT art. 14: la incapacidad permanente parcial mayor al 50% abre la vía civil. "
                    "CSJN, 'Torrillo c/ Gulf Oil' y 'Aquino c/ Cargo Servicios': el art. 39 LRT es "
                    "inconstitucional cuando el empleador actuó con dolo. Art. 1724 CCCN: dolo eventual "
                    "equiparado a dolo directo en materia de responsabilidad. Las inspecciones de 2008 y "
                    "2014 acreditan conocimiento del riesgo. El daño moral es extracontractual y no cubierto "
                    "por el sistema tarifado de la LRT."
                ),
                "jurisprudencia": (
                    "CSJN, 'Aquino c/ Cargo Servicios', Fallos 327:3753. CSJN, 'Torrillo c/ Gulf Oil', "
                    "Fallos 332:709. CSJN, 'Soria c/ Ra-Ma', 2019."
                ),
            },
            {
                "nombre": "Galeno ART S.A. y Cementos Avellaneda S.A.",
                "rol": "demandado",
                "pretension": "IBM del 38%. Rechazo de la acción civil: la LRT es el sistema exclusivo de reparación. No hubo dolo sino culpa leve en las condiciones de trabajo.",
                "fundamentos": (
                    "Art. 39 LRT: la LRT excluye la acción civil salvo dolo del empleador. Las inspecciones "
                    "laborales no acreditan dolo sino incumplimientos que generaron multas ya pagadas. "
                    "La silicosis estadio II no alcanza el 50% de IBM según tablas SRT. El art. 14 LRT "
                    "es constitucional según la reforma de Ley 27.348."
                ),
                "jurisprudencia": (
                    "CNAT Sala X, 'Ferreyra c/ Loma Negra', 2022. CSJN, 'Lucca de Hoz c/ Taddei', 2010."
                ),
            },
        ],
        "cuestiones": [
            "¿Las actas de inspección laboral que documentan incumplimientos reiterados en protección contra sílice son suficientes para acreditar dolo eventual del empleador y habilitar la acción civil del art. 1724 CCCN?",
            "¿La determinación del IBM por encima del 50% de incapacidad habilita automáticamente la acción civil o requiere además probar que el sistema tarifado es insuficiente?",
            "¿El daño moral por enfermedad profesional es resarcible por vía civil aunque la incapacidad no supere el umbral legal?",
        ],
    },
    {
        "titulo": "Caso 5 — Fundación Ambiente c/ Provincia de Salta (Desmonte / Cautelar ambiental)",
        "filename": "Caso05_FundacionAmbiente_c_Salta.docx",
        "caratula": "Fundación Ambiente y Recursos Naturales c/ Provincia de Salta s/ amparo ambiental colectivo",
        "expediente": "FAA-1102/2024",
        "rama": "administrativo",
        "tipo_proceso": "amparo",
        "tribunal": "Juzgado Federal de Salta N° 1",
        "nivel": "N2",
        "hechos": (
            "La Provincia de Salta otorgó en 2023 autorizaciones de desmonte sobre 48.000 hectáreas en el "
            "departamento Rivadavia, zona de bosque nativo de categoría II (amarilla) según la Ley 26.331. "
            "La Fundación FARN impugna las autorizaciones por falta de Evaluación de Impacto Ambiental "
            "acumulativa (se otorgaron 12 permisos individuales de menos de 4.000 ha cada uno, evitando "
            "el umbral que dispara la EIA obligatoria). Las comunidades indígenas Wichí de la zona no "
            "fueron consultadas (art. 75 inc. 17 CN, Convenio 169 OIT). La Provincia sostiene que cada "
            "permiso cumplió la normativa individual y que el fraccionamiento fue legítimo."
        ),
        "partes": [
            {
                "nombre": "Fundación Ambiente y Recursos Naturales (FARN)",
                "rol": "actor",
                "pretension": "Nulidad de las 12 autorizaciones de desmonte. Suspensión cautelar inmediata de las actividades. Obligación de realizar EIA acumulativa y consulta previa libre e informada.",
                "fundamentos": (
                    "Ley 26.331 (Bosques Nativos) art. 22: prohíbe el fraccionamiento artificial para eludir "
                    "la EIA. Principio de no regresión ambiental (LGA art. 4). Convenio 169 OIT art. 6: consulta "
                    "previa obligatoria. Art. 75 inc. 17 CN: derechos de comunidades indígenas. "
                    "Principio precautorio (LGA art. 4): ante incertidumbre de daño grave e irreversible, "
                    "la carga de la prueba se invierte. El desmonte es irreversible."
                ),
                "jurisprudencia": (
                    "CSJN, 'Salas c/ Salta' (Desmonte Pilcomayo), Fallos 332:663. CSJN, 'Comunidad Indígena "
                    "Eben Ezer', Fallos 331:2119. Corte IDH, 'Pueblo Saramaka c/ Surinam', 2007."
                ),
            },
            {
                "nombre": "Provincia de Salta",
                "rol": "demandado",
                "pretension": "Rechazo del amparo. Cada autorización cumplió individualmente la Ley 26.331. El fraccionamiento respondió a solicitudes independientes de distintos propietarios.",
                "fundamentos": (
                    "La Ley 26.331 no establece expresamente la EIA acumulativa. El ordenamiento territorial "
                    "de bosques (OTBN) de Salta fue homologado por la Nación. Las comunidades Wichí fueron "
                    "informadas en audiencias públicas que equivalen a consulta. El amparo no es la vía para "
                    "revisar actos administrativos discrecionales de política ambiental provincial."
                ),
                "jurisprudencia": (
                    "CSJN, 'Mendoza c/ Estado Nacional', Fallos 329:2316. TSJ Salta, 'Ávila c/ PROSA', 2022."
                ),
            },
        ],
        "cuestiones": [
            "¿El fraccionamiento de solicitudes de desmonte en unidades inferiores al umbral legal constituye fraude a la ley (art. 12 CCCN) que invalida las autorizaciones individuales?",
            "¿Las audiencias públicas informativas equivalen a la consulta previa libre e informada del Convenio 169 OIT, o requieren un proceso diferenciado con las comunidades indígenas afectadas?",
            "¿El principio precautorio invierte la carga de la prueba en materia de amparo ambiental, de modo que sea la Provincia quien deba demostrar la inocuidad del desmonte?",
        ],
    },
    {
        "titulo": "Caso 6 — Quiroga c/ PAMI (Prestación médica / Medicamentos oncológicos)",
        "filename": "Caso06_Quiroga_c_PAMI.docx",
        "caratula": "Quiroga, Elena Isabel c/ PAMI - INSSJP s/ amparo de salud",
        "expediente": "AS-44102/2024",
        "rama": "administrativo",
        "tipo_proceso": "amparo",
        "tribunal": "Juzgado Federal de Rosario N° 2",
        "nivel": "N2",
        "hechos": (
            "La actora, afiliada al PAMI de 71 años, padece cáncer de mama estadio III. Su oncóloga tratante "
            "le recetó en enero de 2024 el medicamento Pembrolizumab (inmunoterapia), que no figura en el "
            "vademécum de PAMI. PAMI rechazó la cobertura por no estar en el PMO vigente para su plan. "
            "La actora solicitó amparo por derecho a la salud. PAMI ofrece en su lugar Tamoxifeno "
            "(quimioterapia estándar), que según la oncóloga es significativamente menos eficaz para el "
            "subtipo tumoral específico de la paciente (HER2+). La actora no puede costear el medicamento "
            "($180.000 por ciclo mensual). La ANMAT aprobó el Pembrolizumab para el perfil de la paciente."
        ),
        "partes": [
            {
                "nombre": "Quiroga, Elena Isabel",
                "rol": "actora",
                "pretension": "Cobertura al 100% del Pembrolizumab durante el tratamiento prescripto. Medida cautelar urgente.",
                "fundamentos": (
                    "Art. 42 CN y Ley 23.661 (SNIS): derecho a la salud. Ley 27.455 (PMO): las obras sociales "
                    "deben cubrir toda prestación con evidencia científica para la patología del afiliado. "
                    "El PMO es un piso, no un techo: la CSJN ha ordenado cobertura de medicamentos fuera del "
                    "vademécum cuando son el tratamiento más adecuado. Art. 3 CDN (aplicable analógicamente "
                    "por interés superior en situaciones de vulnerabilidad). ANMAT aprobó el medicamento."
                ),
                "jurisprudencia": (
                    "CSJN, 'Campodónico de Beviacqua c/ Ministerio de Salud', Fallos 323:3229. "
                    "CSJN, 'Asociación Francesa Filantrópica', Fallos 331:2135. "
                    "CNCiv. Sala M, 'Villalba c/ OSDE', 2023."
                ),
            },
            {
                "nombre": "PAMI - INSSJP",
                "rol": "demandado",
                "pretension": "Rechazo del amparo. El PAMI ofrece cobertura del tratamiento estándar aprobado en su vademécum (Tamoxifeno). La sustentabilidad del sistema impide cubrir medicamentos fuera del PMO.",
                "fundamentos": (
                    "El PAMI tiene facultades de auto-regulación de su vademécum (Ley 19.032). El Pembrolizumab "
                    "es de altísimo costo y no está en el PMO nacional. La cobertura universal de inmunoterapias "
                    "haría insostenible el sistema. El médico de cabecera del PAMI puede prescribir alternativas "
                    "igualmente eficaces. El amparo no puede sustituir la política sanitaria del organismo."
                ),
                "jurisprudencia": (
                    "CNCAF Sala V, 'Defensoría del Pueblo c/ PAMI', 2021. CSJN, 'Sánchez c/ Gobierno de la Ciudad', 2012."
                ),
            },
        ],
        "cuestiones": [
            "¿El derecho constitucional a la salud obliga al PAMI a cubrir medicamentos oncológicos fuera del vademécum cuando constituyen el tratamiento de mayor eficacia para el subtipo tumoral específico del afiliado?",
            "¿La sustentabilidad del sistema de salud puede oponerse como defensa frente al derecho individual a la vida de una afiliada en tratamiento oncológico?",
            "¿El Comité Médico del PAMI puede sustituir el criterio terapéutico del oncólogo tratante o el juez debe privilegiar la prescripción especializada?",
        ],
    },
    {
        "titulo": "Caso 7 — Álvarez c/ Municipalidad de Mendoza (Licitación / Conflicto de intereses)",
        "filename": "Caso07_Alvarez_c_MunicipalidadMendoza.docx",
        "caratula": "Álvarez, Gustavo Marcos c/ Municipalidad de Mendoza s/ nulidad de acto administrativo (licitación pública)",
        "expediente": "CA-5531/2024",
        "rama": "administrativo",
        "tipo_proceso": "contencioso_administrativo",
        "tribunal": "Cámara de Apelaciones en lo Contencioso Administrativo de Mendoza",
        "nivel": "N2",
        "hechos": (
            "La Municipalidad de Mendoza licitó la concesión del estacionamiento medido de la ciudad "
            "(contrato por 15 años, canon inicial de $12 millones mensuales). El concejal Ramírez, quien "
            "integró la comisión evaluadora de ofertas, es cuñado del socio mayoritario de la empresa "
            "adjudicataria 'Parking Cuyo S.A.' Esta relación no fue declarada por el concejal ni "
            "informada a la comisión. El actor Álvarez es titular de la empresa segunda ofertante "
            "(diferencia de puntaje: 3 puntos sobre 100). Impugna el acto administrativo de adjudicación "
            "por vicio de conflicto de intereses no declarado. La Municipalidad sostiene que el cuñado "
            "no está incluido en el listado de inhabilidades del Reglamento de Compras municipal."
        ),
        "partes": [
            {
                "nombre": "Álvarez, Gustavo Marcos",
                "rol": "actor",
                "pretension": "Nulidad absoluta del acto de adjudicación. Nueva evaluación sin participación del concejal Ramírez. Indemnización por daño concursal.",
                "fundamentos": (
                    "Ley 25.188 (Ética Pública) arts. 13 y 14: deber de excusación ante interés de parientes "
                    "hasta 4° grado de afinidad (cuñado = 2° grado). Art. 14 inc. f) Ley 19.549 (LNPA): vicio "
                    "en la finalidad del acto. Principio de transparencia (Ley 27.275). La nulidad del acto "
                    "administrativo no requiere demostrar que el concejal influyó en el resultado, basta la "
                    "apariencia de imparcialidad comprometida (CSJN)."
                ),
                "jurisprudencia": (
                    "CSJN, 'Laplacette c/ Provincia de Buenos Aires', Fallos 269:243. "
                    "PTN, Dictamen 137:107. Suprema Corte de Mendoza, 'Fernández c/ Municipalidad', 2021."
                ),
            },
            {
                "nombre": "Municipalidad de Mendoza",
                "rol": "demandado",
                "pretension": "Validez del acto de adjudicación. El cuñado no está incluido en las inhabilidades del Reglamento municipal de compras.",
                "fundamentos": (
                    "El Reglamento de Compras municipal lista taxativamente las relaciones que generan "
                    "inhabilidad: cónyuge, hijos, padres. La Ley 25.188 aplica al Poder Ejecutivo nacional, "
                    "no a los municipios provinciales (autonomía municipal, art. 123 CN). El concejal no "
                    "tenía poder decisorio unilateral: la adjudicación fue votada por la comisión completa "
                    "de 5 miembros. La empresa adjudicataria obtuvo objetivamente la mayor puntuación."
                ),
                "jurisprudencia": (
                    "SCJBA, 'Rosario c/ Municipalidad de La Plata', 2019. PTN, Dictamen 242:89."
                ),
            },
        ],
        "cuestiones": [
            "¿La Ley Nacional de Ética Pública 25.188 aplica a funcionarios municipales de la Provincia de Mendoza o la autonomía municipal genera un régimen exclusivo?",
            "¿El vicio de conflicto de intereses no declarado por un miembro de la comisión evaluadora acarrea nulidad absoluta del acto de adjudicación aunque los demás miembros votaron en la misma dirección?",
            "¿Es suficiente para anular la adjudicación la apariencia de parcialidad o se requiere demostrar que el funcionario incidió efectivamente en el resultado?",
        ],
    },
    {
        "titulo": "Caso 8 — Romano c/ Estado Nacional (Responsabilidad del Estado / Error judicial)",
        "filename": "Caso08_Romano_c_EstadoNacional.docx",
        "caratula": "Romano, Diego Sebastián c/ Estado Nacional (Ministerio de Justicia) s/ daños y perjuicios por error judicial",
        "expediente": "CCA-6612/2024",
        "rama": "administrativo",
        "tipo_proceso": "contencioso_administrativo",
        "tribunal": "Cámara Nacional de Apelaciones en lo Contencioso Administrativo Federal, Sala III",
        "nivel": "N2",
        "hechos": (
            "El actor estuvo detenido preventivamente durante 26 meses (2019-2021) acusado de estafa. "
            "Fue absuelto en juicio oral en agosto de 2021 por 'estado de duda insuperable'. Durante "
            "la detención perdió su empleo de gerente de ventas ($280.000 mensuales), su matrimonio "
            "se disolvió y sufrió un cuadro depresivo severo certificado. La prisión preventiva fue "
            "prorrogada tres veces por el juez instructor sin evaluación de medidas alternativas. "
            "La Cámara de Apelaciones revocó la última prórroga pero el actor ya llevaba 22 meses detenido "
            "en ese momento. El actor demanda al Estado por error judicial y funcionamiento anormal "
            "del servicio de justicia."
        ),
        "partes": [
            {
                "nombre": "Romano, Diego Sebastián",
                "rol": "actor",
                "pretension": "Resarcimiento por lucro cesante ($8.736.000), pérdida de chance laboral ($3.000.000), daño moral ($2.500.000) y daño a la salud ($1.200.000).",
                "fundamentos": (
                    "Ley 24.906 y art. 10 Ley 27.401: responsabilidad del Estado por error judicial. "
                    "Art. 2 Ley 26.944 (LRESP): responsabilidad del Estado por actividad judicial ilegítima. "
                    "La absolución por duda razonable no equivale a error judicial 'reconocido', pero la "
                    "prolongación indebida de la prisión preventiva constituye funcionamiento anormal "
                    "independiente. CADH art. 10: indemnización por error judicial. CPP arts. 280 y 319: "
                    "la prórroga sin análisis de alternativas es contraria al principio de excepcionalidad."
                ),
                "jurisprudencia": (
                    "CSJN, 'Balda c/ Provincia de Buenos Aires', Fallos 318:1990. "
                    "CSJN, 'Rosa c/ Estado Nacional', Fallos 342:2131. "
                    "Corte IDH, 'Chaparro Álvarez y Lapo Íñiguez c/ Ecuador', 2007."
                ),
            },
            {
                "nombre": "Estado Nacional (Ministerio de Justicia)",
                "rol": "demandado",
                "pretension": "Rechazo de la demanda. La absolución por duda no configura error judicial reconocido. La prisión preventiva fue legalmente dictada.",
                "fundamentos": (
                    "Ley 26.944 art. 5: para la responsabilidad por actividad judicial se requiere error "
                    "judicial 'declarado' por sentencia firme o sobreseimiento definitivo por inexistencia "
                    "del hecho. La absolución por duda no equivale a inocencia declarada. La prisión "
                    "preventiva fue dispuesta y prorrogada con fundamento en los méritos de la causa. "
                    "El lucro cesante laboral requiere prueba concreta de causalidad directa."
                ),
                "jurisprudencia": (
                    "CSJN, 'Ramos c/ Provincia de Buenos Aires', Fallos 310:943. "
                    "CNCAF Sala IV, 'Ottino c/ Estado Nacional', 2022."
                ),
            },
        ],
        "cuestiones": [
            "¿La absolución por 'estado de duda insuperable' configura el 'error judicial' que habilita la responsabilidad del Estado conforme la Ley 26.944, o se requiere la declaración expresa de inocencia?",
            "¿La prórroga reiterada de la prisión preventiva sin análisis de medidas alternativas configura 'funcionamiento anormal del servicio de justicia' independiente del error judicial, generando responsabilidad estatal autónoma?",
            "¿Cuál es el estándar de causalidad entre la detención y la pérdida del empleo como gerente para habilitar el lucro cesante?",
        ],
    },
    {
        "titulo": "Caso 9 — Díaz c/ Empresa de Colectivos (Accidente in itinere / Trayecto desviado)",
        "filename": "Caso09_Diaz_c_EmpresaColectivos.docx",
        "caratula": "Díaz, Miriam Graciela c/ Transporte Automotor Río Segundo S.A. y ART Provincia s/ accidente in itinere",
        "expediente": "SRT-22134/2024",
        "rama": "laboral",
        "tipo_proceso": "laboral_ordinario",
        "tribunal": "Cámara Nacional del Trabajo, Sala VI",
        "nivel": "N2",
        "hechos": (
            "La actora sufrió un accidente de tránsito el 14/06/2023 mientras regresaba de su trabajo "
            "en colectivo. Normalmente tomaba la línea 36 directa. Ese día, debido a un paro parcial "
            "de la línea 36, debió tomar la línea 47 que la dejaba a 8 cuadras de su domicilio, "
            "debiendo caminar ese tramo. Mientras caminaba fue atropellada por una motocicleta. "
            "La ART rechazó el siniestro argumentando que el trayecto era 'no habitual' y que la "
            "desviación implica salida del ámbito de cobertura del art. 6 inc. 3 LRT. La actora "
            "sostiene que el desvío fue involuntario, motivado por causas externas (el paro de colectivos). "
            "Sufrió fractura de cadera con 35% de incapacidad."
        ),
        "partes": [
            {
                "nombre": "Díaz, Miriam Graciela",
                "rol": "actora",
                "pretension": "Reconocimiento del accidente como in itinere. Prestaciones dinerarias por incapacidad permanente parcial del 35%. Daño moral.",
                "fundamentos": (
                    "Art. 6 inc. 3 LRT: el accidente in itinere cubre el trayecto entre el domicilio y el "
                    "trabajo en condiciones normales. La doctrina y jurisprudencia admiten 'desvíos razonables' "
                    "por causas de fuerza mayor o necesidades vitales. El paro de colectivos es un hecho "
                    "externo que no supone voluntad de desvío. La finalidad laboral del desplazamiento "
                    "se mantuvo. El principio protectorio (art. 14 bis CN) exige interpretación favorable al trabajador."
                ),
                "jurisprudencia": (
                    "CNAT Sala II, 'Fernández c/ La Meridional ART', 2022. "
                    "CNAT Sala VIII, 'Bolaños c/ Provincia ART', 2021. "
                    "Suprema Corte de Mendoza, 'Rueda c/ Federación Patronal ART', 2019."
                ),
            },
            {
                "nombre": "Transporte Automotor Río Segundo S.A. y ART Provincia",
                "rol": "demandado",
                "pretension": "Rechazo del siniestro. El trayecto por la línea 47 y los 8 cuadras a pie no corresponden al trayecto declarado en el formulario de accidente in itinere.",
                "fundamentos": (
                    "Art. 6 inc. 3 LRT: el trayecto debe ser el habitual declarado. La trabajadora no "
                    "comunicó el cambio de trayecto. La SRT Resolución 1601/2007 exige declaración del "
                    "itinerario habitual. El paro parcial de la línea 36 no es fuerza mayor que justifique "
                    "el desvío sin notificación. El empleador no puede prever riesgos en trayectos no declarados."
                ),
                "jurisprudencia": (
                    "CNAT Sala III, 'Herrera c/ Asociart ART', 2020. "
                    "SCBA, 'Mansilla c/ QBE ART', 2018."
                ),
            },
        ],
        "cuestiones": [
            "¿El cambio de trayecto laboral motivado por un paro parcial de la línea habitual de colectivos constituye un 'desvío razonable' que mantiene la cobertura del accidente in itinere del art. 6 inc. 3 LRT?",
            "¿La falta de declaración previa del trayecto alternativo en el formulario SRT excluye automáticamente la cobertura o cede ante la prueba de la causa justificada del desvío?",
            "¿El principio protectorio del art. 14 bis CN obliga a interpretar restrictivamente las exclusiones del art. 6 LRT?",
        ],
    },
    {
        "titulo": "Caso 10 — Fernández c/ Provincia de Entre Ríos (Policía / Uso excesivo de la fuerza)",
        "filename": "Caso10_Fernandez_c_ProvinciaEntreRios.docx",
        "caratula": "Fernández, Lucas Nahuel c/ Provincia de Entre Ríos s/ daños y perjuicios (responsabilidad del Estado provincial)",
        "expediente": "CA-3312/2024",
        "rama": "administrativo",
        "tipo_proceso": "contencioso_administrativo",
        "tribunal": "Cámara de Apelaciones en lo Civil y Comercial de Paraná, Sala II",
        "nivel": "N2",
        "hechos": (
            "El actor, de 22 años, participó en una manifestación pacífica frente a la Legislatura "
            "provincial el 12/09/2022. Efectivos de la Policía de Entre Ríos dispersaron la manifestación "
            "con balas de goma. El actor fue alcanzado por dos proyectiles en la espalda a distancia de "
            "aproximadamente 8 metros, cuando ya se retiraba. Sufrió contusiones graves y una lesión renal "
            "que requirió internación de 12 días. El peritaje balístico determinó que los disparos se "
            "realizaron a menor distancia de la reglamentaria para ese tipo de munición (mínimo 20 metros). "
            "La Provincia sostiene que los policías actuaron conforme el protocolo de uso de la fuerza "
            "aprobado por Decreto Provincial. El actor no fue detenido ni imputado penalmente."
        ),
        "partes": [
            {
                "nombre": "Fernández, Lucas Nahuel",
                "rol": "actor",
                "pretension": "Daño físico ($850.000), daño moral ($600.000), daño psicológico ($300.000) y declaración de inconstitucionalidad del Decreto provincial de uso de fuerza.",
                "fundamentos": (
                    "Art. 10 de la Ley 26.944: responsabilidad del Estado por actividad lícita e ilícita. "
                    "Art. 3 DUDH y arts. 4 y 5 CADH: derecho a la integridad personal. "
                    "Principios Básicos ONU sobre uso de la fuerza (1990): proporcionalidad y necesidad. "
                    "El protocolo provincial viola el estándar internacional al no exigir agotamiento previo "
                    "de medios no violentos. El actor se retiraba: no había amenaza que justificara el uso "
                    "de la fuerza."
                ),
                "jurisprudencia": (
                    "CSJN, 'Asociación de Derechos Civiles c/ CABA', Fallos 330:3609. "
                    "Corte IDH, 'Zambrano Vélez c/ Ecuador', 2007. "
                    "CEDHA Informe Especial sobre Manifestantes, 2022."
                ),
            },
            {
                "nombre": "Provincia de Entre Ríos",
                "rol": "demandado",
                "pretension": "Rechazo de la demanda. Los policías actuaron dentro del protocolo provincial vigente. La manifestación había derivado en desórdenes.",
                "fundamentos": (
                    "El protocolo provincial fue aprobado por el Poder Ejecutivo en ejercicio de sus "
                    "competencias de seguridad interior. Los efectivos actuaron ante un riesgo real de "
                    "desorden público. La Provincia no puede ser responsabilizada por el cumplimiento "
                    "de protocolos legalmente aprobados (eximente por norma). La distancia de 8 metros "
                    "fue necesaria dadas las circunstancias operativas."
                ),
                "jurisprudencia": (
                    "CSJN, 'Bertorotta c/ Provincia de Buenos Aires', 2018. "
                    "SCJER, 'Lezcano c/ Provincia de Entre Ríos', 2021."
                ),
            },
        ],
        "cuestiones": [
            "¿El cumplimiento de un protocolo provincial de uso de la fuerza exime al Estado de responsabilidad cuando el protocolo viola estándares internacionales de derechos humanos incorporados por el art. 75 inc. 22 CN?",
            "¿El uso de balas de goma a menor distancia de la reglamentaria constituye actividad ilícita del Estado generadora de responsabilidad por fuerza mayor o hay lugar para la responsabilidad por riesgo de la cosa?",
            "¿Puede el juez declarar la inconstitucionalidad del protocolo de uso de la fuerza en el marco de una acción de daños o requiere una acción declarativa autónoma?",
        ],
    },
    {
        "titulo": "Caso 11 — Martínez c/ Netflix Argentina (Contrato digital / Cláusulas abusivas)",
        "filename": "Caso11_Martinez_c_Netflix.docx",
        "caratula": "Martínez, Javier Ernesto c/ Netflix International B.V. s/ nulidad de cláusulas y daños (consumidor digital)",
        "expediente": "CCC-17723/2024",
        "rama": "civil_comercial",
        "tipo_proceso": "ordinario",
        "tribunal": "Juzgado Nacional en lo Civil y Comercial Federal N° 5",
        "nivel": "N2",
        "hechos": (
            "El actor contrató Netflix en 2019 con plan HD a $350/mes. En 2023 Netflix modificó "
            "unilateralmente las condiciones: aumentó el precio a $4.200/mes y eliminó la posibilidad "
            "de compartir contraseñas (que estaba expresamente permitida en las condiciones originales). "
            "La modificación fue notificada por email 30 días antes. El actor impugna: (1) el aumento "
            "de 1.100% acumulado como cláusula abusiva no negociada; (2) la modificación unilateral "
            "del objeto del servicio (ya no puede usarlo en el hogar de su padre anciano como antes); "
            "(3) la jurisdicción pactada (Países Bajos) como cláusula limitativa del acceso a la justicia. "
            "Netflix alega que los términos y condiciones prevén expresamente su facultad de modificación."
        ),
        "partes": [
            {
                "nombre": "Martínez, Javier Ernesto",
                "rol": "actor",
                "pretension": "Nulidad de las cláusulas de modificación unilateral y de jurisdicción extranjera. Restitución de los sobrecargos cobrados. Daño punitivo por práctica abusiva masiva.",
                "fundamentos": (
                    "LDC arts. 37 y 38: nulidad de cláusulas abusivas en contratos de adhesión. "
                    "Art. 1119 CCCN: abusividad por desnivele significativo. "
                    "Art. 1651 CCCN: acuerdo de jurisdicción extranjera en contratos de consumo es nulo. "
                    "Res. 53/2003 SCJ: los aumentos desproporcionados en contratos de tracto sucesivo "
                    "requieren causa justificada. El cambio de la política de contraseñas altera "
                    "sustancialmente el objeto del servicio contratado."
                ),
                "jurisprudencia": (
                    "CNCA Sala I, 'Defensa del Consumidor c/ Movistar', 2022. "
                    "CNCom. Sala D, 'Asociación Consumidores Libres c/ Direct TV', 2021. "
                    "CJUE, 'VB c/ Turkiye Garanti Bankasi AS', C-26/13."
                ),
            },
            {
                "nombre": "Netflix International B.V.",
                "rol": "demandado",
                "pretension": "Incompetencia del tribunal argentino (jurisdicción pactada: Países Bajos). Subsidiariamente, validez de las modificaciones notificadas con 30 días de antelación.",
                "fundamentos": (
                    "Los Términos de Uso prevén expresamente la facultad de modificar el servicio "
                    "y el precio con notificación. El usuario aceptó libremente estos términos al "
                    "suscribir. La cláusula de jurisdicción fue aceptada voluntariamente. "
                    "Los aumentos responden al contexto inflacionario argentino: no son abusivos "
                    "sino necesarios para la sostenibilidad del servicio. La restricción de "
                    "contraseñas es una medida legítima de protección de derechos de autor."
                ),
                "jurisprudencia": (
                    "CNCA, 'Apple c/ DNCI', 2019. TS España, 'González c/ Amazon', 2022."
                ),
            },
        ],
        "cuestiones": [
            "¿La cláusula de jurisdicción extranjera en contratos de consumo digital celebrados con usuarios argentinos es nula por aplicación del art. 1651 CCCN, aunque el usuario la haya aceptado expresamente?",
            "¿La facultad de modificación unilateral de precios prevista en términos y condiciones de contratos de adhesión de servicios digitales es una cláusula abusiva nula conforme el art. 37 LDC?",
            "¿El cambio en la política de uso de contraseñas (restricción de uso familiar) altera el objeto del contrato original y habilita la rescisión sin cargo o la indemnización al consumidor?",
        ],
    },
    {
        "titulo": "Caso 12 — Suárez c/ Fisco (Ganancias / Indemnización por despido)",
        "filename": "Caso12_Suarez_c_Fisco.docx",
        "caratula": "Suárez, Patricia Noemí c/ AFIP - DGI s/ repetición de tributos (impuesto a las ganancias sobre indemnización)",
        "expediente": "TFN-8891/2024",
        "rama": "administrativo",
        "tipo_proceso": "contencioso_administrativo",
        "tribunal": "Cámara Nacional de Apelaciones en lo Contencioso Administrativo Federal, Sala IV",
        "nivel": "N2",
        "hechos": (
            "La actora fue despedida sin causa en 2022 por su empleadora, una multinacional. Recibió "
            "una indemnización total de $18.400.000 (art. 245 LCT + integración mes de despido + SAC + "
            "preaviso). La empleadora retuvo $3.200.000 en concepto de impuesto a las ganancias sobre "
            "la indemnización por preaviso e integración del mes de despido ($4.100.000 sobre los cuales "
            "se calculó el impuesto). La actora reclama la repetición, argumentando que todos los rubros "
            "indemnizatorios por despido están exentos del impuesto a las ganancias conforme el fallo "
            "'Cuevas' de la CSJN y la Ley 27.617. AFIP sostiene que la Ley 27.617 sólo eximió el "
            "rubro 'indemnización por antigüedad' (art. 245 LCT) pero no el preaviso ni la integración."
        ),
        "partes": [
            {
                "nombre": "Suárez, Patricia Noemí",
                "rol": "actora",
                "pretension": "Repetición de $3.200.000 retenidos más intereses. Los rubros de preaviso e integración son indemnizatorios y están exentos.",
                "fundamentos": (
                    "CSJN, 'Cuevas c/ AFIP' (2016): toda indemnización por despido que resarza un daño "
                    "no está alcanzada por ganancias. Ley 27.617 art. 1: exención de 'toda suma que se "
                    "perciba como consecuencia de la extinción de la relación laboral'. El preaviso omitido "
                    "se convierte en indemnización sustitutiva, perdiendo carácter remunerativo. "
                    "El principio de la norma más favorable al trabajador (art. 9 LCT)."
                ),
                "jurisprudencia": (
                    "CSJN, 'Cuevas, Máximo Alberto c/ AFIP', Fallos 339:322. "
                    "TFN Sala D, 'Gómez c/ AFIP', 2023. CNCAF Sala II, 'Benavidez c/ AFIP', 2022."
                ),
            },
            {
                "nombre": "AFIP - DGI",
                "rol": "demandado",
                "pretension": "Rechazo de la repetición. El preaviso es renta del trabajo y no indemnización, por lo que está alcanzado por el impuesto a las ganancias conforme la Ley de IG.",
                "fundamentos": (
                    "Ley 27.617: eximió expresamente el art. 245 LCT (indemnización por antigüedad) pero "
                    "no los demás rubros. El preaviso sustituye la obligación de trabajar: es renta gravada "
                    "conforme art. 79 inc. b) de la Ley de IG. La Circular AFIP 4/2016 (post-'Cuevas') "
                    "mantiene la gravabilidad del preaviso. El principio de legalidad tributaria impide "
                    "la interpretación extensiva de las exenciones."
                ),
                "jurisprudencia": (
                    "CNCAF Sala V, 'Pérez c/ AFIP', 2021. TFN Sala B, 'Rodríguez c/ DGI', 2022."
                ),
            },
        ],
        "cuestiones": [
            "¿Los rubros de preaviso omitido e integración del mes de despido quedan comprendidos en la exención de ganancias establecida por la Ley 27.617 y la doctrina 'Cuevas' de la CSJN?",
            "¿El carácter 'indemnizatorio' de un rubro laboral se determina por su función resarcitoria o por su denominación en la LCT?",
            "¿El principio de legalidad tributaria y la prohibición de analogía en materia de exenciones impide la extensión de la exención a rubros no mencionados expresamente en la Ley 27.617?",
        ],
    },
    {
        "titulo": "Caso 13 — García c/ Obra Social (Fertilización asistida / Ovodonación)",
        "filename": "Caso13_Garcia_c_ObraSocial.docx",
        "caratula": "García, Mónica Alejandra c/ OSDE s/ amparo de salud (fertilización asistida heteróloga)",
        "expediente": "AS-5512/2024",
        "rama": "administrativo",
        "tipo_proceso": "amparo",
        "tribunal": "Juzgado Nacional en lo Civil N° 88 (con competencia en familia)",
        "nivel": "N2",
        "hechos": (
            "La actora tiene 42 años y un diagnóstico de fallo ovárico prematuro. Su única opción "
            "reproductiva es la ovodonación (fertilización in vitro con óvulos de donante anónima). "
            "OSDE rechazó la cobertura del cuarto intento de FIV (los tres anteriores fracasaron) "
            "alegando que la Ley 26.862 cubre hasta tres intentos con diagnóstico de infertilidad, "
            "y que la ovodonación es una técnica de mayor complejidad no especificada en el vademécum. "
            "El Decreto Reglamentario 956/2013 lista la FIV homóloga, no expresamente la heteróloga. "
            "La actora solicita el 4° intento de ovodonación. No tiene capacidad económica para "
            "costear el tratamiento privadamente ($2.800.000 por ciclo)."
        ),
        "partes": [
            {
                "nombre": "García, Mónica Alejandra",
                "rol": "actora",
                "pretension": "Cobertura del cuarto intento de FIV con ovodonación al 100% conforme Ley 26.862.",
                "fundamentos": (
                    "Ley 26.862 art. 8: cobertura de técnicas de reproducción medicamente asistida de "
                    "alta y baja complejidad 'sin límite de intentos'. El Decreto 956/2013 no puede "
                    "restringir lo que la ley garantiza. La ovodonación es una técnica de alta complejidad "
                    "comprendida en el espíritu de la ley. Derecho a formar una familia (art. 14 CN, CADH "
                    "art. 17). El fallo ovárico prematuro es una enfermedad, no una elección."
                ),
                "jurisprudencia": (
                    "CSJN, 'R.A.L. c/ OSDE', 2021. CNCiv. Sala J, 'B.C.G. c/ Swiss Medical', 2022. "
                    "Comité DESC, Observación General N° 14 (derecho a la salud)."
                ),
            },
            {
                "nombre": "OSDE",
                "rol": "demandado",
                "pretension": "Rechazo del amparo. La cobertura de ovodonación no está contemplada expresamente en el decreto reglamentario. El límite de tres intentos es razonable.",
                "fundamentos": (
                    "El Decreto 956/2013 lista taxativamente las técnicas cubiertas: la ovodonación "
                    "requiere tercero donante y tiene implicancias éticas, genéticas y legales distintas "
                    "a la FIV homóloga. La Ley 26.862 dice 'sin límite de intentos' respecto de las "
                    "técnicas indicadas, no de cualquier procedimiento. La sustentabilidad del sistema "
                    "de salud impone límites razonables. El cuarto intento en un perfil de 42 años "
                    "tiene probabilidad de éxito inferior al 15%."
                ),
                "jurisprudencia": (
                    "CNCiv. Sala K, 'P.R.M. c/ OSDE', 2021. CNCAF Sala I, 'Defensoría del Pueblo c/ "
                    "IOSFA', 2020."
                ),
            },
        ],
        "cuestiones": [
            "¿La ovodonación (FIV heteróloga con óvulos de donante) está comprendida dentro de las técnicas de reproducción medicamente asistida cubiertas por la Ley 26.862, aunque no esté listada expresamente en el Decreto Reglamentario 956/2013?",
            "¿La Ley 26.862 garantiza cobertura 'sin límite de intentos' respecto de todos los procedimientos de TRHA o sólo respecto de los expresamente reglamentados?",
            "¿La edad de la paciente y la baja probabilidad estadística de éxito del tratamiento son factores que el juez puede considerar para denegar la cobertura, o vulneran el derecho a la no discriminación por edad?",
        ],
    },
    {
        "titulo": "Caso 14 — Consorcio c/ Propietario (PHorizontal / Obras sin autorización y ruidos)",
        "filename": "Caso14_Consorcio_c_Propietario.docx",
        "caratula": "Consorcio de Propietarios Edificio Torres del Plata c/ Ríos, Juan Pablo s/ cese de perturbación y daños (propiedad horizontal)",
        "expediente": "CCC-6641/2024",
        "rama": "civil_comercial",
        "tipo_proceso": "ordinario",
        "tribunal": "Juzgado Nacional en lo Civil N° 56",
        "nivel": "N2",
        "hechos": (
            "El demandado Ríos es propietario del apartamento 12B del edificio. Sin autorización "
            "del consorcio ni los planos municipales aprobados, realizó en 2022 una ampliación "
            "que implicó la demolición de una pared estructural. La pericia técnica determinó que "
            "la demolición comprometió la resistencia de dos pilares del edificio, generando riesgo "
            "estructural potencial. Además, desde 2021 el demandado desarrolla una academia de "
            "percusión en el departamento entre las 8 y las 22 hs. Las mediciones acústicas "
            "registraron 78 dB promedio en el departamento lindero 12A (el reglamento de copropiedad "
            "fija 45 dB como límite). El demandado alega que el reglamento de copropiedad es "
            "inconstitucional por restringir el uso de la propiedad privada."
        ),
        "partes": [
            {
                "nombre": "Consorcio de Propietarios Torres del Plata",
                "rol": "actor",
                "pretension": "Orden de demolición de la ampliación ilegal, restauración de la pared estructural, cese de la actividad acústica perturbadora y daños al edificio.",
                "fundamentos": (
                    "Art. 2047 CCCN: los propietarios no pueden realizar obras que afecten la estructura "
                    "del edificio. Art. 2046 CCCN: el propietario debe cumplir el reglamento de copropiedad. "
                    "Art. 2069 CCCN: el consorcio tiene legitimación activa para acciones judiciales. "
                    "Art. 14 CN: el derecho de propiedad no es absoluto y cede ante el daño a terceros. "
                    "Los límites acústicos del reglamento son vinculantes (art. 2056 CCCN)."
                ),
                "jurisprudencia": (
                    "CNCiv. Sala A, 'Consorcio Juncal c/ Fliess', 2022. "
                    "CNCiv. Sala E, 'Bercovich c/ Consorcio', 2021. "
                    "SCBA, 'Pereyra c/ Gallo', 2019."
                ),
            },
            {
                "nombre": "Ríos, Juan Pablo",
                "rol": "demandado",
                "pretension": "Rechazo de la demanda. La obra es una mejora que no compromete la estructura. El reglamento de copropiedad es inconstitucional por limitar el uso de su propiedad.",
                "fundamentos": (
                    "Art. 17 CN: la propiedad es inviolable. El reglamento de copropiedad es un contrato "
                    "de adhesión que limita derechos constitucionales más allá de lo razonable. "
                    "La pericia sobre riesgo estructural es especulativa. La actividad musical tiene "
                    "cobertura constitucional (libertad de trabajo y de enseñanza). Los límites acústicos "
                    "del reglamento son más restrictivos que las normas municipales."
                ),
                "jurisprudencia": (
                    "CSJN, 'Ercolano c/ Lanteri de Renshaw', Fallos 136:161. "
                    "CNCiv. Sala H, 'Montoya c/ Consorcio', 2020."
                ),
            },
        ],
        "cuestiones": [
            "¿Las restricciones al uso de la propiedad privada establecidas en el reglamento de copropiedad que van más allá de las normas municipales son inconstitucionales o son obligaciones contractualmente asumidas válidas?",
            "¿La verificación de riesgo estructural potencial (no actual) es suficiente para ordenar la demolición de obras ya realizadas o se requiere la acreditación de daño concreto?",
            "¿El ejercicio de una actividad laboral lícita (academia de música) en una unidad funcional puede ser limitado por el consorcio cuando supera los estándares acústicos del reglamento de copropiedad?",
        ],
    },
    {
        "titulo": "Caso 15 — Molina c/ Empresa de Gas (Explosión / Responsabilidad concurrente)",
        "filename": "Caso15_Molina_c_EmpresaGas.docx",
        "caratula": "Molina, Carlos Eduardo y otros c/ MetroGAS S.A. y Constructora Palermo S.R.L. s/ daños y perjuicios (responsabilidad concurrente)",
        "expediente": "CCC-44312/2024",
        "rama": "civil_comercial",
        "tipo_proceso": "ordinario",
        "tribunal": "Juzgado Nacional en lo Civil N° 101",
        "nivel": "N2",
        "hechos": (
            "En septiembre de 2023 una explosión de gas en un edificio en construcción de Palermo "
            "causó la muerte de dos trabajadores y lesiones graves a otros cuatro, incluyendo el "
            "actor Molina (quemaduras de 2° y 3° grado en el 40% del cuerpo). La investigación "
            "determinó: (1) una pérdida en la cañería de MetroGAS que abastecía el edificio lindera, "
            "nunca reparada pese a reclamos de dos años; (2) el jefe de obras de Constructora Palermo "
            "usó una amoladora eléctrica a 3 metros de la cañería con fuga conocida. MetroGAS alega "
            "que la empresa constructora causó la ignición. Constructora Palermo alega que MetroGAS "
            "fue negligente al no reparar la fuga. Los actores demandan a ambas solidariamente."
        ),
        "partes": [
            {
                "nombre": "Molina, Carlos Eduardo y otros",
                "rol": "actor",
                "pretension": "Responsabilidad solidaria de MetroGAS y Constructora Palermo. Daño físico, psicológico, moral y lucro cesante por incapacidad permanente del 65%.",
                "fundamentos": (
                    "Art. 1751 CCCN: responsabilidad concurrente cuando distintos hechos de diferentes "
                    "sujetos son cada uno causa adecuada del daño. Art. 1757 CCCN: responsabilidad por "
                    "actividad riesgosa (gas a presión). Ley 24.240 art. 40: responsabilidad solidaria "
                    "en la cadena de comercialización. MetroGAS: falta de servicio. Constructora: "
                    "negligencia en la prevención de riesgos laborales (Ley 19.587)."
                ),
                "jurisprudencia": (
                    "CSJN, 'Mosca c/ Buenos Aires', Fallos 330:563. CNCiv. Sala A, 'Paonessa c/ EDENOR', "
                    "2022. CNCiv. Sala C, 'Torres c/ MetroGAS', 2021."
                ),
            },
            {
                "nombre": "MetroGAS S.A. y Constructora Palermo S.R.L.",
                "rol": "demandado",
                "pretension": "MetroGAS: la fuga existía pero la causa eficiente fue el encendido de la amoladora por la constructora. Constructora: la fuga sin encendido no habría causado la explosión.",
                "fundamentos": (
                    "Art. 1726 CCCN: causalidad adecuada. La fuga de gas sin ignición no causa explosión: "
                    "la causa eficiente fue el encendido de la amoladora. La concurrencia de causas no "
                    "impone solidaridad sino contribución proporcional a la causa del daño. MetroGAS: "
                    "los reclamos de fuga no constaban en sus registros oficiales. Constructora: "
                    "el jefe de obra actuó fuera de las instrucciones de seguridad impartidas."
                ),
                "jurisprudencia": (
                    "CNCiv. Sala H, 'González c/ Gas Natural BAN', 2020. "
                    "CNCiv. Sala L, 'Fernández c/ Constructora del Sur', 2022."
                ),
            },
        ],
        "cuestiones": [
            "¿Cuando la explosión requirió causalmente tanto la fuga de gas de MetroGAS como la ignición de la constructora, se configura responsabilidad concurrente (art. 1751 CCCN) o una causalidad excluyente que libera a uno de los demandados?",
            "¿La concurrencia de causas materiales de distintos sujetos genera solidaridad pasiva frente a la víctima o sólo responsabilidad proporcional?",
            "¿La omisión de reparar la fuga de gas pese a reclamos documentados configura responsabilidad por omisión dolosa o culposa de MetroGAS?",
        ],
    },
    {
        "titulo": "Caso 16 — Herrera c/ Estado Nacional (Derecho a la identidad / Adopción)",
        "filename": "Caso16_Herrera_c_EstadoNacional.docx",
        "caratula": "Herrera, Diego Ariel c/ Estado Nacional (RENAPER) s/ acción declarativa de identidad biológica",
        "expediente": "CF-3318/2024",
        "rama": "civil_comercial",
        "tipo_proceso": "acción_declarativa",
        "tribunal": "Juzgado Nacional en lo Civil N° 12",
        "nivel": "N2",
        "hechos": (
            "El actor, de 38 años, fue adoptado a los 3 meses de vida en 1986. Inició búsqueda de "
            "su identidad biológica a través de la Comisión Nacional por el Derecho a la Identidad "
            "(CONADI). Los estudios genéticos del Banco Nacional de Datos Genéticos (BNDG) determinaron "
            "una compatibilidad del 99.97% con una mujer de 60 años (Sra. M.R.) que cotejó su ADN "
            "voluntariamente por sospechar ser madre de un bebé dado en adopción en 1986. Sin embargo, "
            "la Sra. M.R. falleció antes de poder establecer contacto. Los herederos de M.R. se niegan "
            "a reconocer al actor como hijo de la causante, en virtud de sus derechos hereditarios. "
            "El actor solicita el reconocimiento judicial de su filiación post mortem con efectos "
            "identitarios, sin pretensión hereditaria."
        ),
        "partes": [
            {
                "nombre": "Herrera, Diego Ariel",
                "rol": "actor",
                "pretension": "Declaración judicial de filiación biológica con M.R. (fallecida). Rectificación del acta de nacimiento. Sin pretensión hereditaria.",
                "fundamentos": (
                    "Art. 563 CCCN: derecho a conocer los orígenes, irrenunciable e imprescriptible. "
                    "Art. 3 CDN y art. 7 CDN: derecho del niño a conocer a sus padres. Ley 23.511 (BNDG): "
                    "la prueba genética es plena prueba de filiación. La filiación post mortem es admisible "
                    "cuando existe certeza científica (CSJN, doctrina 'Vásquez'). El derecho a la identidad "
                    "es autónomo del derecho hereditario: pueden disociarse sus efectos."
                ),
                "jurisprudencia": (
                    "CSJN, 'Vásquez Ferrá', Fallos 326:3758. CSJN, 'B., L.A. c/ F., A.R.', 2016. "
                    "Corte IDH, 'Gelman c/ Uruguay', 2011."
                ),
            },
            {
                "nombre": "Herederos de M.R.",
                "rol": "demandado",
                "pretension": "Rechazo de la filiación post mortem. Sin consentimiento de la causante, la prueba genética no puede crear estado de familia.",
                "fundamentos": (
                    "Art. 577 CCCN: la acción de reclamación de filiación prescribe. "
                    "La compatibilidad genética del 99.97% no equivale al 100%: existe margen de error. "
                    "La filiación post mortem sin consentimiento de la causante viola su derecho a la "
                    "privacidad y su autonomía (art. 19 CN). Los derechos hereditarios de los herederos "
                    "tienen protección constitucional. La CONADI no tiene facultades para sustituir el "
                    "proceso judicial con garantías."
                ),
                "jurisprudencia": (
                    "CNAT 'C.A.A. c/ Sucesión B.', 2019. TSJ CABA, 'G.A. c/ H.R.', 2022."
                ),
            },
        ],
        "cuestiones": [
            "¿La acción de reclamación de estado de filiación biológica post mortem, fundada en prueba genética con 99.97% de compatibilidad, es prescriptible conforme el art. 577 CCCN o está amparada por la imprescriptibilidad del derecho a la identidad?",
            "¿El derecho a la identidad biológica puede disociarse de los efectos hereditarios de la filiación, de modo que el juez declare la filiación sin abrir la sucesión?",
            "¿Puede establecerse judicialmente la filiación post mortem sin el consentimiento de la causante cuando existe prueba genética concluyente y las circunstancias del caso (adopción 1986) hacen imposible ese consentimiento?",
        ],
    },
    {
        "titulo": "Caso 17 — González c/ Empleadora (Despido / Discriminación por VIH)",
        "filename": "Caso17_Gonzalez_c_Empleadora.docx",
        "caratula": "González, Ernesto Daniel c/ Logística Sur S.A. s/ despido discriminatorio (portador de VIH)",
        "expediente": "CNT-27719/2024",
        "rama": "laboral",
        "tipo_proceso": "laboral_ordinario",
        "tribunal": "Juzgado Nacional del Trabajo N° 3",
        "nivel": "N2",
        "hechos": (
            "El actor, operario logístico, comunicó a la empresa en 2022 su diagnóstico de VIH positivo "
            "para solicitar licencias para tratamiento. En diciembre de 2023, 18 meses después, fue "
            "despedido 'por razones de reorganización'. La empresa sólo despidió a González entre los "
            "12 empleados del sector. Ningún otro empleado fue desvinculado. El actor pide la nulidad "
            "del despido y la reinstalación. La empresa alega causa objetiva económica y niega cualquier "
            "relación con la condición de salud. El actor no tiene constancias escritas de que su "
            "diagnóstico fuera conocido por la gerencia, pero ofrece testimonios de compañeros."
        ),
        "partes": [
            {
                "nombre": "González, Ernesto Daniel",
                "rol": "actor",
                "pretension": "Nulidad del despido, reinstalación y reparación integral del daño moral por despido discriminatorio.",
                "fundamentos": (
                    "Ley 23.798 (SIDA) art. 1: prohibición de discriminar por condición de portador de VIH. "
                    "Ley 23.592 (antidiscriminatoria): el despido discriminatorio es nulo y da derecho a "
                    "la reinstalación. CSJN, 'Álvarez c/ Cencosud': la Ley 23.592 aplica en el ámbito "
                    "laboral privado. Inversión de la carga probatoria: ante indicios suficientes de "
                    "discriminación, la empresa debe probar causa legítima autónoma. Los indicios: "
                    "conocimiento del diagnóstico, unicidad del despido, ausencia de causa acreditada."
                ),
                "jurisprudencia": (
                    "CSJN, 'Álvarez, Maximiliano c/ Cencosud', Fallos 333:2306. "
                    "CNAT Sala II, 'Torres c/ Supermercados Mayoristas', 2022. "
                    "Comité DESC, Observación General N° 20 (no discriminación)."
                ),
            },
            {
                "nombre": "Logística Sur S.A.",
                "rol": "demandado",
                "pretension": "Validez del despido por causa objetiva. Rechazo de la nulidad: la Ley 23.592 requiere prueba directa de la discriminación.",
                "fundamentos": (
                    "El despido se fundó en una reestructuración del sector documentada. La empresa "
                    "no sabía del diagnóstico de VIH: la comunicación fue al médico laboral, no a la "
                    "gerencia. La proximidad temporal (18 meses) no es indicio suficiente. "
                    "La inversión de la carga de la prueba viola el principio de inocencia de la empresa. "
                    "La reinstalación forzosa en una empresa privada viola la libertad de contratación."
                ),
                "jurisprudencia": (
                    "CNAT Sala X, 'Salomone c/ Carrefour', 2021. CSJN, 'Pellicori c/ Colegio Público', "
                    "Fallos 334:1387."
                ),
            },
        ],
        "cuestiones": [
            "¿La comunicación del diagnóstico de VIH al servicio médico laboral de la empresa puede imputarse como 'conocimiento' a la empleadora a efectos de la Ley 23.798?",
            "¿La desvinculación individual de un empleado portador de VIH, 18 meses después de comunicar su diagnóstico, constituye indicio suficiente para invertir la carga de la prueba en materia de discriminación?",
            "¿La reinstalación forzosa prevista por la Ley 23.592 es constitucionalmente válida en el ámbito de una empresa privada, frente a la libertad de contratación del empleador?",
        ],
    },
    {
        "titulo": "Caso 18 — Asociación Civil c/ ANMAT (Alimentos / Etiquetado frontal)",
        "filename": "Caso18_AsocCivil_c_ANMAT.docx",
        "caratula": "Asociación por los Derechos de los Consumidores c/ ANMAT y Estado Nacional s/ acción declarativa de inconstitucionalidad (etiquetado frontal)",
        "expediente": "CA-7712/2024",
        "rama": "administrativo",
        "tipo_proceso": "acción_declarativa",
        "tribunal": "Cámara Nacional de Apelaciones en lo Contencioso Administrativo Federal, Sala II",
        "nivel": "N2",
        "hechos": (
            "La Ley 27.642 (etiquetado frontal de alimentos) fue reglamentada por el Decreto 151/2022, "
            "que estableció un período de excepción de 36 meses para productos con bajo volumen de "
            "ventas (menos de 500.000 unidades anuales). Tres asociaciones de productores de "
            "alimentos ultraprocesados obtuvieron medidas cautelares individuales que en la práctica "
            "paralizaron la vigencia del sistema de octógonos para el 65% de los productos del mercado. "
            "La Asociación actora demanda la nulidad de las cautelares y la plena vigencia de la Ley "
            "27.642 para todos los productos desde la fecha de su entrada en vigor. El Estado Nacional "
            "no ejecutó las multas correspondientes a los productores que incumplieron el etiquetado."
        ),
        "partes": [
            {
                "nombre": "Asociación por los Derechos de los Consumidores",
                "rol": "actor",
                "pretension": "Declaración de vigencia plena e inmediata de la Ley 27.642. Nulidad de las cautelares obtenidas por las cámaras empresariales. Orden de ejecución de multas.",
                "fundamentos": (
                    "Art. 42 CN: derecho a la información de los consumidores. Principio de no regresión "
                    "en derechos del consumidor. Las excepciones reglamentarias no pueden desnaturalizar "
                    "el objeto de la ley. Las cautelares individuales de empresas no pueden oponer efectos "
                    "erga omnes que paralizan una ley de orden público. OPS/OMS: evidencia de impacto "
                    "del etiquetado en la reducción de enfermedades no transmisibles."
                ),
                "jurisprudencia": (
                    "CSJN, 'Asociación de Magistrados c/ Estado Nacional', 2022. "
                    "CNCA Sala I, 'Asociación Protejer c/ ANMAT', 2023. "
                    "Corte Suprema de Chile, 'ACHIGA c/ MINSAL', 2020."
                ),
            },
            {
                "nombre": "ANMAT y Estado Nacional",
                "rol": "demandado",
                "pretension": "Las excepciones reglamentarias son válidas. Las medidas cautelares de terceros son oponibles erga omnes en materia regulatoria.",
                "fundamentos": (
                    "El Decreto 151/2022 ejerció legítimamente la potestad reglamentaria del art. 99 "
                    "inc. 2 CN. La progresividad de la implementación es una decisión de política "
                    "sanitaria razonable. Las cautelares fueron otorgadas por jueces competentes: "
                    "ANMAT debe acatarlas. La acción declarativa de la asociación es inadmisible: "
                    "no tiene legitimación para impugnar cautelares otorgadas en otros expedientes."
                ),
                "jurisprudencia": (
                    "CSJN, 'Droguería del Sud c/ ANMAT', 2018. CNCAF Sala V, 'Cámara Argentina de "
                    "Anunciantes c/ ENACOM', 2022."
                ),
            },
        ],
        "cuestiones": [
            "¿Una asociación de consumidores tiene legitimación para impugnar medidas cautelares obtenidas por cámaras empresariales en expedientes en los que no es parte?",
            "¿El principio de progresividad en la implementación de derechos del consumidor admite excepciones reglamentarias que difieran la vigencia de una ley de orden público por 36 meses?",
            "¿Las medidas cautelares obtenidas individualmente por empresas en causas separadas pueden producir efectos erga omnes sobre la implementación de una ley de salud pública?",
        ],
    },
    {
        "titulo": "Caso 19 — Torres c/ Banco (Fideicomiso de garantía / Ejecución anticipada)",
        "filename": "Caso19_Torres_c_Banco.docx",
        "caratula": "Torres, Raúl Marcelo c/ Banco Ciudad de Buenos Aires y Confianza S.A. Fiduciaria s/ nulidad de ejecución fiduciaria",
        "expediente": "CCC-9912/2024",
        "rama": "civil_comercial",
        "tipo_proceso": "ordinario",
        "tribunal": "Juzgado Nacional en lo Civil N° 62",
        "nivel": "N2",
        "hechos": (
            "El actor contrató un crédito hipotecario con garantía de fideicomiso de garantía "
            "sobre su única vivienda. Ante la mora de dos cuotas (covid-19, período de emergencia "
            "2020), el banco notificó al fideicomisario (Confianza S.A.) quien inició la venta "
            "extrajudicial del inmueble. El actor interpuso medida de no innovar y luego nulidad, "
            "alegando: (1) el fideicomiso de garantía extrajudicial viola el derecho de defensa "
            "del deudor; (2) la vivienda única es inembargable e inejecutable; (3) las cuotas en "
            "mora correspondían al período de emergencia económica (DNU 319/2020 que suspendió "
            "ejecuciones hipotecarias). El banco alega que el DNU 319/2020 venció y que el "
            "fideicomiso de garantía es un contrato válido aceptado libremente."
        ),
        "partes": [
            {
                "nombre": "Torres, Raúl Marcelo",
                "rol": "actor",
                "pretension": "Nulidad de la ejecución extrajudicial. Declaración de inembargabilidad de la vivienda única. Daños por la turbación de la posesión.",
                "fundamentos": (
                    "DNU 319/2020 y Ley 27.541: suspensión de ejecuciones hipotecarias durante la emergencia. "
                    "Art. 244 CCCN: afectación de la vivienda única como bien de familia, inembargable. "
                    "El fideicomiso de garantía no puede sortear las protecciones del ejecutado: "
                    "fraude a la ley (art. 12 CCCN). Arts. 17 y 18 CN: derecho de propiedad y "
                    "defensa en juicio exigen control judicial previo a la ejecución. "
                    "El deudor hipotecario es consumidor financiero (LDC)."
                ),
                "jurisprudencia": (
                    "CSJN, 'Fayt c/ Estado Nacional', 2015. CNCom. Sala A, 'González c/ HSBC Fiduciaria', "
                    "2022. CNCiv. Sala G, 'Aguilar c/ Banco Supervielle', 2021."
                ),
            },
            {
                "nombre": "Banco Ciudad de Buenos Aires y Confianza S.A. Fiduciaria",
                "rol": "demandado",
                "pretension": "Validez de la ejecución fiduciaria extrajudicial. El bien no fue afectado como bien de familia conforme el registro. El DNU 319/2020 venció en 2022.",
                "fundamentos": (
                    "Arts. 1680 y ss. CCCN: el fideicomiso de garantía es un contrato válido con "
                    "ejecución extrajudicial pactada. El bien de familia requiere inscripción registral "
                    "previa (art. 244 CCCN): el actor nunca lo inscribió. El DNU 319/2020 tenía "
                    "vigencia temporal y su vencimiento restableció la facultad de ejecución. "
                    "El actor incurrió en mora dolosa con plena capacidad de pago."
                ),
                "jurisprudencia": (
                    "CNCom. Sala C, 'Pereyra c/ Citibank Fiduciaria', 2021. "
                    "PTN Dictamen 312:89."
                ),
            },
        ],
        "cuestiones": [
            "¿El fideicomiso de garantía con ejecución extrajudicial pactada viola el derecho de defensa del deudor hipotecario (art. 18 CN) cuando se ejecuta la única vivienda del deudor?",
            "¿La protección de la vivienda única del art. 244 CCCN requiere inscripción registral formal previa o es oponible como derecho sustantivo aunque no esté registrada?",
            "¿Las cuotas devengadas durante la vigencia del DNU 319/2020 pueden ser incluidas como causa de mora que habilita la ejecución posterior al vencimiento de la emergencia?",
        ],
    },
    {
        "titulo": "Caso 20 — Villalba c/ Estado Nacional (Derecho al olvido / Internet)",
        "filename": "Caso20_Villalba_c_EstadoNacional.docx",
        "caratula": "Villalba, Sofía Belén c/ Google LLC y Grupo Clarín S.A. s/ hábeas data (derecho al olvido digital)",
        "expediente": "HD-4421/2024",
        "rama": "civil_comercial",
        "tipo_proceso": "hábeas_data",
        "tribunal": "Juzgado Nacional en lo Civil y Comercial Federal N° 2",
        "nivel": "N2",
        "hechos": (
            "La actora fue procesada en 2007 por un delito de estafa menor. En 2009 fue sobreseída "
            "definitivamente. Las noticias periodísticas de su procesamiento permanecen indexadas "
            "en Google y publicadas en el archivo digital de Clarín. En 2024 la actora, que trabaja "
            "como docente, sufre perjuicios laborales porque al buscarse su nombre en Google "
            "aparecen las noticias del procesamiento sin la noticia del sobreseimiento. Solicitó "
            "a Google la desindexación y a Clarín la rectificación o baja del contenido. Ambas "
            "empresas rechazaron el pedido: Google invoca libertad de acceso a la información; "
            "Clarín invoca la libertad de prensa y el interés histórico del archivo periodístico."
        ),
        "partes": [
            {
                "nombre": "Villalba, Sofía Belén",
                "rol": "actora",
                "pretension": "Desindexación de Google de los resultados vinculados al procesamiento. Rectificación o actualización por parte de Clarín para incluir el sobreseimiento.",
                "fundamentos": (
                    "Ley 25.326 arts. 16 y 33: derecho de rectificación, actualización y supresión de datos "
                    "inexactos, desactualizados o que afecten la dignidad. El sobreseimiento definitivo "
                    "cancela la información de procesamiento: el dato es hoy 'inexacto' por incompleto. "
                    "CSJN, 'Rodríguez c/ Google' (2014): el buscador es responsable si conociendo el "
                    "carácter dañoso del enlace no lo elimina. Derecho al olvido europeo (CJUE, 'Google "
                    "c/ AEPD', C-131/12). Art. 19 CN y CADH art. 11: derecho a la privacidad."
                ),
                "jurisprudencia": (
                    "CSJN, 'Rodríguez, María Belén c/ Google Inc.', Fallos 337:1174. "
                    "CJUE, 'Google Spain c/ AEPD', C-131/12 (2014). "
                    "CNCom. Sala E, 'Gimbutas c/ Yahoo!', 2021."
                ),
            },
            {
                "nombre": "Google LLC y Grupo Clarín S.A.",
                "rol": "demandado",
                "pretension": "Rechazo del hábeas data. La información sobre procesos judiciales es de interés público y está amparada por la libertad de prensa y el acceso a la información.",
                "fundamentos": (
                    "La libertad de prensa (art. 14 CN, art. 13 CADH) ampara el archivo periodístico "
                    "histórico. Los procesos judiciales son actos públicos que no pueden ser 'borrados'. "
                    "Google: es un intermediario técnico, no el autor del contenido; la responsabilidad "
                    "recae en Clarín. La CSJN en 'Rodríguez' exige notificación fehaciente de daño "
                    "manifiesto para responsabilizar al buscador: la actora no lo acreditó fehacientemente. "
                    "El 'derecho al olvido' no existe como tal en el derecho argentino."
                ),
                "jurisprudencia": (
                    "CSJN, 'Rodríguez c/ Google', Fallos 337:1174. CSJN, 'Belén Rodríguez c/ Yahoo!', "
                    "2015. TS España, 'B.B.B. c/ Google', 2021."
                ),
            },
        ],
        "cuestiones": [
            "¿La información sobre un procesamiento penal que fue sobreseído definitivamente 15 años atrás constituye un 'dato inexacto' o 'desactualizado' que habilita la acción de hábeas data correctivo conforme la Ley 25.326?",
            "¿Existe en el derecho argentino un 'derecho al olvido digital' que permita exigir la desindexación de información periodística verdadera pero dañosa por obsolescencia, y cuál es su fundamento normativo?",
            "¿Google LLC responde como responsable del tratamiento de datos al indexar noticias de medios periodísticos, o sólo actúa como intermediario técnico sin responsabilidad autónoma?",
        ],
    },
]


# ============================================================
# FUNCIÓN: CREAR WORD
# ============================================================
def crear_word(caso, filepath):
    doc = Document()

    # Título principal
    titulo = doc.add_heading(caso['titulo'], level=1)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Datos del caso
    doc.add_heading('DATOS DEL CASO', level=2)
    tabla = doc.add_table(rows=6, cols=2)
    tabla.style = 'Table Grid'
    datos = [
        ('Carátula', caso['caratula']),
        ('Expediente', caso['expediente']),
        ('Tribunal', caso['tribunal']),
        ('Rama', caso['rama']),
        ('Proceso', caso['tipo_proceso']),
        ('Nivel', caso['nivel']),
    ]
    for i, (k, v) in enumerate(datos):
        tabla.rows[i].cells[0].text = k
        tabla.rows[i].cells[1].text = v
        tabla.rows[i].cells[0].paragraphs[0].runs[0].bold = True
    doc.add_paragraph()

    # Hechos
    doc.add_heading('HECHOS PROBADOS', level=2)
    doc.add_paragraph(caso['hechos'])
    doc.add_paragraph()

    # Partes
    doc.add_heading('POSICIONES DE LAS PARTES', level=2)
    for parte in caso['partes']:
        doc.add_heading(parte['nombre'] + f" ({parte['rol'].upper()})", level=3)
        tabla_p = doc.add_table(rows=3, cols=2)
        tabla_p.style = 'Table Grid'
        filas = [
            ('Pretensión', parte['pretension']),
            ('Fundamentos', parte['fundamentos']),
            ('Jurisprudencia', parte['jurisprudencia']),
        ]
        for i, (k, v) in enumerate(filas):
            tabla_p.rows[i].cells[0].text = k
            tabla_p.rows[i].cells[1].text = v
            tabla_p.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        doc.add_paragraph()

    # Cuestiones
    doc.add_heading('CUESTIONES A RESOLVER', level=2)
    for i, c in enumerate(caso['cuestiones'], 1):
        doc.add_paragraph(f"{i}. {c}", style='List Number')
    doc.add_paragraph()

    # Footer
    fp = doc.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run('Generado por JUSTICIA ARGENTINA — Sistema de Resolución Asistida')
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.save(filepath)
    print(f'  Guardado: {os.path.basename(filepath)}')


# ============================================================
# PASO 1: CREAR LOS 20 WORD EN casos3/
# ============================================================
print('=' * 60)
print('PASO 1: Creando 20 archivos Word en ~/Downloads/casos3/')
print('=' * 60)
for caso in CASOS:
    filepath = os.path.join(CASOS3_DIR, caso['filename'])
    crear_word(caso, filepath)
print(f'\n20 archivos Word guardados en: {CASOS3_DIR}')


# ============================================================
# PASO 2: CORRER EL MOTOR CONTRA LOS 20 CASOS
# ============================================================
print()
print('=' * 60)
print('PASO 2: Ejecutando motor de sentencias...')
print('=' * 60)

from generador.motor_sentencias import GeneradorSentencias, CasoInput
from generador.clasificador import clasificar_caso

def extraer_normas(texto):
    patrones = [
        r'[Aa]rt(?:\u00edculo|iculo)?\.?\s*\d+(?:\s*bis)?',
        r'[Ll]ey\s+\d{4,6}',
        r'L\.?C\.?T\.?', r'CCCN|C\u00f3digo Civil',
        r'PIDESC|CADH|Convenio\s+\d+',
    ]
    encontradas = set()
    for p in patrones:
        for m in re.finditer(p, texto, re.IGNORECASE):
            encontradas.add(m.group(0).strip().upper()[:30])
    return encontradas

gen = GeneradorSentencias()
print(f'Modelo activo: {gen.gemini_model}')

campos_validos = set(CasoInput.__dataclass_fields__.keys())
resultados = []

for i, caso in enumerate(CASOS):
    print(f'\n[{i+1}/20] {caso["titulo"][:65]}')

    data = {
        'caratula': caso['caratula'],
        'expediente': caso['expediente'],
        'rama': caso['rama'],
        'tipo_proceso': caso['tipo_proceso'],
        'tribunal': caso['tribunal'],
        'hechos_probados': caso['hechos'],
        'partes': [
            {
                'nombre': p['nombre'],
                'rol': p['rol'],
                'pretension': p['pretension'],
                'fundamentos_juridicos': p['fundamentos'],
                'jurisprudencia_citada': p['jurisprudencia'],
            }
            for p in caso['partes']
        ],
        'cuestiones_a_resolver': caso['cuestiones'],
    }

    cl = clasificar_caso(data)
    nivel = cl['nivel']
    data['nivel_complejidad'] = nivel

    caso_filtrado = {k: v for k, v in data.items() if k in campos_validos}
    caso_obj = CasoInput(**caso_filtrado)
    sentencia = gen.generar_sentencia(caso_obj)

    texto = sentencia.texto_completo or ''
    tu = texto.upper()
    estructura = ((1 if 'VISTOS' in tu else 0) +
                  (1 if 'CONSIDERANDO' in tu else 0) +
                  (1 if 'RESUELVO' in tu else 0)) / 3
    normas = extraer_normas(texto)
    print(f'  Nivel: {nivel} | Estructura: {estructura*100:.0f}% | Normas: {len(normas)}')

    # Guardar sentencia como Word en casos3/
    sent_filename = caso['filename'].replace('.docx', '_SENTENCIA.docx')
    sent_path = os.path.join(CASOS3_DIR, sent_filename)
    doc_sent = Document()
    doc_sent.add_heading(caso['caratula'], level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc_sent.add_heading('PROYECTO DE SENTENCIA', level=2).alignment = WD_ALIGN_PARAGRAPH.CENTER
    info = doc_sent.add_paragraph()
    info.add_run(f'Nivel: {nivel} | Estructura: {estructura*100:.0f}% | Normas detectadas: {len(normas)}').italic = True
    doc_sent.add_paragraph()
    for linea in texto.split('\n'):
        p = doc_sent.add_paragraph(linea)
        if linea.strip() in ('VISTOS:', 'CONSIDERANDO:', 'RESUELVO:',
                              'VISTOS', 'CONSIDERANDO', 'RESUELVO'):
            p.runs[0].bold = True if p.runs else None
    doc_sent.save(sent_path)
    print(f'  Sentencia guardada: {sent_filename}')

    resultados.append({
        'caso': caso['titulo'],
        'nivel': nivel,
        'estructura': estructura,
        'normas_count': len(normas),
        'texto_preview': texto[:500],
    })

    if i < len(CASOS) - 1:
        time.sleep(3)

# Resumen final
print()
print('=' * 60)
print('RESUMEN FINAL — 20 CASOS')
print('=' * 60)
for r in resultados:
    print(f'{r["caso"][:55]:55s} | Struct={r["estructura"]*100:.0f}% | Normas={r["normas_count"]}')

results_path = os.path.join(CASOS3_DIR, 'resultados_20casos.json')
with open(results_path, 'w', encoding='utf-8') as f:
    json.dump(resultados, f, ensure_ascii=False, indent=2)

print(f'\nTodo guardado en: {CASOS3_DIR}')
print(f'Resultados JSON: {results_path}')
