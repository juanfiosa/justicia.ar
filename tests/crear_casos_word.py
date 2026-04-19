"""
Crea 10 archivos Word con los casos N2, los corre contra el motor,
y genera 10 casos nuevos guardados en ~/Downloads/casos2/
"""
import sys, os, json, re, time
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', 'src'))
# Scripts de testing usan la key de testing (cuota separada de producción)
os.environ['GEMINI_API_KEY'] = 'AIzaSyDwQp6HcIiqamO0KFeiexL2T560-zvw0_w'

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOWNLOADS = os.path.expanduser('~/Downloads')
CASOS1_DIR = DOWNLOADS
CASOS2_DIR = os.path.join(DOWNLOADS, 'casos2')
os.makedirs(CASOS2_DIR, exist_ok=True)

# ============================================================
# DEFINICION DE LOS 10 CASOS N2
# ============================================================
CASOS = [
    {
        "titulo": "Caso 1 — Flores c/ Rappi (Plataformas digitales)",
        "filename": "Caso1_Flores_c_Rappi.docx",
        "caratula": "Flores, Diego Ezequiel c/ Rappi Argentina S.R.L. s/ despido",
        "expediente": "SD-14832/2024",
        "rama": "Laboral",
        "tipo_proceso": "Ordinario",
        "tribunal": "Juzgado Nacional del Trabajo N° 12",
        "nivel": "N2 — Difícil",
        "hechos": (
            "El actor prestó servicios de delivery para Rappi Argentina S.R.L. durante 3 años mediante la "
            "aplicación móvil de la empresa. Utilizaba su propia bicicleta y teléfono celular. Tenía libertad "
            "de conectarse cuando quisiera, pero el algoritmo de la plataforma lo penalizaba con menor "
            "asignación de pedidos si su tasa de aceptación era inferior al 80% o si cancelaba pedidos ya "
            "tomados. Fue bloqueado de la aplicación sin notificación previa en agosto de 2024. Rappi "
            "facturaba los servicios como si el actor fuera un contratista independiente bajo la figura de "
            "'socio repartidor'. El actor nunca firmó contrato de trabajo pero recibió capacitaciones de la "
            "empresa, usaba uniforme con logo de Rappi y seguía instrucciones de ruta del sistema."
        ),
        "partes": [
            {
                "nombre": "Flores, Diego Ezequiel",
                "rol": "Actor",
                "pretension": (
                    "Reconocimiento de la relación de dependencia laboral y pago de indemnización por despido "
                    "incausado, SAC, vacaciones proporcionales y diferencias salariales por todo el período trabajado."
                ),
                "fundamentos": (
                    "Art. 23 LCT (presunción de relación laboral). Art. 21 LCT (definición de contrato de trabajo). "
                    "Art. 14 bis CN (protección del trabajador). El control algorítmico constituye subordinación técnica "
                    "y jurídica encubierta. La Ley 27.742 no elimina la presunción del art. 23 LCT."
                ),
                "jurisprudencia": (
                    "CNAT Sala II, 'Vdefer c/ Glovo', 2023. CNAT Sala VII, 'Mantovani c/ PedidosYa', 2022. "
                    "Tribunal Supremo de España, 'Ley Rider', 2021."
                ),
            },
            {
                "nombre": "Rappi Argentina S.R.L.",
                "rol": "Demandado",
                "pretension": (
                    "Rechazo total de la demanda. Los repartidores son trabajadores independientes que usan "
                    "la plataforma como herramienta tecnológica."
                ),
                "fundamentos": (
                    "El actor era libre de conectarse o no, sin horarios fijos ni exclusividad. No había ajenidad "
                    "en la organización. Ley 27.742 art. 96 ter reconoce la figura del trabajador independiente "
                    "de plataformas. Ausencia de subordinación jurídica efectiva."
                ),
                "jurisprudencia": (
                    "CNAT Sala III, 'Barrios c/ PedidosYa', 2021. CNFed. Civ. y Com., 'Urbano Express c/ AFIP', 2022."
                ),
            },
        ],
        "cuestiones": [
            "¿El control algorítmico de aceptación de pedidos configura subordinación jurídica suficiente para presumir relación de dependencia bajo el art. 23 LCT?",
            "¿La libertad de horario es incompatible con el contrato de trabajo cuando los demás indicios apuntan a la dependencia?",
            "¿La figura del 'socio repartidor' desvirtúa la presunción del art. 23 LCT o constituye fraude laboral?",
        ],
    },
    {
        "titulo": "Caso 2 — Giménez c/ Banco Hipotecario (Teletrabajo)",
        "filename": "Caso2_Gimenez_c_BancoHipotecario.docx",
        "caratula": "Giménez, Marcela Andrea c/ Banco Hipotecario S.A. s/ acción de reinstalación y daños",
        "expediente": "CNT-5521/2024",
        "rama": "Laboral",
        "tipo_proceso": "Ordinario",
        "tribunal": "Cámara Nacional del Trabajo, Sala IV",
        "nivel": "N2 — Difícil",
        "hechos": (
            "La actora trabajó en modalidad presencial en la sucursal Palermo del Banco Hipotecario S.A. desde 2015. "
            "En marzo de 2020, ante la pandemia de COVID-19, fue migrada a teletrabajo. En septiembre de 2023 el banco "
            "notificó el retorno obligatorio a la presencialidad a partir de noviembre de 2023. La actora manifestó "
            "por escrito que el teletrabajo era esencial para el cuidado de su hijo con discapacidad (TEA, "
            "certificado médico adjunto) y solicitó mantener la modalidad remota. El banco rechazó la solicitud "
            "por razones de organización interna y la despidió con causa al no presentarse. La actora tiene dos "
            "hijos a cargo, uno con TEA, y demostró que el trabajo presencial le impediría cumplir sus "
            "obligaciones de cuidado sin asistencia externa."
        ),
        "partes": [
            {
                "nombre": "Giménez, Marcela Andrea",
                "rol": "Actora",
                "pretension": (
                    "Nulidad del despido y reinstalación en modalidad teletrabajo, o en su defecto indemnización "
                    "agravada por despido discriminatorio. Daño moral."
                ),
                "fundamentos": (
                    "Ley 27.555 (teletrabajo) art. 8 y art. 6: reversibilidad acordada. El retorno obligatorio sin "
                    "acuerdo viola el espíritu de la ley. Ley 26.485 en su dimensión laboral. Ley 26.378 (CDPD) "
                    "respecto del hijo con discapacidad. Art. 242 LCT: la negativa al retorno no configura injuria "
                    "si existe causa justificada documentada."
                ),
                "jurisprudencia": (
                    "CNAT Sala II, 'López c/ Telecom Argentina', 2023. Comité CEDAW, Recomendación General 19. "
                    "OIT, Convenio 156 sobre trabajadores con responsabilidades familiares."
                ),
            },
            {
                "nombre": "Banco Hipotecario S.A.",
                "rol": "Demandado",
                "pretension": (
                    "Ratificación del despido con causa. La negativa injustificada a retornar a la presencialidad "
                    "configura abandono de trabajo e injuria grave."
                ),
                "fundamentos": (
                    "Ley 27.555 art. 7: la reversibilidad requiere acuerdo, pero las condiciones que motivaron "
                    "el teletrabajo (pandemia) ya no existen. Ius variandi del empleador (art. 66 LCT): el retorno "
                    "a condiciones originales no implica alteración sustancial. El banco cumplió el plazo de preaviso."
                ),
                "jurisprudencia": (
                    "CNAT Sala VI, 'Hermida c/ HSBC Bank', 2023. SCJBA, 'Giordano c/ YPF', 2019."
                ),
            },
        ],
        "cuestiones": [
            "¿La Ley 27.555 otorga al trabajador un derecho adquirido a mantener el teletrabajo pactado pandémicamente, o el empleador puede revertirlo unilateralmente?",
            "¿Las responsabilidades de cuidado de un hijo con discapacidad constituyen 'causa justificada' suficiente para negarse al retorno presencial?",
            "¿El despido por negativa al retorno presencial en estas circunstancias configura discriminación indirecta por razón de género y discapacidad?",
        ],
    },
    {
        "titulo": "Caso 3 — Rodríguez c/ Banco Santander (Phishing bancario)",
        "filename": "Caso3_Rodriguez_c_Santander.docx",
        "caratula": "Rodríguez, Carlos Alberto c/ Banco Santander Argentina S.A. s/ daños y perjuicios",
        "expediente": "CCC-88234/2024",
        "rama": "Civil",
        "tipo_proceso": "Ordinario",
        "tribunal": "Juzgado Nacional en lo Civil y Comercial Federal N° 3",
        "nivel": "N2 — Difícil",
        "hechos": (
            "El actor recibió un llamado telefónico de quien se identificó como empleado del banco, informándole "
            "una supuesta transacción sospechosa. Mediante ingeniería social (spoofing del número oficial del banco), "
            "el interlocutor le solicitó los tokens de seguridad generados por la aplicación. El actor, creyendo "
            "hablar con el banco, proporcionó los códigos. Se realizaron transferencias por $4.800.000 a cuentas de "
            "terceros. El banco rechazó el reclamo argumentando culpa exclusiva del usuario. El actor es jubilado "
            "de 67 años sin antecedentes de incidentes previos. Pericias informáticas determinaron que el llamado "
            "utilizó técnica de spoofing del número oficial del banco."
        ),
        "partes": [
            {
                "nombre": "Rodríguez, Carlos Alberto",
                "rol": "Actor",
                "pretension": (
                    "Restitución de $4.800.000 más intereses, daño moral y daño punitivo por conducta "
                    "negligente del banco."
                ),
                "fundamentos": (
                    "Ley 24.240 art. 40: responsabilidad objetiva del proveedor. Art. 1757 CCCN (actividad riesgosa). "
                    "El banco tiene deber de seguridad reforzado frente a adultos mayores. El spoofing del número "
                    "oficial creó una apariencia de legitimidad que el consumidor no podía detectar."
                ),
                "jurisprudencia": (
                    "CNCom. Sala D, 'Fernández c/ BBVA', 2023. CNCom. Sala A, 'Pérez c/ Galicia', 2022. "
                    "CSJN, 'Mosca c/ Buenos Aires', Fallos 330:563."
                ),
            },
            {
                "nombre": "Banco Santander Argentina S.A.",
                "rol": "Demandado",
                "pretension": "Rechazo de la demanda por culpa exclusiva de la víctima al entregar voluntariamente sus credenciales.",
                "fundamentos": (
                    "Art. 1729 CCCN: culpa de la víctima como eximente. El banco cumplió todos los estándares BCRA "
                    "(Com. A 6017). El token fue generado correctamente. Las advertencias sobre no compartir "
                    "credenciales estaban disponibles en la app."
                ),
                "jurisprudencia": (
                    "CNCom. Sala B, 'Aquino c/ Citibank', 2022. BCRA Comunicación A 6017/2017."
                ),
            },
        ],
        "cuestiones": [
            "¿El banco tiene deber de seguridad frente al phishing por spoofing de su número oficial, aunque las credenciales fueran entregadas voluntariamente por el usuario?",
            "¿La entrega de tokens bajo engaño mediante suplantación del número oficial configura culpa exclusiva de la víctima o concurrencia de culpas?",
            "¿Procede el daño punitivo del art. 52 bis LDC cuando el banco rechaza sistemáticamente este tipo de reclamos?",
        ],
    },
    {
        "titulo": "Caso 4 — Sánchez Propiedades c/ Moreno (Alquileres / DNU 70/2023)",
        "filename": "Caso4_SanchezPropiedades_c_Moreno.docx",
        "caratula": "Sánchez Propiedades S.R.L. c/ Moreno, Lucia Valentina s/ cobro de alquileres y desalojo",
        "expediente": "JCC-22109/2024",
        "rama": "Civil",
        "tipo_proceso": "Ordinario",
        "tribunal": "Juzgado Civil y Comercial N° 8 de Rosario",
        "nivel": "N2 — Difícil",
        "hechos": (
            "En agosto de 2022 se celebró contrato de locación de vivienda por 3 años con actualización semestral "
            "por ICL (Índice de Contratos de Locación) según Ley 27.551. En diciembre de 2023 el DNU 70/2023 "
            "derogó el art. 14 de la Ley 27.551, eliminando el ICL como índice obligatorio. El locador notificó "
            "que a partir de enero de 2024 la actualización sería mensual por IPC, resultando en un aumento del "
            "214% en el primer semestre de 2024. La locataria impugnó el cambio unilateral y continuó pagando "
            "según ICL. Al momento de la demanda existe una diferencia de $1.450.000 entre lo reclamado y lo pagado."
        ),
        "partes": [
            {
                "nombre": "Sánchez Propiedades S.R.L.",
                "rol": "Actor",
                "pretension": (
                    "Cobro de diferencias de alquileres por $1.450.000 y, subsidiariamente, desalojo por falta de pago. "
                    "Aplicación del índice IPC desde la entrada en vigor del DNU 70/2023."
                ),
                "fundamentos": (
                    "DNU 70/2023 art. 247: derogación del art. 14 Ley 27.551 con efectos inmediatos sobre contratos "
                    "en curso. Libertad de contratación (art. 958 CCCN). La actualización por ICL implicaría una "
                    "pérdida del 60% del valor real del alquiler."
                ),
                "jurisprudencia": (
                    "Juzgado Civil 15 CABA, 'Inmobiliaria Paz c/ Díaz', enero 2024. CNCiv. Sala H, "
                    "'Construyendo SA c/ García', marzo 2024."
                ),
            },
            {
                "nombre": "Moreno, Lucia Valentina",
                "rol": "Demandado",
                "pretension": (
                    "Rechazo de la demanda. El contrato fue celebrado bajo la Ley 27.551 y sus condiciones deben "
                    "regir hasta su vencimiento natural. Inconstitucionalidad del DNU 70/2023 en materia locativa."
                ),
                "fundamentos": (
                    "Art. 3 CCCN: la ley nueva no afecta derechos adquiridos. Art. 17 CN: inviolabilidad de los "
                    "contratos. El DNU 70/2023 viola el art. 99 inc. 3 CN al regular materia de derecho común "
                    "sin urgencia que impida el trámite legislativo."
                ),
                "jurisprudencia": (
                    "CSJN, 'Cerealera San Martín', Fallos 341:1132. Cám. Civil Neuquén, 'Hidalgo c/ López', feb. 2024."
                ),
            },
        ],
        "cuestiones": [
            "¿El DNU 70/2023 se aplica a contratos de locación celebrados bajo la Ley 27.551 y en plena ejecución al momento de su dictado?",
            "¿El cambio unilateral del índice de actualización por el locador, sin acuerdo del locatario, constituye una alteración sustancial del contrato?",
            "¿El DNU 70/2023 es constitucional en cuanto al requisito de necesidad y urgencia para modificar normas de derecho común locativo?",
        ],
    },
    {
        "titulo": "Caso 5 — Ferreyra c/ Clínica San Lucas (Consentimiento informado digital)",
        "filename": "Caso5_Ferreyra_c_ClinicaSanLucas.docx",
        "caratula": "Ferreyra, Ana Beatriz c/ Clínica San Lucas S.A. y otro s/ daños y perjuicios",
        "expediente": "CCC-71203/2024",
        "rama": "Civil",
        "tipo_proceso": "Ordinario",
        "tribunal": "Cámara Nacional de Apelaciones en lo Civil, Sala J",
        "nivel": "N2 — Difícil",
        "hechos": (
            "La actora fue operada de hernia de disco lumbar en enero de 2024 en la Clínica San Lucas. Previo a "
            "la cirugía firmó un consentimiento informado digital mediante tableta, marcando una casilla de "
            "'he leído y acepto'. El documento tenía 48 páginas en letra 8 y no contaba con traducción a lenguaje "
            "llano. La actora tiene nivel educativo primario completo. Postoperatoriamente sufrió una lesión del "
            "nervio ciático izquierdo con pie caído permanente, complicación mencionada en la página 31 del "
            "documento pero no explicada verbalmente. Peritos médicos confirman que la técnica quirúrgica fue "
            "correcta pero la lesión era una complicación conocida con incidencia del 2-4%."
        ),
        "partes": [
            {
                "nombre": "Ferreyra, Ana Beatriz",
                "rol": "Actora",
                "pretension": (
                    "Indemnización por incapacidad parcial permanente (35%), daño moral, daño estético y pérdida "
                    "de chance laboral. Nulidad del consentimiento informado por vicio en la formación del acto."
                ),
                "fundamentos": (
                    "Ley 26.529 (derechos del paciente) arts. 5 y 6: el consentimiento requiere comprensión efectiva, "
                    "no mera firma. Art. 59 CCCN. El formato digital de 48 páginas no garantiza comprensión real "
                    "ante paciente de baja escolaridad. Obligación de informar verbalmente los riesgos relevantes."
                ),
                "jurisprudencia": (
                    "CSJN, 'Valdés c/ IOSE', Fallos 335:996. CNCiv. Sala L, 'Gómez c/ Hospital Alemán', 2022."
                ),
            },
            {
                "nombre": "Clínica San Lucas S.A.",
                "rol": "Demandado",
                "pretension": (
                    "Rechazo de la demanda. Existió consentimiento informado válido. La complicación es inherente "
                    "a la cirugía y fue correctamente ejecutada."
                ),
                "fundamentos": (
                    "El art. 6 Ley 26.529 no exige información verbal: la firma del documento digitalizado cumple "
                    "el requisito legal. No hay nexo causal entre el consentimiento y el daño: la lesión es una "
                    "complicación inherente ejecutada sin culpa."
                ),
                "jurisprudencia": (
                    "CNCiv. Sala I, 'Torres c/ Sanatorio Güemes', 2023. SCJBA, 'Peralta c/ Hospital Privado', 2021."
                ),
            },
        ],
        "cuestiones": [
            "¿El consentimiento informado prestado mediante firma digital en documento de 48 páginas, sin explicación verbal de riesgos, es válido frente a una paciente de baja escolaridad?",
            "¿La validez del consentimiento en formato digital requiere acreditar comprensión efectiva o es suficiente la firma electrónica?",
            "¿Existe nexo causal entre la insuficiencia del consentimiento y el daño cuando la complicación era inherente al procedimiento?",
        ],
    },
    {
        "titulo": "Caso 6 — Martínez c/ Supermercados La Cosecha (Despido por redes sociales)",
        "filename": "Caso6_Martinez_c_LaCosecha.docx",
        "caratula": "Martínez, Javier Hernán c/ Supermercados La Cosecha S.A. s/ despido",
        "expediente": "CNT-19874/2024",
        "rama": "Laboral",
        "tipo_proceso": "Ordinario",
        "tribunal": "Juzgado Nacional del Trabajo N° 5",
        "nivel": "N2 — Difícil",
        "hechos": (
            "El actor se desempeñaba como repositor desde hacía 8 años. En agosto de 2024 publicó en su cuenta "
            "personal de Instagram (perfil privado, 340 seguidores) un video mostrando condiciones de trabajo: "
            "ropa de protección insuficiente en cámaras de frío, falta de guantes y sobrecarga de cajas. La "
            "publicación fue compartida por un compañero fuera del grupo privado y llegó a medios locales. "
            "El supermercado lo despidió con causa a los 3 días, invocando 'violación del deber de fidelidad "
            "y daño a la imagen empresaria'. No existía reglamento interno que prohibiera publicaciones en redes. "
            "Los hechos denunciados resultaron verídicos según peritos."
        ),
        "partes": [
            {
                "nombre": "Martínez, Javier Hernán",
                "rol": "Actor",
                "pretension": (
                    "Indemnización por despido incausado más daño moral por despido discriminatorio en represalia "
                    "por denuncia de irregularidades laborales."
                ),
                "fundamentos": (
                    "Art. 14 CN: libertad de expresión fuera del horario laboral. La denuncia de irregularidades "
                    "en seguridad e higiene es un acto legítimo protegido. Art. 242 LCT: la causa de despido debe "
                    "ser injuria de gravedad suficiente. Ley 24.013 art. 47: protección frente a represalias."
                ),
                "jurisprudencia": (
                    "CNAT Sala V, 'Ruiz c/ Cencosud', 2022. TEDH, 'Herbai c/ Hungría', 2019. "
                    "OIT, Convenio 158 sobre terminación del empleo."
                ),
            },
            {
                "nombre": "Supermercados La Cosecha S.A.",
                "rol": "Demandado",
                "pretension": (
                    "Ratificación del despido con causa. La difusión de información interna dañó gravemente "
                    "la imagen de la empresa y constituye injuria suficiente."
                ),
                "fundamentos": (
                    "Art. 85 LCT: deber de fidelidad. Art. 88 LCT: deber de no perjudicar al empleador. "
                    "La veracidad no elimina la injuria: el medio (redes sociales con masificación) fue "
                    "desproporcionado frente a los canales internos disponibles."
                ),
                "jurisprudencia": (
                    "CNAT Sala III, 'Gorosito c/ McDonald's', 2020. SCJBA, 'Pedernera c/ Supermercado Norte', 2019."
                ),
            },
        ],
        "cuestiones": [
            "¿El deber de fidelidad del art. 85 LCT limita la libertad de expresión del trabajador en redes sociales personales respecto de condiciones de trabajo verídicas?",
            "¿La difusión de condiciones laborales irregulares en redes sociales constituye injuria grave cuando los hechos son verdaderos y el trabajador actuó desde un perfil privado?",
            "¿Existe acto discriminatorio cuando el despido sigue inmediatamente a una denuncia pública de irregularidades en seguridad e higiene?",
        ],
    },
    {
        "titulo": "Caso 7 — Villalba c/ ANSES (Jubilación / Tareas insalubres y género)",
        "filename": "Caso7_Villalba_c_ANSES.docx",
        "caratula": "Villalba, Rosa Carmen c/ ANSES s/ reajuste de haberes y reconocimiento de servicios",
        "expediente": "FRO-88432/2024",
        "rama": "Administrativo",
        "tipo_proceso": "Ordinario",
        "tribunal": "Juzgado Federal de la Seguridad Social N° 6",
        "nivel": "N2 — Difícil",
        "hechos": (
            "La actora se desempeñó durante 22 años como enfermera en guardia de emergencias del Hospital Público "
            "de Rosario. Cumplió 60 años en marzo de 2024 y solicitó jubilación ordinaria con reducción de edad "
            "por tareas insalubres (Decreto 4257/68 y Ley 24.241 art. 157). ANSES rechazó el beneficio argumentando "
            "que enfermería no está incluida en el listado del Decreto 4257/68, cuyo catálogo data de 1968 y no "
            "fue actualizado. La actora aportó pericias médicas que demuestran exposición a radiaciones, productos "
            "químicos, agentes biológicos y estrés postraumático acumulado."
        ),
        "partes": [
            {
                "nombre": "Villalba, Rosa Carmen",
                "rol": "Actora",
                "pretension": (
                    "Reconocimiento del carácter insalubre de sus tareas de enfermería en guardia de emergencias "
                    "y reducción de la edad jubilatoria. Subsidiariamente, declaración de inconstitucionalidad "
                    "del Decreto 4257/68 por discriminación."
                ),
                "fundamentos": (
                    "Art. 14 bis CN: derecho a jubilación. Ley 24.241 art. 157 (tareas insalubres). El catálogo "
                    "del Decreto 4257/68 viola el principio de igualdad por haber sido diseñado para trabajos "
                    "masculinos históricos (minería, pesca) sin contemplar trabajos del cuidado de alto riesgo. "
                    "CEDAW y Convenio OIT 189."
                ),
                "jurisprudencia": (
                    "CSJN, 'Ascua c/ SOMISA', Fallos 333:2306. CFSeg. Social Sala II, 'Blanco c/ ANSES', 2022. "
                    "CSJN, 'Brizuela', Fallos 330:3765."
                ),
            },
            {
                "nombre": "ANSES",
                "rol": "Demandado",
                "pretension": (
                    "Ratificación del rechazo. La determinación de tareas insalubres es de resorte exclusivo "
                    "del Poder Ejecutivo. La inclusión de nuevas actividades requiere acto administrativo expreso."
                ),
                "fundamentos": (
                    "El Poder Judicial no puede sustituir al Ejecutivo en la determinación de actividades "
                    "insalubres. La actora puede acceder a la jubilación ordinaria al cumplir los años de "
                    "aportes. Los beneficios especiales deben ser de interpretación estricta."
                ),
                "jurisprudencia": (
                    "CSJN, 'Itzcovich c/ ANSES', Fallos 328:566. "
                    "CFSeg. Social Sala I, 'Rodríguez c/ ANSES', 2023."
                ),
            },
        ],
        "cuestiones": [
            "¿Puede el juez reconocer el carácter insalubre de tareas de enfermería en emergencias por vía analógica cuando la actividad no está incluida en el Decreto 4257/68?",
            "¿El Decreto 4257/68, al excluir históricamente trabajos del cuidado de alto riesgo, viola el principio de igualdad y no discriminación por razón de género?",
            "¿La omisión del Poder Ejecutivo de actualizar por 55 años el catálogo de tareas insalubres genera una laguna que el Poder Judicial debe integrar?",
        ],
    },
    {
        "titulo": "Caso 8 — Suárez c/ Provincia de Buenos Aires (Femicidio y omisión estatal)",
        "filename": "Caso8_Suarez_c_ProvBsAs.docx",
        "caratula": "Suárez, María Eugenia (en representación de sus hijos menores) c/ Provincia de Buenos Aires s/ daños y perjuicios",
        "expediente": "CCA-34102/2024",
        "rama": "Civil",
        "tipo_proceso": "Ordinario",
        "tribunal": "Cámara de Apelaciones en lo Contencioso Administrativo de La Plata, Sala II",
        "nivel": "N2 — Difícil",
        "hechos": (
            "La madre de la actora fue asesinada por su ex pareja en septiembre de 2023. En los 8 meses previos "
            "la víctima había realizado 6 denuncias penales y obtenido 2 medidas de exclusión del hogar y 1 "
            "prohibición de acercamiento. La policía no verificó el cumplimiento de las medidas. El agresor violó "
            "la prohibición al menos 3 veces sin que la policía actuara. La víctima solicitó en 3 oportunidades "
            "acompañamiento policial y tobillera electrónica, que nunca le fueron asignados pese a estar "
            "disponibles. La Fiscalía tampoco solicitó prisión preventiva pese al historial de violencia."
        ),
        "partes": [
            {
                "nombre": "Suárez, María Eugenia (por sí y en representación de sus hijos menores)",
                "rol": "Actora",
                "pretension": (
                    "Indemnización por daño moral, daño material, pérdida de chance y daño al proyecto de vida "
                    "de los hijos menores por la muerte de su madre. Responsabilidad del Estado provincial "
                    "por omisión en el deber de protección."
                ),
                "fundamentos": (
                    "Art. 1763 CCCN (responsabilidad del Estado por omisión). Ley 26.485 arts. 7 y 16: "
                    "obligación estatal de protección efectiva. Convención de Belém do Pará art. 7: "
                    "debida diligencia. El Estado conocía el riesgo y tenía medios disponibles."
                ),
                "jurisprudencia": (
                    "Corte IDH, 'González y otras (Campo Algodonero) c/ México', 2009. "
                    "CSJN, 'Asociación de Magistrados', Fallos 329:2316. CCAyT CABA, 'S.V. c/ GCBA', 2022."
                ),
            },
            {
                "nombre": "Provincia de Buenos Aires",
                "rol": "Demandado",
                "pretension": (
                    "Rechazo de la demanda. El Estado adoptó todas las medidas razonables. La responsabilidad "
                    "del homicidio es exclusiva del autor material."
                ),
                "fundamentos": (
                    "Art. 1729 CCCN: el hecho del tercero rompe el nexo causal. El Estado no es garante absoluto "
                    "de la seguridad personal. La omisión policial no fue la causa eficiente: el femicidio "
                    "hubiera podido ocurrir igualmente."
                ),
                "jurisprudencia": (
                    "SCBA, 'B.M.A. c/ Provincia de Buenos Aires', 2021. "
                    "CSJN, 'Mosca c/ Buenos Aires', Fallos 330:563."
                ),
            },
        ],
        "cuestiones": [
            "¿La omisión del Estado en verificar el cumplimiento de medidas cautelares de protección frente a violencia de género configura responsabilidad civil extracontractual?",
            "¿El hecho del autor del femicidio interrumpe el nexo causal entre la omisión estatal y el daño, o ambas causas concurren?",
            "¿El estándar de 'debida diligencia' de la Convención de Belém do Pará es directamente aplicable como parámetro de antijuridicidad de la conducta estatal provincial?",
        ],
    },
    {
        "titulo": "Caso 9 — Constructora Del Valle c/ Inversiones Barracas (Deuda en dólares)",
        "filename": "Caso9_ConstructoraDelValle_c_InversionesBarracas.docx",
        "caratula": "Constructora Del Valle S.A. c/ Inversiones Barracas S.R.L. s/ cobro de pesos",
        "expediente": "CCC-41872/2024",
        "rama": "Civil",
        "tipo_proceso": "Ordinario",
        "tribunal": "Juzgado Nacional en lo Comercial N° 22",
        "nivel": "N2 — Difícil",
        "hechos": (
            "En enero de 2019 las partes celebraron un contrato de obra privada por USD 2.500.000, pagaderos "
            "en cuotas mensuales al tipo de cambio oficial vigente a cada vencimiento. Las primeras 12 cuotas "
            "se pagaron normalmente. En septiembre de 2019, ante la imposición del cepo cambiario (DECNU 609/2019), "
            "el deudor comenzó a pagar al tipo oficial ($63/USD) aunque el mercado paralelo cotizaba a $75. "
            "En 2024, al momento del juicio, restan 18 cuotas; el tipo oficial es de $900/USD y el paralelo "
            "de $1.100/USD. El acreedor exige pago al tipo libre o indexación. El deudor argumenta que el "
            "pago al tipo oficial es liberatorio."
        ),
        "partes": [
            {
                "nombre": "Constructora Del Valle S.A.",
                "rol": "Actora",
                "pretension": (
                    "Cobro de cuotas impagas al tipo de cambio MEP/CCL, o subsidiariamente con indexación "
                    "por CER desde cada vencimiento, más intereses."
                ),
                "fundamentos": (
                    "Art. 765 CCCN: el 'tipo de cambio oficial' debe interpretarse como el tipo al que el "
                    "acreedor puede efectivamente hacerse de los dólares pactados (MEP/CCL). Art. 1091 CCCN: "
                    "imprevisión contractual. La brecha cambiaria superior al 100% altera extraordinariamente "
                    "la ecuación del contrato."
                ),
                "jurisprudencia": (
                    "CNCom. Sala A, 'Trusso c/ Ares Inmobiliaria', 2023. "
                    "CSJN, 'Rinaldi c/ Guzmán', Fallos 330:855."
                ),
            },
            {
                "nombre": "Inversiones Barracas S.R.L.",
                "rol": "Demandado",
                "pretension": (
                    "Declaración de que el pago al tipo de cambio oficial BNA es plenamente liberatorio. "
                    "Rechazo de cualquier indexación adicional."
                ),
                "fundamentos": (
                    "Art. 765 CCCN: el deudor se libera pagando el equivalente en pesos al tipo oficial. "
                    "Ley 23.928 y Ley 25.561: prohibición de indexar. El riesgo cambiario es inherente "
                    "a las obligaciones en moneda extranjera y no configura imprevisión."
                ),
                "jurisprudencia": (
                    "CNCom. Sala F, 'Ríos c/ Navarro', 2023. CNCom. Sala B, 'Construcciones Omega c/ Torres', 2022."
                ),
            },
        ],
        "cuestiones": [
            "¿En una obligación contractual privada pactada en dólares, el pago al tipo de cambio oficial BNA es liberatorio cuando existe una brecha superior al 100% con los tipos financieros legales (MEP/CCL)?",
            "¿La imposibilidad legal de acceder a los dólares al tipo oficial (cepo cambiario) configura un supuesto de imprevisión del art. 1091 CCCN?",
            "¿La prohibición de indexar de la Ley 23.928 aplica cuando la discusión no es sobre actualización sino sobre el tipo de cambio para la conversión?",
        ],
    },
    {
        "titulo": "Caso 10 — Frigorífico San Cayetano c/ López (IA como prueba judicial)",
        "filename": "Caso10_FrigorificoSanCayetano_c_Lopez.docx",
        "caratula": "Frigorífico San Cayetano S.A. c/ López, Héctor Ramón s/ despido con causa",
        "expediente": "CNT-28834/2024",
        "rama": "Laboral",
        "tipo_proceso": "Ordinario",
        "tribunal": "Cámara Nacional del Trabajo, Sala VIII",
        "nivel": "N2 — Difícil",
        "hechos": (
            "El demandado fue despedido con causa por supuestamente haber filtrado secretos comerciales. "
            "Como prueba principal, el empleador presentó un informe pericial generado íntegramente mediante "
            "GPT-4 que analizó capturas de WhatsApp del teléfono corporativo y emitió un dictamen sobre el "
            "'tono' e 'intención' de los mensajes, concluyendo que el trabajador 'probablemente' transmitió "
            "información confidencial. El perito oficial cuestionó la metodología: el sistema de IA no tiene "
            "acceso a metadatos originales, no puede verificar la autenticidad de las capturas y opera sobre "
            "probabilidades estadísticas. El trabajador niega haber enviado los mensajes y alega manipulación."
        ),
        "partes": [
            {
                "nombre": "Frigorífico San Cayetano S.A.",
                "rol": "Actor",
                "pretension": (
                    "Ratificación del despido con causa por violación del deber de fidelidad (art. 85 LCT). "
                    "Validez del informe pericial generado por IA como prueba científica admisible."
                ),
                "fundamentos": (
                    "Art. 386 CPCCN: libre apreciación de la prueba. El informe de IA es prueba pericial "
                    "que el juez puede valorar libremente. La verosimilitud y consistencia del análisis "
                    "crea presunción grave. Art. 242 LCT: la causa no requiere certeza absoluta."
                ),
                "jurisprudencia": (
                    "CNCom. Sala C, 'Google Argentina c/ Da Cunha', 2023. "
                    "Daubert v. Merrell Dow Pharmaceuticals (admisibilidad prueba científica, EEUU)."
                ),
            },
            {
                "nombre": "López, Héctor Ramón",
                "rol": "Demandado",
                "pretension": (
                    "Nulidad del informe pericial de IA. Reconocimiento del despido como incausado. "
                    "Indemnización completa más daño moral."
                ),
                "fundamentos": (
                    "Art. 18 CN: garantía de defensa en juicio requiere prueba verificable y reproducible. "
                    "Un informe de IA sobre 'intenciones probables' no es prueba científica: no tiene "
                    "metodología reproducible, no puede ser contrapericiado y opera sobre sesgos estadísticos. "
                    "Art. 472 CPCCN: el perito debe explicar fundamentos."
                ),
                "jurisprudencia": (
                    "CNAT Sala VII, 'Medina c/ Telecom Argentina', 2023. TEDH, 'Lekić v. Slovenia', 2019."
                ),
            },
        ],
        "cuestiones": [
            "¿Un informe pericial generado íntegramente por IA que analiza capturas de pantalla y emite conclusiones sobre 'intención probable' de mensajes constituye prueba pericial válida en proceso laboral?",
            "¿La admisión de un dictamen de IA sobre autenticidad e intención de comunicaciones, sin posibilidad de contrapericia técnica equivalente, viola la garantía del art. 18 CN?",
            "¿Qué estándares mínimos de verificabilidad, reproducibilidad y transparencia algorítmica debe cumplir una prueba basada en IA para ser admisible en proceso judicial argentino?",
        ],
    },
]

# ============================================================
# FUNCION PARA CREAR WORD
# ============================================================
def crear_word(caso, filepath):
    doc = Document()

    # Estilos
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # Título principal
    titulo = doc.add_heading(level=0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo.add_run('JUSTICIA ARGENTINA')
    run.font.color.rgb = RGBColor(0x1a, 0x36, 0x5d)
    run.font.size = Pt(16)

    subtitulo = doc.add_heading(level=1)
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = subtitulo.add_run(caso['titulo'])
    run2.font.size = Pt(13)

    doc.add_paragraph()

    # Datos identificatorios
    tabla_id = doc.add_table(rows=6, cols=2)
    tabla_id.style = 'Table Grid'
    campos_id = [
        ('Carátula', caso['caratula']),
        ('Expediente', caso['expediente']),
        ('Rama del derecho', caso['rama']),
        ('Tipo de proceso', caso['tipo_proceso']),
        ('Tribunal', caso['tribunal']),
        ('Nivel de complejidad', caso['nivel']),
    ]
    for i, (campo, valor) in enumerate(campos_id):
        tabla_id.rows[i].cells[0].text = campo
        tabla_id.rows[i].cells[1].text = valor
        tabla_id.rows[i].cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    # Hechos probados
    h = doc.add_heading('HECHOS PROBADOS', level=2)
    doc.add_paragraph(caso['hechos'])
    doc.add_paragraph()

    # Partes
    doc.add_heading('POSICIONES PROCESALES DE LAS PARTES', level=2)
    for parte in caso['partes']:
        p_titulo = doc.add_heading(f"{parte['nombre']} — {parte['rol']}", level=3)
        tabla_p = doc.add_table(rows=3, cols=2)
        tabla_p.style = 'Table Grid'
        campos_p = [
            ('Pretensión', parte['pretension']),
            ('Fundamentos jurídicos', parte['fundamentos']),
            ('Jurisprudencia citada', parte['jurisprudencia']),
        ]
        for i, (campo, valor) in enumerate(campos_p):
            tabla_p.rows[i].cells[0].text = campo
            tabla_p.rows[i].cells[1].text = valor
            tabla_p.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        doc.add_paragraph()

    # Cuestiones a resolver
    doc.add_heading('CUESTIONES A RESOLVER', level=2)
    for i, c in enumerate(caso['cuestiones'], 1):
        doc.add_paragraph(f"{i}. {c}", style='List Number')

    doc.add_paragraph()

    # Footer
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_f = footer_p.add_run('Generado por JUSTICIA ARGENTINA — Sistema de Resolución Asistida')
    run_f.font.size = Pt(9)
    run_f.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.save(filepath)
    print(f'  Guardado: {os.path.basename(filepath)}')


# ============================================================
# 1. CREAR LOS 10 WORD EN DESCARGAS
# ============================================================
print('=' * 60)
print('PASO 1: Creando archivos Word en Descargas...')
print('=' * 60)
for caso in CASOS:
    filepath = os.path.join(CASOS1_DIR, caso['filename'])
    crear_word(caso, filepath)
print(f'\n10 archivos Word guardados en: {CASOS1_DIR}')


# ============================================================
# 2. CORRER EL MOTOR CONTRA LOS 10 CASOS
# ============================================================
print()
print('=' * 60)
print('PASO 2: Ejecutando motor de sentencias...')
print('=' * 60)

from generador.motor_sentencias import GeneradorSentencias, CasoInput
from generador.clasificador import clasificar_caso

def extraer_normas(texto):
    patrones = [
        r'[Aa]rt(?:\u00edculo|iculo)?\.?\s*\d+(?:\s*bis)?',
        r'[Ll]ey\s+\d{4,6}',
        r'L\.?C\.?T\.?',
        r'PIDESC|Pacto\s+Internacional',
        r'CADH|Convenci\u00f3n\s+Americana',
        r'OIT|Convenio\s+\d+',
        r'CCCN|C\u00f3digo\s+Civil',
    ]
    encontradas = set()
    for p in patrones:
        for m in re.finditer(p, texto, re.IGNORECASE):
            encontradas.add(m.group(0).strip().upper()[:30])
    return encontradas

gen = GeneradorSentencias()
print(f'Modelo: {gen.gemini_model}')

campos_validos = set(CasoInput.__dataclass_fields__.keys())
resultados = []

for i, caso in enumerate(CASOS):
    print(f'\n[{i+1}/10] {caso["titulo"]}')

    # Armar el input para el motor
    data = {
        'caratula': caso['caratula'],
        'expediente': caso['expediente'],
        'rama': caso['rama'].lower(),
        'tipo_proceso': caso['tipo_proceso'].lower(),
        'tribunal': caso['tribunal'],
        'hechos_probados': caso['hechos'],
        'partes': [
            {
                'nombre': p['nombre'],
                'rol': p['rol'].lower(),
                'pretension': p['pretension'],
                'fundamentos_juridicos': p['fundamentos'],
                'jurisprudencia_citada': p['jurisprudencia'],
            }
            for p in caso['partes']
        ],
        'cuestiones_a_resolver': caso['cuestiones'],
    }

    cl = clasificar_caso(data)
    nivel = cl['nivel']
    data['nivel_complejidad'] = nivel

    caso_filtrado = {k: v for k, v in data.items() if k in campos_validos}
    caso_obj = CasoInput(**caso_filtrado)
    sentencia = gen.generar_sentencia(caso_obj)

    texto = sentencia.texto_completo or ''
    tu = texto.upper()
    estructura = ((1 if 'VISTOS' in tu else 0) + (1 if 'CONSIDERANDO' in tu else 0) + (1 if 'RESUELVO' in tu else 0)) / 3
    normas = extraer_normas(texto)

    print(f'  Nivel: {nivel} | Estructura: {estructura*100:.0f}% | Normas detectadas: {len(normas)}')

    # Extraer considerandos y resuelvo del texto_completo si están vacíos en el objeto
    considerandos = sentencia.considerandos or ''
    resuelvo = sentencia.resuelvo or ''
    if not considerandos:
        idx_c = tu.find('CONSIDERANDO')
        idx_r = tu.find('RESUELVO')
        if idx_c >= 0:
            considerandos = texto[idx_c:idx_r].strip() if idx_r > idx_c else texto[idx_c:].strip()
    if not resuelvo:
        idx_r = tu.find('RESUELVO')
        if idx_r >= 0:
            resuelvo = texto[idx_r:].strip()

    print(f'  Considerandos: {len(considerandos)} chars | Resuelvo: {len(resuelvo)} chars')

    resultados.append({
        'caso': caso['titulo'],
        'nivel': nivel,
        'estructura': estructura,
        'normas_count': len(normas),
        'considerandos_preview': considerandos[:400],
        'resuelvo_preview': resuelvo[:300],
        'texto_completo': texto,
    })

    if i < len(CASOS) - 1:
        time.sleep(4)

# Resumen
print()
print('=' * 60)
print('RESUMEN EJECUCION')
print('=' * 60)
for r in resultados:
    print(f'{r["caso"][:52]:52s} | Struct={r["estructura"]*100:.0f}% | Normas={r["normas_count"]}')

# Guardar resultados JSON
results_path = os.path.join(CASOS1_DIR, 'resultados_10casos.json')
with open(results_path, 'w', encoding='utf-8') as f:
    json.dump(resultados, f, ensure_ascii=False, indent=2)
print(f'\nResultados guardados en: {results_path}')


# ============================================================
# 3. GENERAR 10 CASOS NUEVOS Y GUARDARLOS EN casos2/
# ============================================================
print()
print('=' * 60)
print('PASO 3: Creando 10 casos nuevos en ~/Downloads/casos2/')
print('=' * 60)

CASOS2 = [
    {
        "titulo": "Caso 11 — Oviedo c/ Municipalidad de Córdoba (Acoso laboral en empleo público)",
        "filename": "Caso11_Oviedo_c_Municipalidad.docx",
        "caratula": "Oviedo, Claudia Beatriz c/ Municipalidad de Córdoba s/ amparo laboral",
        "expediente": "AMC-14203/2024",
        "rama": "Administrativo",
        "tipo_proceso": "Amparo",
        "tribunal": "Cámara Contencioso Administrativo de Córdoba, Sala 1",
        "nivel": "N2 — Difícil",
        "hechos": (
            "La actora, auxiliar administrativa de la Municipalidad de Córdoba con 14 años de antigüedad, "
            "denunció a su jefe directo por acoso laboral (mobbing) sostenido durante 2 años: cambio reiterado "
            "de funciones a tareas degradantes, aislamiento físico en un cubículo sin ventanas, exclusión de "
            "reuniones de equipo y descalificación verbal documentada por 6 testigos. Presentó dos denuncias "
            "internas ante la Dirección de Recursos Humanos, sin respuesta formal en 90 días. Solicitó mediante "
            "amparo su reubicación inmediata en otra dependencia. La Municipalidad argumenta que las medidas "
            "organizacionales fueron razonables y que no existió sumario previo que acredite el acoso."
        ),
        "partes": [
            {
                "nombre": "Oviedo, Claudia Beatriz",
                "rol": "Actora",
                "pretension": (
                    "Reubicación inmediata en otra dependencia municipal, instrucción de sumario disciplinario "
                    "contra el acosador y reparación del daño psicológico documentado."
                ),
                "fundamentos": (
                    "Ley 26.485 (violencia laboral). Ley nacional 24.185 (convenciones colectivas sector "
                    "público). Art. 14 bis CN: condiciones dignas de trabajo. Ley provincial 9.224 (acoso "
                    "laboral Córdoba). El amparo procede por urgencia ante daño psicológico en curso. La "
                    "omisión administrativa de investigar el acoso es en sí misma una conducta antijurídica."
                ),
                "jurisprudencia": (
                    "CSJN, 'Ramos c/ Estado Nacional', Fallos 333:311. TSJ Córdoba, 'Pinto c/ Municipalidad "
                    "de Río Cuarto', 2022. Corte IDH, 'Lagos del Campo c/ Perú', 2017."
                ),
            },
            {
                "nombre": "Municipalidad de Córdoba",
                "rol": "Demandado",
                "pretension": (
                    "Rechazo del amparo. El empleado público no puede ser reubicado sin sumario previo. "
                    "Las medidas organizacionales obedecieron a necesidades del servicio."
                ),
                "fundamentos": (
                    "El amparo no es la vía procesal adecuada: existe la vía contencioso-administrativa. "
                    "Las medidas de reorganización interna son facultades discrecionales del empleador "
                    "público. Sin sumario concluido no existe acoso acreditado. El pase de dependencia "
                    "requiere trámite administrativo previo."
                ),
                "jurisprudencia": (
                    "CSJN, 'Schnaiderman c/ Estado Nacional', Fallos 331:735. TSJ Córdoba, "
                    "'González c/ Municipalidad de Córdoba', 2021."
                ),
            },
        ],
        "cuestiones": [
            "¿El amparo es la vía procesal adecuada para obtener la reubicación urgente de un empleado público víctima de mobbing documentado cuando la vía administrativa fue agotada sin respuesta?",
            "¿La omisión del empleador público de investigar denuncias de acoso laboral en el plazo reglamentario configura conducta antijurídica que legitima la acción de amparo?",
            "¿Puede el juez ordenar medidas organizacionales internas (reubicación, sumario) al empleador público sin violar el principio de separación de poderes?",
        ],
    },
    {
        "titulo": "Caso 12 — Torres c/ Instituto Educativo Privado San Martín (Despido docente / Maternidad)",
        "filename": "Caso12_Torres_c_InstitutoSanMartin.docx",
        "caratula": "Torres, Florencia Daniela c/ Instituto Educativo Privado San Martín S.A. s/ despido discriminatorio",
        "expediente": "CNT-33421/2024",
        "rama": "Laboral",
        "tipo_proceso": "Ordinario",
        "tribunal": "Juzgado Nacional del Trabajo N° 18",
        "nivel": "N2 — Difícil",
        "hechos": (
            "La actora era docente de nivel primario en el Instituto San Martín desde 2018. Notificó su embarazo "
            "en marzo de 2024. En abril de 2024 el instituto no renovó su contrato para el ciclo lectivo siguiente, "
            "invocando 'reestructuración curricular' que eliminó el cargo. En los hechos, el cargo fue cubierto "
            "por otra docente sin embarazo a los 15 días del despido. La actora tiene a cargo a su madre "
            "con enfermedad crónica. El instituto alega que el estatuto del docente privado (Ley 13.047) "
            "habilita la no renovación al final del ciclo lectivo sin expresión de causa."
        ),
        "partes": [
            {
                "nombre": "Torres, Florencia Daniela",
                "rol": "Actora",
                "pretension": (
                    "Nulidad del despido por discriminación por embarazo. Reinstalación o indemnización "
                    "agravada (art. 182 LCT). Daño moral. Aplicación de la presunción del art. 178 LCT."
                ),
                "fundamentos": (
                    "Art. 178 LCT: presunción de despido por embarazo cuando ocurre dentro de los 7,5 meses "
                    "anteriores o posteriores al parto. Art. 182 LCT: indemnización especial. Ley 23.592 "
                    "(antidiscriminación): el cargo fue inmediatamente cubierto por una no embarazada. "
                    "El Estatuto del Docente Privado no puede derogar protecciones de orden público."
                ),
                "jurisprudencia": (
                    "CSJN, 'Álvarez c/ Cencosud', Fallos 333:2306. CNAT Sala II, 'Reyes c/ Colegio "
                    "Del Sol', 2023. CEDAW, Recomendación General 28."
                ),
            },
            {
                "nombre": "Instituto Educativo Privado San Martín S.A.",
                "rol": "Demandado",
                "pretension": (
                    "Rechazo de la demanda. La no renovación al fin del ciclo no es despido. El estatuto "
                    "docente regula específicamente la relación y desplaza la LCT."
                ),
                "fundamentos": (
                    "Ley 13.047 (Estatuto del Docente Privado) art. 11: la no renovación al fin del ciclo "
                    "es una facultad expresa del empleador. El estatuto especial prevalece sobre la LCT "
                    "por especialidad. No hubo acto positivo de despido: solo expiración contractual. "
                    "La reestructuración curricular fue real y documentada."
                ),
                "jurisprudencia": (
                    "CNAT Sala VI, 'Rodríguez c/ Instituto Belgrano', 2022. SCJBA, 'Pereyra c/ Colegio "
                    "Norteño', 2020."
                ),
            },
        ],
        "cuestiones": [
            "¿La no renovación de contrato docente al fin del ciclo lectivo, cuando ocurre dentro del período de protección del art. 178 LCT, activa la presunción de despido por embarazo?",
            "¿El Estatuto del Docente Privado (Ley 13.047) desplaza las protecciones de la LCT en materia de despido por maternidad, o ambas normas deben armonizarse?",
            "¿La inmediata cobertura del cargo por una docente no embarazada constituye indicio suficiente de discriminación bajo la Ley 23.592?",
        ],
    },
    {
        "titulo": "Caso 13 — Herrera c/ Obra Social OSECAC (Cobertura de tratamiento experimental)",
        "filename": "Caso13_Herrera_c_OSECAC.docx",
        "caratula": "Herrera, Martín Alejandro c/ Obra Social de Empleados de Comercio (OSECAC) s/ amparo de salud",
        "expediente": "FAM-22104/2024",
        "rama": "Civil",
        "tipo_proceso": "Amparo",
        "tribunal": "Juzgado Federal Civil y Comercial N° 1 de Mendoza",
        "nivel": "N2 — Difícil",
        "hechos": (
            "El actor, de 38 años, fue diagnosticado con esclerosis lateral amiotrófica (ELA) en estadio inicial. "
            "Su médico tratante prescribió un tratamiento con células madre autólogas que tiene aprobación "
            "compasiva en fase II en Argentina (ANMAT Disposición 1429/2023) pero no está incluido en el "
            "Programa Médico Obligatorio (PMO). El costo del tratamiento es de USD 45.000. OSECAC rechazó la "
            "cobertura por no ser tratamiento aprobado definitivamente por ANMAT y no estar en el PMO. Peritos "
            "médicos indican que sin el tratamiento la progresión de la ELA será irreversible en 18-24 meses. "
            "Hay evidencia científica preliminar de beneficio en el 60% de los casos en fase II."
        ),
        "partes": [
            {
                "nombre": "Herrera, Martín Alejandro",
                "rol": "Actor",
                "pretension": (
                    "Cobertura íntegra del tratamiento de células madre autólogas por parte de OSECAC, "
                    "con carácter urgente. Declaración de que la limitación del PMO no puede oponerse "
                    "al derecho a la vida frente a una enfermedad sin cura."
                ),
                "fundamentos": (
                    "Art. 42 CN: derecho a la salud. Ley 23.661 (obras sociales) art. 28: cobertura de "
                    "enfermedades crónicas. Ley 26.689 (enfermedades poco frecuentes): ELA es EPF, obliga "
                    "cobertura aunque no esté en PMO. Art. 75 inc. 22 CN: PIDESC art. 12 (derecho al nivel "
                    "más alto de salud). El PMO es un piso, no un techo."
                ),
                "jurisprudencia": (
                    "CSJN, 'Campodónico de Beviacqua c/ Estado Nacional', Fallos 323:3229. "
                    "CSJN, 'Asociación Francesa Filantrópica c/ OSECAC', Fallos 324:754. "
                    "CNCiv. Sala M, 'Rivero c/ Swiss Medical', 2023."
                ),
            },
            {
                "nombre": "Obra Social de Empleados de Comercio (OSECAC)",
                "rol": "Demandado",
                "pretension": (
                    "Rechazo del amparo. La obra social no está obligada a cubrir tratamientos no "
                    "incluidos en el PMO ni aprobados definitivamente por ANMAT."
                ),
                "fundamentos": (
                    "El PMO es el catálogo mínimo obligatorio fijado por el Estado. OSECAC cumple con "
                    "todas las prestaciones del PMO. Un tratamiento en fase II no tiene eficacia demostrada "
                    "y su financiamiento con fondos solidarios del sistema perjudicaría al resto de los "
                    "afiliados. Corresponde al Estado, no a la obra social, financiar tratamientos experimentales."
                ),
                "jurisprudencia": (
                    "CNCiv. Sala K, 'González c/ OSDE', 2022 (rechazó cobertura tratamiento experimental). "
                    "CSJN, 'Hospital Británico c/ IOMA', Fallos 328:4640."
                ),
            },
        ],
        "cuestiones": [
            "¿La Ley 26.689 de enfermedades poco frecuentes obliga a las obras sociales a cubrir tratamientos en fase II de aprobación compasiva no incluidos en el PMO cuando son la única alternativa disponible para una ELA?",
            "¿El principio de reserva solidaria del sistema de obras sociales puede oponerse al derecho individual a la salud y a la vida frente a una enfermedad degenerativa irreversible?",
            "¿El PMO actúa como límite máximo o como piso mínimo de las obligaciones de cobertura de las obras sociales?",
        ],
    },
    {
        "titulo": "Caso 14 — Asociación Proteger c/ Municipalidad de Tigre (Derecho ambiental urbano)",
        "filename": "Caso14_AsocProteger_c_Tigre.docx",
        "caratula": "Asociación Civil Proteger el Delta c/ Municipalidad de Tigre s/ acción ambiental colectiva",
        "expediente": "FAM-9823/2024",
        "rama": "Civil",
        "tipo_proceso": "Amparo",
        "tribunal": "Juzgado Federal de Primera Instancia de San Isidro N° 2",
        "nivel": "N2 — Difícil",
        "hechos": (
            "La Municipalidad de Tigre otorgó en 2022 un permiso de construcción para un complejo "
            "de 8 torres residenciales en una parcela de 12 hectáreas en el Delta del Tigre. El "
            "proyecto implica el relleno de 4 hectáreas de humedales nativos con infraestructura de "
            "hormigón. La Asociación demandante acredita con pericias que los humedales cumplen funciones "
            "de amortiguación de inundaciones para 15.000 familias ribereñas y son hábitat de 12 "
            "especies en categoría de amenaza. El EIA municipal aprobó el proyecto sin consulta pública "
            "ni audiencia ambiental. La obra está al 30% de avance."
        ),
        "partes": [
            {
                "nombre": "Asociación Civil Proteger el Delta",
                "rol": "Actora",
                "pretension": (
                    "Suspensión inmediata de la obra. Nulidad del permiso municipal por omisión de "
                    "audiencia pública y EIA insuficiente. Restauración de los humedales afectados."
                ),
                "fundamentos": (
                    "Art. 41 CN: derecho a un ambiente sano. Ley 25.675 (General del Ambiente) art. 20: "
                    "audiencia pública obligatoria. Ley 25.688 (aguas): protección de humedales. Principio "
                    "precautorio (art. 4 LGA). El daño ambiental es de imposible reparación ulterior: "
                    "procede la medida cautelar innovativa. Legitimación de la asociación (art. 43 CN)."
                ),
                "jurisprudencia": (
                    "CSJN, 'Mendoza c/ Estado Nacional' (Riachuelo), Fallos 331:1622. "
                    "CSJN, 'Salas c/ Salta', Fallos 332:663. CApel. SM, 'Ambientalistas c/ ACUMAR', 2023."
                ),
            },
            {
                "nombre": "Municipalidad de Tigre",
                "rol": "Demandado",
                "pretension": (
                    "Rechazo del amparo. El permiso fue otorgado conforme el código de ordenamiento "
                    "urbano y el EIA aprobado. La paralización de la obra causaría daño económico "
                    "desproporcionado al inversor privado."
                ),
                "fundamentos": (
                    "El poder de policía urbana es de competencia municipal exclusiva (art. 123 CN). "
                    "El EIA fue aprobado por organismos técnicos competentes. La Ley 25.675 no exige "
                    "audiencia pública para proyectos privados en jurisdicción municipal. Principio "
                    "de proporcionalidad: el daño económico de la paralización supera el ambiental."
                ),
                "jurisprudencia": (
                    "CSJN, 'Acumar c/ Inversiones Delta SA', 2021. CApel. San Martín, 'Desarrolladora "
                    "Paraná c/ Municipalidad de Tigre', 2022."
                ),
            },
        ],
        "cuestiones": [
            "¿La omisión de audiencia pública en el EIA de un proyecto que afecta humedales con funciones ecosistémicas críticas invalida el permiso municipal, aun cuando el código urbano local no la exija expresamente?",
            "¿El principio precautorio de la Ley 25.675 justifica la suspensión cautelar de una obra al 30% de avance ante la amenaza de daño ambiental irreversible sobre humedales?",
            "¿El poder de policía urbana municipal puede admitir el relleno de humedales sin coordinación con las autoridades de cuenca y sin aplicar la jerarquía normativa ambiental federal?",
        ],
    },
    {
        "titulo": "Caso 15 — Ramos c/ Empresa de Transporte Vía Uno (Accidente de tránsito / Seguro insuficiente)",
        "filename": "Caso15_Ramos_c_ViaUno.docx",
        "caratula": "Ramos, Santiago Ignacio c/ Empresa de Transporte Vía Uno S.A. y Citio Seguros S.A. s/ daños y perjuicios",
        "expediente": "CCC-55029/2024",
        "rama": "Civil",
        "tipo_proceso": "Ordinario",
        "tribunal": "Juzgado Nacional Civil N° 44",
        "nivel": "N2 — Difícil",
        "hechos": (
            "El actor sufrió lesiones graves (paraplejia parcial, 65% de incapacidad) al ser embestido por "
            "un colectivo de la línea 203 operado por Vía Uno S.A. en CABA. La sentencia de primera instancia "
            "determinó responsabilidad de la empresa y condenó al pago de $85.000.000 en concepto de "
            "incapacidad sobreviniente, daño moral y gastos médicos futuros. La aseguradora Citio Seguros S.A. "
            "tiene una póliza cuyo límite de cobertura es de $8.000.000 (monto fijado en 2019 y nunca "
            "actualizado). La aseguradora invoca el límite contractual. La empresa transportista alega "
            "insolvencia. La víctima reclama el monto total a la aseguradora por oponibilidad limitada de "
            "la franquicia frente a terceros damnificados."
        ),
        "partes": [
            {
                "nombre": "Ramos, Santiago Ignacio",
                "rol": "Actor",
                "pretension": (
                    "Cobro de la totalidad de $85.000.000 de la aseguradora Citio Seguros, sin limitación "
                    "al tope contractual de $8.000.000. Subsidiariamente, declaración de inoponibilidad "
                    "del límite por desactualización y lesión al derecho de la víctima."
                ),
                "fundamentos": (
                    "Art. 118 Ley 17.418 (seguros): el tercero damnificado puede citar a la aseguradora y "
                    "ejecutar la sentencia contra ella. El límite de $8.000.000 fijado en 2019 es írrito: "
                    "su valor real es el 9% de la condena, lo que lo convierte en cobertura inexistente. "
                    "Arts. 1757 y 1758 CCCN: responsabilidad objetiva. Obligación de la aseguradora de "
                    "mantener límites razonables conforme la función social del seguro obligatorio."
                ),
                "jurisprudencia": (
                    "CSJN, 'Buffoni c/ Castro', Fallos 337:329. CNCiv. Sala H, 'López c/ Bus Norte', 2023. "
                    "CSJN, 'Flores c/ Transportes Rio de la Plata', Fallos 340:765."
                ),
            },
            {
                "nombre": "Citio Seguros S.A.",
                "rol": "Demandado",
                "pretension": (
                    "Limitación de su responsabilidad al tope contractual de $8.000.000. El contrato "
                    "de seguro y sus límites son oponibles al tercero damnificado."
                ),
                "fundamentos": (
                    "Art. 118 Ley 17.418: la aseguradora responde 'en la medida del seguro'. El tope "
                    "contractual es válido y fue aceptado por el transportista al contratar. La "
                    "actualización de límites corresponde al Poder Ejecutivo (SSN), no puede ser declarada "
                    "por el juez. La insolvencia del asegurado no traslada el riesgo empresario a la aseguradora."
                ),
                "jurisprudencia": (
                    "CNCom. Sala D, 'Pedraza c/ Seguros San Cristóbal', 2022. SSN Resolución 39.927/2016."
                ),
            },
        ],
        "cuestiones": [
            "¿El límite de cobertura de una póliza de seguro obligatorio de transporte público, fijado en 2019 y nunca actualizado, es oponible al tercero damnificado cuando su valor real representa el 9% de la condena judicial?",
            "¿Puede el juez declarar inoponible o incrementar el límite contractual del seguro cuando su desactualización inflacionaria lo convierte en cobertura materialmente inexistente?",
            "¿La función social del seguro obligatorio de transporte público impone a la aseguradora obligaciones que trascienden el límite contractual frente a víctimas con incapacidades severas?",
        ],
    },
    {
        "titulo": "Caso 16 — Pereyra c/ Universidad Nacional del Sur (Acceso a cargo docente / Discapacidad)",
        "filename": "Caso16_Pereyra_c_UNS.docx",
        "caratula": "Pereyra, Lucas Emmanuel c/ Universidad Nacional del Sur s/ amparo por discriminación",
        "expediente": "FBB-7231/2024",
        "rama": "Administrativo",
        "tipo_proceso": "Amparo",
        "tribunal": "Juzgado Federal N° 1 de Bahía Blanca",
        "nivel": "N2 — Difícil",
        "hechos": (
            "El actor, abogado con discapacidad motriz (usuario de silla de ruedas), obtuvo el primer "
            "lugar en el concurso docente para Profesor Adjunto de Derecho Constitucional en la UNS. "
            "El jurado, en dictamen fundado, lo calificó como mejor candidato. La Junta Electoral de la "
            "UNS rechazó su designación argumentando que el edificio de la Facultad de Derecho no tiene "
            "accesibilidad total: el aula asignada al cargo está en el primer piso sin ascensor. La UNS "
            "propuso designar al segundo candidato 'hasta que se resuelvan las obras de accesibilidad', "
            "sin fecha comprometida. El actor tiene certificado de discapacidad y cupo del 4% del art. 8 "
            "Ley 22.431 le es aplicable."
        ),
        "partes": [
            {
                "nombre": "Pereyra, Lucas Emmanuel",
                "rol": "Actor",
                "pretension": (
                    "Designación inmediata en el cargo de Profesor Adjunto ganado por concurso. "
                    "Orden a la UNS de proveer adaptaciones razonables (aula accesible en planta baja). "
                    "Declaración de que la falta de accesibilidad no puede privarlo del derecho ganado."
                ),
                "fundamentos": (
                    "Ley 22.431 art. 8: cupo del 4% y obligación de adaptar el puesto. Ley 26.378 (CDPD) "
                    "art. 2: obligación de 'ajustes razonables'. Art. 16 CN: igualdad ante la ley. "
                    "La falta de accesibilidad es responsabilidad del Estado, no una causa que prive al "
                    "actor del cargo ganado. El ajuste razonable (cambiar de aula) no implica carga "
                    "desproporcionada para la universidad."
                ),
                "jurisprudencia": (
                    "CSJN, 'Recurso de hecho: Q.C., S.Y.', Fallos 335:452. CNCiv. Sala A, 'Barrionuevo "
                    "c/ GCBA', 2023. CDPD Comité, Observación General N°6."
                ),
            },
            {
                "nombre": "Universidad Nacional del Sur",
                "rol": "Demandado",
                "pretension": (
                    "Rechazo del amparo. La decisión de postergar la designación fue adoptada para "
                    "proteger al propio candidato. La universidad no puede cumplir el cargo sin "
                    "la infraestructura adecuada."
                ),
                "fundamentos": (
                    "La universidad actúa dentro de sus facultades de autogobierno (art. 75 inc. 19 CN). "
                    "Las obras de accesibilidad requieren presupuesto que depende del Ministerio de "
                    "Educación. El cambio de aula afecta la organización académica del departamento. "
                    "La postergación no es discriminación sino adecuación razonable al interés del actor."
                ),
                "jurisprudencia": (
                    "CSJN, 'UBA s/ autonomía universitaria', Fallos 329:3066. CApel. Cont. Adm. Federal, "
                    "'Díaz c/ UBA', 2022."
                ),
            },
        ],
        "cuestiones": [
            "¿La ausencia de accesibilidad física en el edificio universitario puede válidamente postergar la designación de un candidato que ganó el concurso docente por sus méritos, en lugar de obligar a la universidad a proveer ajustes razonables?",
            "¿El principio de ajuste razonable de la CDPD (Ley 26.378) impone a la universidad reasignar el aula a planta baja como obligación de resultado, no como opción discrecional?",
            "¿La autonomía universitaria del art. 75 inc. 19 CN puede invocarse para postergar el cumplimiento de obligaciones de no discriminación por discapacidad?",
        ],
    },
    {
        "titulo": "Caso 17 — García Hermanos SRL c/ AFIP (Multa fiscal / Error excusable)",
        "filename": "Caso17_GarciaHermanos_c_AFIP.docx",
        "caratula": "García Hermanos S.R.L. c/ AFIP-DGI s/ recurso de apelación — multa por defraudación fiscal",
        "expediente": "CAF-18823/2024",
        "rama": "Administrativo",
        "tipo_proceso": "Ordinario",
        "tribunal": "Cámara Nacional de Apelaciones en lo Contencioso Administrativo Federal, Sala IV",
        "nivel": "N2 — Difícil",
        "hechos": (
            "García Hermanos S.R.L. declaró incorrectamente el IVA de 18 operaciones durante el ejercicio "
            "2022, omitiendo $4.200.000 en débitos fiscales. La AFIP determinó la deuda y aplicó una multa "
            "por defraudación fiscal (art. 46 Ley 11.683) equivalente al 200% del impuesto omitido "
            "($8.400.000). La empresa acredita que el error se debió a una falla en el sistema de "
            "facturación electrónica homologado por la propia AFIP (RECE), que clasificó incorrectamente "
            "las operaciones como exentas. El software con el bug fue utilizado por 2.300 contribuyentes "
            "y la AFIP emitió una circular correctiva 8 meses después del período fiscalizado. La empresa "
            "pagó el impuesto omitido en cuanto fue intimada, sin resistencia."
        ),
        "partes": [
            {
                "nombre": "García Hermanos S.R.L.",
                "rol": "Actora",
                "pretension": (
                    "Nulidad de la multa por defraudación fiscal. Subsidiariamente, reducción a multa "
                    "por omisión culposa (art. 45 Ley 11.683). Reconocimiento del error excusable "
                    "por falla en sistema oficial de AFIP."
                ),
                "fundamentos": (
                    "Art. 45 y 46 Ley 11.683: la defraudación requiere dolo. La empresa actuó de buena "
                    "fe confiando en el sistema homologado por la AFIP. Error excusable (art. 45 in fine): "
                    "la falla del RECE es causa externa ajena a la empresa. Principio de culpabilidad "
                    "en derecho sancionatorio. La AFIP no puede sancionar con multa dolosa un error "
                    "causado por su propio sistema."
                ),
                "jurisprudencia": (
                    "CSJN, 'Lapiduz c/ DGI', Fallos 321:1043. CNCont. Adm. Fed. Sala III, 'Importadora "
                    "del Sur c/ AFIP', 2023. CSJN, 'Parafina del Plata c/ DGI', Fallos 271:297."
                ),
            },
            {
                "nombre": "AFIP-DGI",
                "rol": "Demandado",
                "pretension": (
                    "Confirmación de la multa por defraudación. El contribuyente tenía la obligación "
                    "de verificar sus declaraciones independientemente del sistema de facturación."
                ),
                "fundamentos": (
                    "Art. 46 Ley 11.683: la configuración del tipo no requiere prueba del dolo subjetivo: "
                    "basta el dolo eventual. El contribuyente es responsable de sus declaraciones juradas "
                    "con independencia del software utilizado. La circular correctiva fue emitida dentro "
                    "de los plazos razonables. El error no fue detectado en 18 operaciones consecutivas, "
                    "lo que excluye el carácter excusable."
                ),
                "jurisprudencia": (
                    "CSJN, 'Casa Elen-Valmi c/ DGI', Fallos 321:842. TFN Sala A, 'Supermercados "
                    "Walmart c/ AFIP', 2022."
                ),
            },
        ],
        "cuestiones": [
            "¿El error en la declaración de IVA originado en una falla del sistema de facturación electrónica homologado por la propia AFIP constituye 'error excusable' que excluye la multa por defraudación del art. 46 Ley 11.683?",
            "¿La multa por defraudación fiscal requiere la acreditación de dolo subjetivo o es suficiente la configuración objetiva de la omisión cuando el contribuyente prueba buena fe y causa ajena?",
            "¿La responsabilidad del Estado por los errores de su propio sistema de facturación homologado puede invocarse como causal exculpatoria en materia de infracciones tributarias?",
        ],
    },
    {
        "titulo": "Caso 18 — Fundación Recuperar c/ Instituto Nacional del Cáncer (Acceso a medicamento oncológico)",
        "filename": "Caso18_Fundacion_c_INC.docx",
        "caratula": "Fundación Recuperar (en representación de pacientes oncológicos) c/ Instituto Nacional del Cáncer s/ amparo colectivo",
        "expediente": "FAM-31209/2024",
        "rama": "Civil",
        "tipo_proceso": "Amparo",
        "tribunal": "Juzgado Federal Civil N° 5 de Buenos Aires",
        "nivel": "N2 — Difícil",
        "hechos": (
            "El Instituto Nacional del Cáncer (INC) incluyó en su Formulario Terapéutico Nacional (FTN) "
            "el medicamento Pembrolizumab para tratamiento de cáncer de pulmón de células no pequeñas en "
            "segunda línea. Sin embargo, por restricciones presupuestarias, el INC distribuyó solo el 40% "
            "de las dosis requeridas durante 2024, generando una lista de espera de 1.847 pacientes. "
            "La Fundación Recuperar representa a 213 pacientes en lista de espera, todos con pronóstico "
            "de sobrevida de 6 a 18 meses sin el tratamiento. El medicamento está aprobado por ANMAT. "
            "El INC argumenta que la asignación sigue criterios clínicos de priorización aprobados "
            "por el Comité de Ética."
        ),
        "partes": [
            {
                "nombre": "Fundación Recuperar (en representación de 213 pacientes)",
                "rol": "Actora",
                "pretension": (
                    "Provisión inmediata de Pembrolizumab a todos los pacientes en lista de espera. "
                    "Declaración de que la restricción presupuestaria no puede justificar la privación "
                    "de un tratamiento incluido en el FTN. Orden de partida presupuestaria urgente."
                ),
                "fundamentos": (
                    "Art. 42 CN: derecho a la salud. Ley 27.285 (Instituto Nacional del Cáncer): "
                    "obligación de provisión de medicamentos del FTN. PIDESC art. 12. La inclusión "
                    "en el FTN genera una obligación estatal de resultado. El criterio de priorización "
                    "no es válido cuando excluye al 60% de los pacientes con indicación médica."
                ),
                "jurisprudencia": (
                    "CSJN, 'Campodónico de Beviacqua', Fallos 323:3229. CSJN, 'Asociación Benghalensis "
                    "c/ Estado Nacional', Fallos 323:1339. Corte IDH, 'Poblete Vilches c/ Chile', 2018."
                ),
            },
            {
                "nombre": "Instituto Nacional del Cáncer",
                "rol": "Demandado",
                "pretension": (
                    "Rechazo del amparo. La distribución de medicamentos sigue criterios de priorización "
                    "clínica razonables. La ejecución presupuestaria depende del Congreso y el Ejecutivo."
                ),
                "fundamentos": (
                    "El INC actúa dentro de sus recursos asignados por ley de presupuesto. Los criterios "
                    "de priorización son médicamente fundados y éticamente validados. El poder judicial "
                    "no puede imponer gastos no previstos en el presupuesto nacional. La solución "
                    "estructural corresponde al proceso legislativo, no al amparo individual."
                ),
                "jurisprudencia": (
                    "CSJN, 'Badaro c/ ANSES', Fallos 329:3089 (límites control judicial gasto público). "
                    "CNCiv. Sala F, 'González c/ INC', 2023."
                ),
            },
        ],
        "cuestiones": [
            "¿La inclusión de un medicamento en el Formulario Terapéutico Nacional genera una obligación de resultado del Estado de suministrarlo a todos los pacientes con indicación médica, o admite limitaciones presupuestarias?",
            "¿Los criterios de priorización clínica del INC son compatibles con el derecho a la igualdad cuando excluyen al 60% de los pacientes con indicación médica validada?",
            "¿Puede el Poder Judicial ordenar la asignación de partidas presupuestarias para garantizar el acceso a medicamentos oncológicos incluidos en el formulario terapéutico oficial?",
        ],
    },
    {
        "titulo": "Caso 19 — Consorcio Torre Libertad c/ Propietario Airbnb (Conflicto de usos en propiedad horizontal)",
        "filename": "Caso19_Consorcio_c_PropietarioAirbnb.docx",
        "caratula": "Consorcio de Propietarios Torre Libertad c/ Della Valle, Marco Antonio s/ prohibición de uso y daños",
        "expediente": "JCC-44123/2024",
        "rama": "Civil",
        "tipo_proceso": "Ordinario",
        "tribunal": "Juzgado Civil N° 63 de la Ciudad Autónoma de Buenos Aires",
        "nivel": "N2 — Difícil",
        "hechos": (
            "El propietario del piso 12° B del edificio Torre Libertad alquila la unidad en la plataforma "
            "Airbnb en forma continua desde 2021, con una rotación de 4 a 8 huéspedes distintos por semana. "
            "El reglamento de copropiedad del edificio (año 2003) establece 'uso exclusivamente residencial "
            "permanente'. El consorcio demanda la cesación del uso turístico alegando violación del "
            "reglamento, deterioro de las partes comunes y inseguridad. El demandado sostiene que el alquiler "
            "temporario está expresamente regulado en el CCCN (arts. 1199 y 1223) y que el reglamento no "
            "puede prohibirlo. La Legislatura porteña aprobó en 2023 la Ley 6.255 que regula los alquileres "
            "turísticos pero no prohíbe la actividad en propiedad horizontal."
        ),
        "partes": [
            {
                "nombre": "Consorcio de Propietarios Torre Libertad",
                "rol": "Actor",
                "pretension": (
                    "Cese inmediato del uso turístico de la unidad. Daños y perjuicios por deterioro "
                    "de partes comunes. Multa diaria por incumplimiento (art. 2069 CCCN)."
                ),
                "fundamentos": (
                    "Art. 2056 CCCN: el reglamento puede establecer restricciones al uso de unidades. "
                    "La cláusula 'uso residencial permanente' excluye el alquiler turístico rotativo. "
                    "Art. 2069 CCCN: el consorcio puede solicitar al juez la prohibición. La "
                    "afectación a las partes comunes y la inseguridad son daños reales y documentados."
                ),
                "jurisprudencia": (
                    "CNCiv. Sala C, 'Consorcio Palermo Soho c/ Brizuela', 2023. CNCiv. Sala E, "
                    "'Consorcio Av. Santa Fe c/ García', 2022."
                ),
            },
            {
                "nombre": "Della Valle, Marco Antonio",
                "rol": "Demandado",
                "pretension": (
                    "Rechazo de la demanda. El alquiler temporario es una forma de uso residencial "
                    "legítimo y está expresamente contemplado en el CCCN."
                ),
                "fundamentos": (
                    "Arts. 1199 y 1223 CCCN: la locación con fines turísticos es un contrato típico "
                    "y lícito. El reglamento de 2003 no podía prever ni prohibir Airbnb. Las restricciones "
                    "al dominio deben interpretarse restrictivamente (art. 1884 CCCN). La Ley porteña "
                    "6.255 reguló y reconoció la actividad, sin prohibirla en PH."
                ),
                "jurisprudencia": (
                    "CNCiv. Sala K, 'Consorcio Retiro c/ López', 2023 (admitió Airbnb). "
                    "CSJN, 'Madorran c/ Administración Nacional de Aduanas', Fallos 330:1189."
                ),
            },
        ],
        "cuestiones": [
            "¿La cláusula de 'uso residencial permanente' en un reglamento de copropiedad anterior al CCCN prohíbe el alquiler temporario turístico tipo Airbnb?",
            "¿El CCCN (art. 1223) que regula el contrato de locación con fines turísticos prevalece sobre el reglamento de copropiedad que establece uso residencial, o ambas normas son conciliables?",
            "¿El juez puede ordenar el cese definitivo del alquiler turístico en propiedad horizontal sin que el reglamento lo prohíba expresamente, con fundamento en la afectación a las partes comunes?",
        ],
    },
    {
        "titulo": "Caso 20 — Núñez c/ Banco Nación (Crédito UVA / Imprevisión)",
        "filename": "Caso20_Nunez_c_BancoNacion.docx",
        "caratula": "Núñez, Patricia Mónica c/ Banco de la Nación Argentina s/ revisión de contrato de mutuo hipotecario UVA",
        "expediente": "CCC-61934/2024",
        "rama": "Civil",
        "tipo_proceso": "Ordinario",
        "tribunal": "Juzgado Nacional en lo Civil N° 88",
        "nivel": "N2 — Difícil",
        "hechos": (
            "La actora tomó en 2017 un crédito hipotecario UVA por el equivalente a $1.800.000 (100.000 UVAs) "
            "para adquirir su vivienda única. El índice UVA se ajusta por el CER (Coeficiente de Estabilización "
            "de Referencia). Al momento del contrato, 1 UVA equivalía a $18. En 2024, 1 UVA equivale a $1.050, "
            "lo que elevó el capital adeudado a $105.000.000. La cuota mensual pasó de $14.000 a $820.000, "
            "representando el 68% del ingreso familiar. La actora solicitó la revisión contractual por "
            "imprevisión (art. 1091 CCCN) y la adecuación de las cuotas. El banco rechazó la solicitud "
            "invocando la naturaleza del sistema UVA aprobado por ley."
        ),
        "partes": [
            {
                "nombre": "Núñez, Patricia Mónica",
                "rol": "Actora",
                "pretension": (
                    "Revisión judicial del contrato de mutuo hipotecario UVA con adecuación de cuotas "
                    "al índice salarial (CVS) en lugar del CER. Suspensión de ejecución hipotecaria "
                    "hasta resolución definitiva."
                ),
                "fundamentos": (
                    "Art. 1091 CCCN: imprevisión por alteración extraordinaria de las circunstancias. "
                    "La inflación del 5.733% en 7 años excede cualquier previsión razonable de 2017. "
                    "Art. 14 bis CN: protección de la vivienda familiar. Ley 25.798 y 26.313 (créditos "
                    "hipotecarios): el Estado tuvo que intervenir en 2019 congelando UVAs, reconociendo "
                    "el desequilibrio sistémico. El aumento del 7.500% en la cuota viola el art. 1198 CC "
                    "y art. 1091 CCCN."
                ),
                "jurisprudencia": (
                    "CNCiv. Sala I, 'Acosta c/ BNA', 2023 (revisó crédito UVA). CNCiv. Sala H, "
                    "'Ferreyra c/ BBVA', 2023. CSJN, 'Rinaldi c/ Guzmán', Fallos 330:855."
                ),
            },
            {
                "nombre": "Banco de la Nación Argentina",
                "rol": "Demandado",
                "pretension": (
                    "Rechazo de la demanda. El sistema UVA fue aprobado por ley y el riesgo inflacionario "
                    "es inherente al contrato libremente asumido."
                ),
                "fundamentos": (
                    "Ley 27.271 (UVA): el sistema fue expresamente diseñado para indexar al costo de "
                    "construcción/CER. La actora aceptó libre y conscientemente el riesgo inflacionario. "
                    "El art. 1091 CCCN requiere que la alteración sea imprevisible: en Argentina la "
                    "inflación es estructuralmente previsible. La revisión judicial de créditos UVA "
                    "afectaría el sistema financiero y a miles de deudores/acreedores."
                ),
                "jurisprudencia": (
                    "CNCiv. Sala M, 'López c/ Banco Galicia', 2023 (rechazó imprevisión). "
                    "BCRA Comunicación A 6823/2019 (congelamiento transitorio UVA)."
                ),
            },
        ],
        "cuestiones": [
            "¿La inflación acumulada del 5.733% desde 2017 que elevó el capital adeudado en un crédito UVA en un 5.733% configura una 'alteración extraordinaria e imprevisible de las circunstancias' en los términos del art. 1091 CCCN?",
            "¿La naturaleza inflacionaria estructural de la economía argentina hace que la inflación sea siempre 'previsible', excluyendo la imprevisión, o la magnitud puede ser igualmente imprevisible?",
            "¿Puede el juez sustituir el índice CER por el CVS (índice salarial) en un contrato UVA cuando la cuota supera el 60% del ingreso familiar y afecta la vivienda única?",
        ],
    },
]

print(f'Creando {len(CASOS2)} archivos Word en: {CASOS2_DIR}')
for caso in CASOS2:
    filepath = os.path.join(CASOS2_DIR, caso['filename'])
    crear_word(caso, filepath)
print(f'\n10 archivos Word guardados en: {CASOS2_DIR}')
print('\nTODO COMPLETADO.')
print(f'  Casos 1-10 (Word): {CASOS1_DIR}')
print(f'  Casos 11-20 (Word): {CASOS2_DIR}')
print(f'  Resultados motor: {results_path}')
