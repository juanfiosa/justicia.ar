"""
MÓDULO 3: API + SERVIDOR
=========================
API FastAPI que conecta la interfaz del juez con el indexador y el generador.

Endpoints:
    GET  /                      → Interfaz web (frontend)
    POST /api/generar            → Genera sentencia a partir del caso
    POST /api/buscar-precedentes → Busca precedentes similares
    GET  /api/estadisticas       → Estadísticas del sistema
    GET  /api/ramas              → Lista ramas del derecho disponibles
    GET  /api/health             → Health check
"""

import os
import sys
import json
import hmac
import hashlib
import time
import threading
import secrets
from collections import defaultdict
from pathlib import Path

# Agregar directorios al path
ROOT_DIR = str(Path(__file__).parent.parent.parent)
sys.path.insert(0, ROOT_DIR)
sys.path.insert(0, str(Path(__file__).parent.parent))

from fastapi import FastAPI, Request, UploadFile, File, Response, Form
from fastapi.responses import HTMLResponse, JSONResponse, RedirectResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel, Field
from typing import Optional
import io

from generador.motor_sentencias import generar, GeneradorSentencias, CasoInput
from generador.prompts import _extraer_datos_sentencia_grado as _extraer_datos_sentencia_inline
from generador.verificador import verificar_sentencia
from generador.clasificador import clasificar_caso
from arquitectura.ramas_derecho import RAMAS_DERECHO, normalizar_materia_saij

# Intentar importar indexadores (solo si hay datos locales disponibles)
# En deploy cloud (Render) no hay ChromaDB data, así que se deshabilita para ahorrar RAM
MODO_CLOUD = os.environ.get('RENDER', '') or os.environ.get('MODO_CLOUD', '')

indexador = None
INDEXADOR_DISPONIBLE = False
indexador_infoleg = None
INFOLEG_DISPONIBLE = False

# Intentar indexadores Qdrant (funcionan tanto local como en cloud)
try:
    from indexador.pipeline_qdrant import IndexadorJurisprudenciaQdrant
    indexador = IndexadorJurisprudenciaQdrant()
    INDEXADOR_DISPONIBLE = True
    print("Indexador jurisprudencia Qdrant: OK")
except Exception as e:
    print(f"ADVERTENCIA: Indexador jurisprudencia no disponible ({e}).")
    # Fallback a ChromaDB local si no está en cloud
    if not MODO_CLOUD:
        try:
            from indexador.pipeline import IndexadorJurisprudencia
            indexador = IndexadorJurisprudencia()
            INDEXADOR_DISPONIBLE = True
        except Exception as e2:
            print(f"ADVERTENCIA: ChromaDB tampoco disponible ({e2}).")

try:
    from indexador.infoleg_qdrant import IndexadorInfoLegQdrant
    indexador_infoleg = IndexadorInfoLegQdrant()
    INFOLEG_DISPONIBLE = True
    print("Indexador InfoLeg Qdrant: OK")
except Exception as e:
    print(f"ADVERTENCIA: Indexador InfoLeg no disponible ({e}).")
    if not MODO_CLOUD:
        try:
            from indexador.infoleg import IndexadorInfoLeg
            indexador_infoleg = IndexadorInfoLeg()
            INFOLEG_DISPONIBLE = True
        except Exception as e2:
            print(f"ADVERTENCIA: InfoLeg ChromaDB tampoco disponible ({e2}).")

# Indexador legislacion provincial (Qdrant)
indexador_provincial = None
PROVINCIAL_DISPONIBLE = False
try:
    from indexador.provincial_qdrant import IndexadorProvincialQdrant
    indexador_provincial = IndexadorProvincialQdrant()
    PROVINCIAL_DISPONIBLE = True
    print("Indexador legislacion provincial Qdrant: OK")
except Exception as e:
    print(f"ADVERTENCIA: Indexador provincial no disponible ({e}).")


# ================================================================
# MODELOS PYDANTIC (validación de entrada)
# ================================================================

class ParteInput(BaseModel):
    rol: str = Field(..., description="actor, demandado, querellante, defensa")
    nombre: str
    letrado: str = ""
    pretension: str = ""
    fundamentos_juridicos: str = ""
    argumentos: str = ""
    jurisprudencia_citada: str = ""


class CasoRequest(BaseModel):
    caratula: str = Field(..., description="Carátula del expediente")
    expediente: str = ""
    rama: str = Field("civil_comercial", description="Rama del derecho")
    subrama: str = ""
    tipo_proceso: str = "ordinario"
    tribunal: str = ""
    jurisdiccion: str = "Nacional"
    instancia: str = "primera"
    hechos_probados: str = Field(..., description="Hechos que el juez tiene por probados")
    prueba_valorada: str = ""   # campo legacy — usar valoracion_prueba
    valoracion_prueba: str = Field(
        "",
        description=(
            "Lectura del juez sobre la prueba producida. Opcional, solo para primera instancia. "
            "No es el listado de pruebas sino la conclusión del juez sobre lo que cada una acredita "
            "y cómo resuelve los hechos controvertidos. "
            "Ej: 'Las cámaras muestran que fue Villalba quien sacó el cuchillo primero. "
            "El testigo Ferreyra fue categórico. Tengo por no acreditada la legítima defensa.'"
        )
    )
    partes: list[ParteInput] = Field(default_factory=list)
    cuestiones_a_resolver: list[str] = Field(default_factory=list)
    hay_cuestion_constitucional: bool = False
    descripcion_cuestion_constitucional: str = ""
    nivel_complejidad: Optional[int] = Field(None, ge=1, le=3, description="Si es None, el sistema clasifica automáticamente. 1=Rutinario, 2=Difícil, 3=Constitucional")
    # Segunda instancia (Cámara)
    sentencia_primera_instancia: str = ""
    agravios: str = ""
    # TSJ / Casación provincial
    tipo_recurso_casacion: str = ""   # sustancial | formal | inaplicabilidad | nulidad
    causal_casacion: str = ""
    sentencia_recurrida: str = ""     # texto de la sentencia impugnada (TSJ o CSJN)
    recurso_casacion_texto: str = ""  # síntesis del recurso (casación o REF)
    contestacion_recurso: str = ""
    # CSJN — Recurso Extraordinario Federal
    tipo_cuestion_federal: str = ""   # simple | compleja_directa | compleja_indirecta | arbitrariedad | gravedad_institucional | per_saltum
    cuestion_federal: str = ""        # cuestión federal articulada
    es_queja: bool = False            # True si llega por queja art. 285
    introduccion_oportuna: str = ""   # cuándo y cómo se introdujo la cuestión federal
    # Campos penales
    imputado: str = ""
    delito_imputado: str = ""
    pena_solicitada: str = ""
    atenuantes: str = ""
    agravantes: str = ""
    prision_preventiva: str = ""
    magistrados: str = ""  # Nombres de los vocales separados por '; ' (para tribunal colegiado)
    # Tipo de órgano (para diseño institucional específico: juzgado_control, camara_acusacion, etc.)
    tipo_organo: str = ""


class ValoracionPruebaRequest(BaseModel):
    """Análisis de razonamiento probatorio para casación"""
    caratula: str = ""
    tribunal_recurrido: str = ""      # Tribunal cuya sentencia se impugna
    # Prueba producida y tenida en cuenta
    prueba_transcripta: str = ""      # Texto de la prueba: testimonios, pericias, docs
    # Razonamiento del tribunal a quo sobre esa prueba
    razonamiento_probatorio: str = "" # Qué dijo la sentencia sobre la prueba y cómo llegó al hecho
    hecho_probado: str = ""           # El hecho que el tribunal tuvo por acreditado
    # Posición del impugnante
    agravios_valoracion: str = ""     # Qué argumenta el impugnante sobre la prueba
    # Contexto procesal
    tipo_proceso: str = "penal"
    jurisdiccion: str = ""


class CoherenciaRequest(BaseModel):
    sentencia: str
    caratula: str = ""
    cuestiones_a_resolver: list[str] = Field(default_factory=list)


class VocalConfig(BaseModel):
    nombre: str
    posicion: str  # "adhesion" | "concurrencia" | "disidencia_parcial" | "disidencia_total"
    fundamento: str = ""  # En qué difiere (para concurrencia/disidencia)


class VotoRequest(BaseModel):
    caso: dict
    sentencia_ponencia: str   # El voto/sentencia ya generado (la ponencia)
    vocal: VocalConfig


class ExportarPDFRequest(BaseModel):
    texto: str
    caratula: str = ""
    tribunal: str = ""
    instancia: str = ""


class UsuarioLogin(BaseModel):
    tribunal: str
    password: str


class BusquedaRequest(BaseModel):
    query: str = Field(..., description="Texto de búsqueda")
    rama: Optional[str] = None
    n_results: int = Field(10, ge=1, le=50)
    jurisdiccion: Optional[str] = None
    incluir_legislacion: bool = Field(True, description="Incluir normas nacionales y provinciales en resultados")
    provincia: Optional[str] = Field(None, description="Filtrar legislación provincial por provincia")


class MensajeChat(BaseModel):
    rol: str  # "user" | "assistant"
    contenido: str


class ChatRequest(BaseModel):
    caso: dict
    sentencia_actual: str
    historial: list[MensajeChat] = Field(default_factory=list)
    mensaje: str


class RegenerarRequest(BaseModel):
    caso: dict
    sentencia_actual: str
    historial: list[MensajeChat] = Field(default_factory=list)


class PonenciaCreateRequest(BaseModel):
    """Crea una ponencia compartible para que otros vocales puedan votar."""
    caso: dict                # Datos del caso (CasoRequest dict)
    sentencia: dict           # Resultado completo de /api/generar
    vocal_ponente: str        # Nombre del vocal preopinante


class VotoColegiadoRequest(BaseModel):
    """Voto de un vocal sobre una ponencia ya generada."""
    tipo_voto: str            # "adhesion_total" | "adhesion_parcial" | "propio"
    nombre_vocal: str
    # Para adhesion_parcial:
    considerandos_propios: list[str] = []    # Secciones donde tiene voto propio ej: ["III","VI"]
    fundamento_disidencia: str = ""          # Descripción breve de la divergencia


# ================================================================
# PONENCIAS (store en memoria con TTL 24h)
# ================================================================

_ponencias: dict[str, dict] = {}      # token → registro
_PONENCIA_TTL = 24 * 3600             # 24 horas


def _limpiar_ponencias_vencidas() -> None:
    ahora = time.time()
    vencidas = [k for k, v in _ponencias.items() if ahora - v["created_at"] > _PONENCIA_TTL]
    for k in vencidas:
        del _ponencias[k]


def _parsear_considerandos(texto: str) -> dict[str, str]:
    """Divide el bloque de considerandos en secciones romanas."""
    import re
    # Captura numerales romanos seguidos de punto o punto y espacio al inicio de línea
    partes = re.split(r'\n((?:I|II|III|IV|V|VI|VII|VIII|IX|X)(?:\.|\.\s))', texto)
    secciones: dict[str, str] = {}
    for i in range(1, len(partes) - 1, 2):
        num = partes[i].rstrip('. ')
        cuerpo = partes[i + 1] if i + 1 < len(partes) else ''
        # Cortar en el próximo numeral (por si el split no lo tomó)
        cuerpo = re.split(r'\n(?:I|II|III|IV|V|VI|VII|VIII|IX|X)(?:\.|\.\s)', cuerpo)[0].strip()
        secciones[num] = cuerpo
    return secciones


def _vocales_sin_voto(registro: dict) -> list[str]:
    """Devuelve los vocales del caso que aún no emitieron voto."""
    caso = registro["caso"]
    mags = [m.strip() for m in caso.get("magistrados", "").split(";") if m.strip()]
    votados = {v["nombre"] for v in registro["votos"]}
    votados.add(registro["vocal_ponente"])
    return [m for m in mags if m not in votados]


# ================================================================
# RATE LIMITER (sin dependencias externas)
# ================================================================

class _RateLimiter:
    """Ventana deslizante por IP. Thread-safe."""
    def __init__(self, max_requests: int, window_seconds: int):
        self.max_requests = max_requests
        self.window = window_seconds
        self._lock = threading.Lock()
        self._history: dict[str, list[float]] = defaultdict(list)

    def is_allowed(self, key: str) -> bool:
        now = time.monotonic()
        cutoff = now - self.window
        with self._lock:
            calls = self._history[key]
            # Descartar timestamps fuera de la ventana
            calls[:] = [t for t in calls if t > cutoff]
            if len(calls) >= self.max_requests:
                return False
            calls.append(now)
            return True

# 10 generaciones por IP por minuto — suficiente para uso normal, bloquea bucles accidentales
_generar_limiter = _RateLimiter(max_requests=10, window_seconds=60)


# ================================================================
# SISTEMA DE ACCESO POR INVITACIÓN
# ================================================================

# Códigos válidos: variable de entorno INVITE_CODES=GARCIA-2026,MARTINEZ-2026,...
# Si no está configurada, el sistema queda abierto (útil para desarrollo local)
_INVITE_CODES_RAW = os.environ.get('INVITE_CODES', '')
INVITE_CODES: set[str] = (
    {c.strip().upper() for c in _INVITE_CODES_RAW.split(',') if c.strip()}
    if _INVITE_CODES_RAW else set()
)
AUTH_HABILITADO = bool(INVITE_CODES)

# Clave para firmar cookies de sesión (generar una fija en Render con SESSION_SECRET)
SESSION_SECRET = os.environ.get('SESSION_SECRET', secrets.token_hex(32))
SESSION_COOKIE = 'justia_session'
SESSION_DAYS   = 30

# Registro en memoria de accesos (código → lista de {ip, timestamp})
_accesos: dict[str, list[dict]] = defaultdict(list)


def _firmar(payload: str) -> str:
    """HMAC-SHA256 del payload con SESSION_SECRET."""
    return hmac.new(SESSION_SECRET.encode(), payload.encode(), hashlib.sha256).hexdigest()


def _crear_cookie(codigo: str) -> str:
    """Genera el valor de la cookie de sesión."""
    ts = str(int(time.time()))
    sig = _firmar(f"{codigo}:{ts}")
    return f"{codigo}:{ts}:{sig}"


def _validar_cookie(cookie_val: str) -> str | None:
    """Devuelve el código si la cookie es válida y no expiró, o None."""
    try:
        codigo, ts, sig = cookie_val.split(':', 2)
    except ValueError:
        return None
    if not hmac.compare_digest(sig, _firmar(f"{codigo}:{ts}")):
        return None
    if time.time() - int(ts) > SESSION_DAYS * 86400:
        return None
    return codigo.upper()


def _esta_autenticado(request: Request) -> bool:
    """True si el request tiene una sesión válida (o si auth está deshabilitado)."""
    if not AUTH_HABILITADO:
        return True
    cookie = request.cookies.get(SESSION_COOKIE, '')
    return _validar_cookie(cookie) is not None


# HTML de la página de login
_LOGIN_HTML = """<!DOCTYPE html>
<html lang="es">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>JUSTIA — Acceso</title>
    <style>
        * {{ margin:0; padding:0; box-sizing:border-box; }}
        body {{
            font-family: Georgia, 'Times New Roman', serif;
            background: #f7fafc;
            min-height: 100vh;
            display: flex;
            flex-direction: column;
            align-items: center;
            justify-content: center;
        }}
        .card {{
            background: white;
            border-radius: 12px;
            box-shadow: 0 8px 32px rgba(0,0,0,0.12);
            padding: 2.5rem 2rem;
            width: 100%;
            max-width: 400px;
            text-align: center;
        }}
        .escudo {{ font-size: 2.5rem; margin-bottom: 0.5rem; }}
        h1 {{
            color: #1a365d;
            font-size: 1.6rem;
            letter-spacing: 3px;
            margin-bottom: 0.3rem;
        }}
        .subtitulo {{
            color: #4a5568;
            font-size: 0.85rem;
            font-style: italic;
            margin-bottom: 2rem;
        }}
        label {{
            display: block;
            text-align: left;
            font-size: 0.9rem;
            color: #2d3748;
            margin-bottom: 0.4rem;
            font-weight: bold;
        }}
        input[type=text] {{
            width: 100%;
            padding: 0.75rem 1rem;
            border: 2px solid #e2e8f0;
            border-radius: 6px;
            font-size: 1.1rem;
            letter-spacing: 2px;
            text-transform: uppercase;
            text-align: center;
            margin-bottom: 1.2rem;
            font-family: monospace;
            transition: border-color 0.2s;
        }}
        input[type=text]:focus {{
            outline: none;
            border-color: #2b6cb0;
        }}
        button {{
            width: 100%;
            padding: 0.85rem;
            background: #1a365d;
            color: white;
            border: none;
            border-radius: 6px;
            font-size: 1rem;
            font-family: Georgia, serif;
            letter-spacing: 1px;
            cursor: pointer;
            transition: background 0.2s;
        }}
        button:hover {{ background: #2b6cb0; }}
        .error {{
            background: #fff5f5;
            border: 1px solid #fc8181;
            color: #c53030;
            padding: 0.7rem 1rem;
            border-radius: 6px;
            font-size: 0.9rem;
            margin-bottom: 1rem;
        }}
        .footer {{
            margin-top: 1.5rem;
            font-size: 0.75rem;
            color: #a0aec0;
        }}
    </style>
</head>
<body>
    <div class="card">
        <div class="escudo">⚖️</div>
        <h1>JUSTIA</h1>
        <p class="subtitulo">Sistema de Asistencia Judicial con IA</p>
        {error_html}
        <form method="POST" action="/api/login">
            <label>Código de acceso</label>
            <input type="text" name="codigo" placeholder="XXXX-0000"
                   autocomplete="off" autofocus maxlength="30">
            <button type="submit">INGRESAR</button>
        </form>
        <p class="footer">Acceso restringido a magistrados y funcionarios autorizados.<br>
        Si no tiene código, solicítelo al administrador del sistema.</p>
    </div>
</body>
</html>"""


# ================================================================
# APP FASTAPI
# ================================================================

app = FastAPI(
    title="JUSTICIA ARGENTINA",
    description="Sistema de Resolución Asistida por IA para Jueces Argentinos",
    version="0.2.0"
)


# ================================================================
# ENDPOINTS
# ================================================================

@app.api_route("/api/health", methods=["GET", "HEAD"])
async def health_check():
    gemini_ok = bool(os.environ.get('GEMINI_API_KEY', ''))
    anthropic_ok = bool(os.environ.get('ANTHROPIC_API_KEY', ''))
    return {
        "status": "healthy",
        "service": "JUSTICIA ARGENTINA API",
        "version": "0.2.0",
        "indexador_disponible": INDEXADOR_DISPONIBLE,
        "motor_ia": "gemini" if gemini_ok else ("anthropic" if anthropic_ok else "demo"),
        "gemini_configurado": gemini_ok,
        "anthropic_configurado": anthropic_ok,
    }


@app.get("/api/ramas")
async def listar_ramas():
    """Lista las ramas del derecho disponibles con sus subramas"""
    ramas = []
    for rama_id, rama in RAMAS_DERECHO.items():
        subramas = []
        for sub_id, sub in rama.get('subramas', {}).items():
            subramas.append({
                'id': sub_id,
                'nombre': sub['nombre'],
                'descripcion': sub.get('descripcion', '')
            })
        ramas.append({
            'id': rama_id,
            'nombre': rama['nombre'],
            'fueros': rama.get('fuero_judicial', []),
            'subramas': subramas
        })
    return {"ramas": ramas}


@app.post("/api/buscar-precedentes")
async def buscar_precedentes(req: BusquedaRequest):
    """Busca precedentes de jurisprudencia y normas aplicables"""
    respuesta = {"query": req.query, "jurisprudencia": [], "legislacion": []}

    # Jurisprudencia
    if INDEXADOR_DISPONIBLE:
        filtros = {}
        if req.jurisdiccion:
            filtros['jurisdiccion'] = req.jurisdiccion
        resultados = indexador.buscar_precedentes(
            query=req.query,
            rama=req.rama,
            n_results=req.n_results,
            filtros=filtros if filtros else None
        )
        respuesta["jurisprudencia"] = resultados

    # Legislación nacional
    if req.incluir_legislacion and INFOLEG_DISPONIBLE:
        try:
            normas_nac = indexador_infoleg.buscar_normas(
                query=req.query,
                rama=req.rama,
                n_results=5
            )
            respuesta["legislacion"].extend([
                {"fuente": "nacional", "texto": n} for n in normas_nac
            ])
        except Exception as e:
            print(f"Error buscando InfoLeg: {e}")

    # Legislación provincial
    if req.incluir_legislacion and PROVINCIAL_DISPONIBLE:
        try:
            normas_prov = indexador_provincial.buscar(
                query=req.query,
                provincia=req.provincia,
                n_results=5
            )
            for n in normas_prov:
                meta = n.get("metadata", {})
                respuesta["legislacion"].append({
                    "fuente": "provincial",
                    "provincia": meta.get("provincia", ""),
                    "tipo": meta.get("tipo", ""),
                    "nombre": meta.get("nombre", ""),
                    "texto": n.get("texto", "")[:500],
                    "score": n.get("score", 0),
                })
        except Exception as e:
            print(f"Error buscando provincial: {e}")

    respuesta["total"] = len(respuesta["jurisprudencia"]) + len(respuesta["legislacion"])
    # Compatibilidad hacia atrás
    respuesta["precedentes"] = respuesta["jurisprudencia"]
    return respuesta


@app.post("/api/clasificar")
async def clasificar_complejidad(req: CasoRequest):
    """Clasifica automáticamente el nivel de complejidad sin generar sentencia"""
    caso_dict = req.model_dump()
    caso_dict['partes'] = [p.model_dump() for p in req.partes]
    resultado = clasificar_caso(caso_dict)
    return resultado


def _limpiar_placeholders_resultado(resultado: dict, req: 'CasoRequest') -> None:
    """
    Post-procesamiento de la sentencia generada:
    1. Reemplaza placeholders [xxx], XXX, (a determinar), etc. con datos reales
    2. Valida que el VISTOS contenga los datos clave de la sentencia de grado
       y los inyecta si el LLM los inventó o los omitió
    Modifica el dict resultado en lugar (in-place).
    """
    import re
    from datetime import date

    # ── Datos del caso disponibles ──────────────────────────────────────────
    tribunal = req.tribunal or ''
    meses = ['enero','febrero','marzo','abril','mayo','junio',
             'julio','agosto','septiembre','octubre','noviembre','diciembre']
    d = date.today()
    hoy = f"{d.day} de {meses[d.month-1]} de {d.year}"

    # ── Extraer datos de la sentencia de grado ──────────────────────────────
    # Reutilizar la función ya definida en prompts.py si está disponible
    sentencia_grado = req.sentencia_primera_instancia or ''
    tribunal_grado = ''
    fecha_grado = ''
    pena_grado = ''
    nro_sentencia = ''

    if sentencia_grado:
        try:
            datos = _extraer_datos_sentencia_inline(sentencia_grado)
            tribunal_grado = datos.get("Tribunal que dictó la sentencia", '')
            fecha_grado    = datos.get("Fecha de la sentencia", '')
            pena_grado     = datos.get("Pena impuesta", '')
            nro_sentencia  = datos.get("Número de sentencia", '')
        except Exception:
            pass  # fallback: los campos quedan vacíos

    # ── Palabras que indican que un fragmento es placeholder ────────────────
    PALABRAS_PH = (r'nombre|fecha|pena|tribunal|grado|plazo|t[eé]rmino|'
                   r'readecua|sentencia|inhabilitaci|especificar|determinar|'
                   r'indicar|insertar|completar|xxx|n/a')

    # ── Paso 1: reemplazos específicos [xxx] ────────────────────────────────
    reemplazos = [
        (r'\[Nombre de la C[aá]mara[^\]]*\]',   tribunal or 'este Tribunal'),
        (r'\[Nombre del [Tt]ribunal[^\]]*\]',    tribunal or 'este Tribunal'),
        (r'\[Nombre del TSJ[^\]]*\]',            tribunal or 'este Tribunal'),
        (r'\[fecha\b[^\]]*\]',                   hoy),
        (r'\[Tribunal de grado\]',               tribunal_grado or 'el tribunal de grado'),
        (r'\[tribunal de grado\]',               tribunal_grado or 'el tribunal de grado'),
        (r'\[fecha de la sentencia[^\]]*\]',     fecha_grado or 'según consta en autos'),
        (r'\[pena[^\]]{0,60}\]',                 pena_grado or 'la pena impuesta en la sentencia apelada'),
        (r'\[t[eé]rmino de la inhabilitaci[oó]n[^\]]*\]', 'el plazo fijado en la sentencia de grado'),
        (r'\[plazo de inhabilitaci[oó]n[^\]]*\]',         'el plazo fijado en la sentencia de grado'),
        (r'\[inhabilitaci[oó]n[^\]]*\]',         'inhabilitación según consta en autos'),
        (r'\[\d+\]',                             ''),
    ]

    # ── Paso 2: patrones sin corchetes que también son placeholders ─────────
    # "XXX", "N/A", "(a determinar)", "(a especificar)", "fs. XXX"
    reemplazos_sin_corchetes = [
        (r'\bXXX\b',                             'según consta en autos'),
        (r'N/A\b',                               'según consta en autos'),
        (r'\(a\s+determinar\)',                  'según consta en autos'),
        (r'\(a\s+especificar\)',                 'según consta en autos'),
        (r'\(a\s+indicar\)',                     'según consta en autos'),
        (r'fs\.\s*XXX',                          'fs. que obran en autos'),
        (r'Nro\.\s*XXX',                         f'Nro. {nro_sentencia}' if nro_sentencia else 'según consta en autos'),
        (r'fecha\s+XXX',                         fecha_grado or 'según consta en autos'),
        # Typos comunes del LLM
        (r'\bIMONER\b',                          'IMPONER'),
        (r'\bIMPONOR\b',                         'IMPONER'),
        (r'\bRESOLVER\b(?=\s*:)',                'RESUELVO'),
    ]

    campos = ['vistos', 'considerandos', 'resuelvo', 'texto_completo']
    for campo in campos:
        texto = resultado.get(campo, '')
        if not texto:
            continue

        for patron, reemplazo in reemplazos:
            texto = re.sub(patron, reemplazo, texto, flags=re.IGNORECASE)

        for patron, reemplazo in reemplazos_sin_corchetes:
            texto = re.sub(patron, reemplazo, texto, flags=re.IGNORECASE)

        # Barrido genérico: cualquier [texto descriptivo largo]
        def _reemplazar_generico(m):
            contenido = m.group(1)
            if re.search(PALABRAS_PH, contenido, re.IGNORECASE) and len(contenido) > 5:
                return 'según consta en autos'
            return m.group(0)
        texto = re.sub(r'\[([^\]]{6,80})\]', _reemplazar_generico, texto, flags=re.IGNORECASE)

        resultado[campo] = texto

    # ── Paso 3: validar y corregir datos clave en el VISTOS ─────────────────
    # Solo aplicar si el VISTOS fue generado por LLM (no por Python — ya tiene datos correctos)
    if pena_grado and resultado.get('vistos') and not resultado.get('vistos_generado_python'):
        vistos = resultado['vistos']
        # Detectar mención de pena en el VISTOS (cualquier número de años)
        pena_patron = re.compile(
            r'((?:UN|UNO|DOS|TRES|CUATRO|CINCO|SEIS|SIETE|OCHO|NUEVE|DIEZ|\d+)'
            r'(?:\s*\(\d+\))?\s*(?:a[ñn]o[s]?|mes(?:es)?)\s+de\s+prisi[oó]n'
            r'(?:[^\.]{0,60})?)',
            re.IGNORECASE
        )
        pena_en_vistos = pena_patron.search(vistos)
        if pena_en_vistos:
            pena_generada = pena_en_vistos.group(1).strip()
            # Normalizar para comparar (minúsculas, sin espacios extras)
            def _norm(s):
                return re.sub(r'\s+', ' ', s.lower().strip())
            if _norm(pena_generada) != _norm(pena_grado):
                # El LLM inventó una pena diferente → reemplazar
                vistos_corregido = pena_patron.sub(pena_grado, vistos, count=1)
                resultado['vistos'] = vistos_corregido
                # Actualizar texto_completo también
                if resultado.get('texto_completo'):
                    resultado['texto_completo'] = resultado['texto_completo'].replace(
                        pena_generada, pena_grado, 1
                    )
                resultado['pena_corregida'] = True


@app.post("/api/generar")
async def generar_sentencia(req: CasoRequest, request: Request):
    """Genera un proyecto de sentencia para el caso dado"""

    # Rate limiting por IP
    client_ip = request.client.host if request.client else "unknown"
    if not _generar_limiter.is_allowed(client_ip):
        return JSONResponse(
            status_code=429,
            content={"error": "Demasiadas solicitudes. Espere un momento antes de generar otra sentencia."}
        )

    # Construir query de búsqueda
    query_busqueda = f"{req.hechos_probados} {' '.join(req.cuestiones_a_resolver)}"
    # Agregar fundamentos de las partes para búsqueda más rica
    for p in req.partes:
        if p.fundamentos_juridicos:
            query_busqueda += f" {p.fundamentos_juridicos}"

    # Buscar precedentes de jurisprudencia
    precedentes = []
    if INDEXADOR_DISPONIBLE:
        resultados_raw = indexador.buscar_precedentes(
            query=query_busqueda,
            rama=req.rama,
            n_results=6
        )
        for r in resultados_raw:
            precedentes.append({
                'caratula': r['metadata'].get('caratula', 'N/A'),
                'tribunal': r['metadata'].get('tribunal_tipo', 'N/A'),
                'fecha': r['metadata'].get('fecha', 'N/A'),
                'texto': r['texto'],
                'relevancia': 1.0 - r.get('distancia', 0.5)
            })

    # Buscar normas aplicables: InfoLeg nacional + legislacion provincial
    normas_contexto = []

    # 1. Normas nacionales (InfoLeg)
    if INFOLEG_DISPONIBLE:
        try:
            normas_raw = indexador_infoleg.buscar_normas(
                query=query_busqueda,
                rama=req.rama,
                n_results=4
            )
            normas_contexto.extend(normas_raw)
        except Exception as e:
            print(f"Error buscando normas InfoLeg: {e}")

    # 2. Legislacion provincial — prioriza la provincia del caso si está definida
    if PROVINCIAL_DISPONIBLE:
        try:
            # Inferir provincia del campo jurisdiccion o tribunal
            provincia_filtro = None
            jurisdiccion = req.jurisdiccion or ""
            # Si dice "Provincial" o nombra una provincia, filtrar
            PROVINCIAS_CONOCIDAS = [
                "Buenos Aires", "Córdoba", "Santa Fe", "Mendoza", "Tucumán",
                "Entre Ríos", "Salta", "Misiones", "Chaco", "Corrientes",
                "Santiago del Estero", "San Juan", "Jujuy", "Río Negro",
                "Neuquén", "Formosa", "Chubut", "San Luis", "Catamarca",
                "La Rioja", "La Pampa", "Santa Cruz", "Tierra del Fuego",
                "Ciudad Autónoma de Buenos Aires",
            ]
            for prov in PROVINCIAS_CONOCIDAS:
                if prov.lower() in jurisdiccion.lower() or prov.lower() in (req.tribunal or "").lower():
                    provincia_filtro = prov
                    break

            normas_prov = indexador_provincial.buscar(
                query=query_busqueda,
                provincia=provincia_filtro,
                n_results=3
            )
            # Formatear igual que normas nacionales
            for n in normas_prov:
                meta = n.get("metadata", {})
                texto_norma = (
                    f"[{meta.get('provincia','Provincial')}] "
                    f"{meta.get('nombre', meta.get('titulo','Norma'))}: "
                    f"{n.get('texto','')[:500]}"
                )
                normas_contexto.append(texto_norma)
        except Exception as e:
            print(f"Error buscando normas provinciales: {e}")

    # Convertir a dict para el generador
    caso_dict = req.model_dump()
    caso_dict['partes'] = [p.model_dump() for p in req.partes]

    # Clasificación automática si el juez no especificó nivel
    clasificacion = None
    if req.nivel_complejidad is None:
        clasificacion = clasificar_caso(caso_dict)
        caso_dict['nivel_complejidad'] = clasificacion['nivel']
    else:
        caso_dict['nivel_complejidad'] = req.nivel_complejidad

    # Generar sentencia
    resultado = generar(
        caso_dict=caso_dict,
        precedentes=precedentes,
        normas=normas_contexto,
        api_key=os.environ.get('ANTHROPIC_API_KEY', '')
    )

    resultado['precedentes_usados'] = precedentes
    if clasificacion:
        resultado['clasificacion_automatica'] = clasificacion

    # Reemplazar VISTOS con versión generada en Python (datos exactos, sin alucinaciones)
    if req.instancia in ('primera', 'segunda', 'casacion', 'extraordinaria'):
        try:
            vistos_python = _generar_vistos(req)
        except Exception as e_vistos:
            vistos_python = ''
            resultado['vistos_error'] = str(e_vistos)
            print(f"[VISTOS] Error generando VISTOS Python: {e_vistos}")
        if vistos_python:
            resultado['vistos'] = vistos_python
            resultado['vistos_generado_python'] = True
            # Reconstruir texto_completo con el vistos correcto
            tc = resultado.get('texto_completo', '')
            idx_consid = tc.upper().find('Y CONSIDERANDO')
            if idx_consid == -1:
                idx_consid = tc.upper().find('CONSIDERANDO')
            if idx_consid != -1:
                resultado['texto_completo'] = (
                    'VISTOS:\n\n' + vistos_python + '\n\nY\n\n' + tc[idx_consid:]
                )

    # Post-procesamiento: reemplazar placeholders [xxx] que el LLM puede generar
    _limpiar_placeholders_resultado(resultado, req)

    # Corrección de coherencia entre CONSIDERANDOS y RESUELVO
    if resultado.get('considerandos') and resultado.get('resuelvo'):
        resuelvo_corregido = _corregir_coherencia_resuelvo(
            considerandos=resultado['considerandos'],
            resuelvo=resultado['resuelvo'],
            caratula=req.caratula,
            tribunal=req.tribunal or '',
            instancia=req.instancia or '',
            agravios=req.agravios or ''
        )
        if resuelvo_corregido != resultado['resuelvo']:
            resultado['resuelvo'] = resuelvo_corregido
            # Reconstruir texto_completo con el resuelvo corregido
            tc = resultado.get('texto_completo', '')
            idx = tc.upper().find('RESUELVO')
            if idx != -1:
                resultado['texto_completo'] = tc[:idx] + resuelvo_corregido
            resultado['resuelvo_corregido'] = True

    # Verificar citas normativas en la sentencia generada
    texto_sentencia = resultado.get('texto_completo', '')
    if texto_sentencia:
        try:
            resultado['verificacion'] = verificar_sentencia(
                texto_sentencia,
                indexador_infoleg=indexador_infoleg if INFOLEG_DISPONIBLE else None,
                indexador_jurisprudencia=indexador if INDEXADOR_DISPONIBLE else None
            )
        except Exception as e:
            print(f"Error en verificación: {e}")
            resultado['verificacion'] = None

    return resultado


# ================================================================
# GENERACIÓN PYTHON DEL VISTOS (datos exactos, sin LLM)
# ================================================================

def _generar_vistos(req: 'CasoRequest') -> str:
    """
    Genera la sección VISTOS en Python con datos exactos del caso.
    Elimina las alucinaciones del LLM en el encabezado formal.
    Solo para instancias de alzada (segunda, casación, extraordinaria).
    """
    import re
    from datetime import date

    meses = ['enero','febrero','marzo','abril','mayo','junio',
             'julio','agosto','septiembre','octubre','noviembre','diciembre']
    d = date.today()
    hoy = f"{d.day} de {meses[d.month-1]} de {d.year}"

    caratula = req.caratula or ''
    expediente = req.expediente or ''
    tribunal = req.tribunal or 'este Tribunal'
    instancia = req.instancia or ''

    # ── Extraer datos de la sentencia recurrida (inline, sin imports externos) ──
    if instancia == 'segunda':
        sentencia_src = req.sentencia_primera_instancia or ''
    else:
        sentencia_src = req.sentencia_recurrida or ''

    datos = _extraer_datos_sentencia_inline(sentencia_src) if sentencia_src else {}
    tribunal_grado = datos.get("Tribunal que dictó la sentencia", '') or 'el tribunal de grado'
    fecha_grado    = datos.get("Fecha de la sentencia", '') or 'la fecha que consta en autos'
    nro_sentencia  = datos.get("Número de sentencia", '')
    pena_grado     = datos.get("Pena impuesta", '')
    inhabilitacion = datos.get("Inhabilitación impuesta", '')
    calificacion   = datos.get("Calificación legal", '')

    # ── Identificar partes ───────────────────────────────────────────────────
    recurrente_nombre = ''
    recurrente_rol = ''
    recurrido_nombre = ''
    for p in req.partes:
        rol = p.rol.lower() if hasattr(p, 'rol') else p.get('rol', '').lower()
        nombre = p.nombre if hasattr(p, 'nombre') else p.get('nombre', '')
        if rol in ('defensa', 'imputado', 'apelante', 'recurrente'):
            recurrente_nombre = nombre
            recurrente_rol = 'la defensa'
        elif rol in ('fiscal', 'fiscalia', 'ministerio público', 'recurrido'):
            recurrido_nombre = nombre

    imputado = req.imputado or ''
    delito = req.delito_imputado or calificacion or ''

    # ── Construir referencia a la sentencia apelada ──────────────────────────
    ref_sentencia = f'la Sentencia'
    if nro_sentencia:
        ref_sentencia += f' Nro. {nro_sentencia}'
    ref_sentencia += f' de fecha {fecha_grado}'
    ref_sentencia += f', dictada por el {tribunal_grado}'

    # Nombre corto del imputado (solo el nombre, no la biografía completa)
    # req.imputado puede ser "Carlos Alberto Vega, 45 años, ..." → tomar solo hasta la primera coma
    imputado_nombre = imputado.split(',')[0].strip() if imputado else 'el/la imputado/a'

    # ── Construir descripción de la condena ──────────────────────────────────
    partes_condena = []
    if pena_grado:
        partes_condena.append(f'la pena de {pena_grado}')
    if inhabilitacion:
        partes_condena.append(inhabilitacion)
    condena_str = ' y '.join(partes_condena) if partes_condena else ''

    condena_txt = ''
    if condena_str:
        condena_txt = f', mediante la cual se condenó a {imputado_nombre} a {condena_str}'
    if delito:
        condena_txt += f', como autor/a penalmente responsable del delito de {delito}'
    if condena_txt:
        condena_txt += '; y'

    # ── Síntesis de los agravios ─────────────────────────────────────────────
    agravios_raw = req.agravios or ''
    # Contar agravios: detecta "1)" o "1." al inicio de línea O después de espacio/punto (inline)
    n_agravios = len(re.findall(r'(?:^|[\s\.\;])\s*\d+[\)\.]\s', agravios_raw, re.MULTILINE))
    if n_agravios > 1:
        agravios_resumen = f'{n_agravios} agravios'
    elif agravios_raw:
        agravios_resumen = 'el agravio que surge de la posición de la recurrente'
    else:
        agravios_resumen = 'los agravios articulados'

    # ── Armado del VISTOS ────────────────────────────────────────────────────
    if instancia == 'primera':
        es_penal = (req.rama or '').lower() == 'penal'
        tipo_proceso_txt = req.tipo_proceso or 'proceso ordinario'
        exp_txt = f', Expediente Nro. {expediente}' if expediente else ''

        if es_penal:
            # Partes penales
            imputado_txt = imputado_nombre or 'el/la imputado/a'
            delito_txt = req.delito_imputado or calificacion or ''
            delito_frase = f' por el delito de {delito_txt}' if delito_txt else ''

            # Extraer magistrados si está disponible (para tribunal colegiado)
            magistrados_str = getattr(req, 'magistrados', '') or ''
            magistrados_lista = [m.strip() for m in magistrados_str.split(';') if m.strip()] if magistrados_str else []


            # Formato diferenciado: colegiado vs unipersonal
            if len(magistrados_lista) >= 2:
                # Tribunal colegiado: formato formal con Sentencia Número y Vocales
                vocales_str = '; '.join(magistrados_lista[:3])
                vistos = (
                    f'SENTENCIA NÚMERO [número correlativo]\n\n'
                    f'En la ciudad de [ciudad], a los {hoy}, reunido en Acuerdo el tribunal\n'
                    f'"{tribunal}", integrado por los señores Vocales {vocales_str},\n'
                    f'a fin de dictar sentencia en la causa caratulada "{caratula}",\n'
                    f'Expediente Nro. {expediente}, elevada a esta sede para la celebración del juicio\n'
                    f'oral y público que tuvo lugar en audiencia(s) del/los día(s) [fecha(s) del debate].\n'
                    f'Intervino el Ministerio Público Fiscal y la defensa técnica del imputado\n'
                    f'{imputado_txt} a cargo del Dr./Dra. [nombre defensor]. Concluida la recepción de\n'
                    f'prueba, los alegatos de las partes y la última palabra del imputado, el tribunal\n'
                    f'deliberó y emitió veredicto, encontrándose en estado de dictar sentencia definitiva.\n\n'
                    f'VISTOS: [Síntesis de lo actuado en el debate: prueba testimonial — quiénes declararon\n'
                    f'y puntos clave de sus testimonios —; prueba pericial — qué pericias se rindieron y sus\n'
                    f'conclusiones —; prueba documental e incorporada por lectura; alegatos del Ministerio\n'
                    f'Público Fiscal, la querella si la hubiere, y la defensa.]'
                )
            else:
                # Unipersonal o sin magistrados: formato simple
                vistos = (
                    f'Que en las presentes actuaciones caratuladas "{caratula}"{exp_txt}'
                    f', tramitadas ante el {tribunal} por el procedimiento de {tipo_proceso_txt},'
                    f' se llevó a cabo el debate oral y público en relación a {imputado_txt}'
                    f'{delito_frase}.'
                    f'\n\nConcluida la audiencia de debate, deliberado el veredicto'
                    f' y encontrándose la causa en estado de dictar sentencia, el tribunal RESUELVE:'
                )
        else:
            # Partes civiles/comerciales
            actores = []
            demandados = []
            for p in req.partes:
                rol = (p.rol if hasattr(p, 'rol') else p.get('rol', '')).lower()
                nombre = p.nombre if hasattr(p, 'nombre') else p.get('nombre', '')
                if rol in ('actor', 'actora', 'demandante', 'querellante'):
                    actores.append(nombre)
                elif rol in ('demandado', 'demandada', 'accionado'):
                    demandados.append(nombre)
            actores_txt = ' y '.join(actores) if actores else 'la parte actora'
            demandados_txt = ' y '.join(demandados) if demandados else 'la parte demandada'
            subrama = req.subrama or ''
            objeto = f' en concepto de {subrama}' if subrama else ''
            vistos = (
                f'La demanda promovida por {actores_txt} contra {demandados_txt}{objeto},'
                f' tramitada ante el {tribunal} bajo el Expediente Nro. {req.expediente or "s/n"}'
                f' por el procedimiento {tipo_proceso_txt};\n\n'
                f'Sustanciado el proceso con el trámite de ley, producida la prueba ofrecida'
                f' por las partes y encontrándose los autos en estado de dictar sentencia, el/la'
                f' titular del {tribunal} DISPONE:'
            )
        exp_txt = ''  # ya incluido en el cuerpo
        encabezado = ''
        return vistos

    elif instancia == 'segunda':
        letra_recurrente = recurrente_rol or 'la parte recurrente'
        letra_letrado = f', patrocinado/a por {recurrente_nombre},' if recurrente_nombre else ','
        vistos = (
            f'El recurso de apelación interpuesto por {letra_recurrente}'
            f'{letra_letrado} articulando {agravios_resumen},'
            f' contra {ref_sentencia}'
            f'{condena_txt}'
        )
        if recurrido_nombre:
            vistos += (
                f'\n\nLa parte recurrida, representada por {recurrido_nombre},'
                f' solicitó la confirmación íntegra del fallo.'
            )

    elif instancia == 'casacion':
        tipo_recurso = req.tipo_recurso_casacion or ''
        mapa_tipo = {
            'sustancial': 'casación sustancial (error in iudicando)',
            'formal': 'casación formal (error in procedendo)',
            'inaplicabilidad': 'inaplicabilidad de ley',
            'nulidad': 'nulidad'
        }
        tipo_label = mapa_tipo.get(tipo_recurso, tipo_recurso or 'casación')
        vistos = (
            f'El recurso de {tipo_label} interpuesto contra {ref_sentencia}'
            f'{condena_txt}'
        )

    elif instancia == 'extraordinaria':
        vistos = (
            f'El recurso extraordinario federal interpuesto contra {ref_sentencia}'
            f'{condena_txt}'
        )

    else:
        return ''

    # ── Encabezado con carátula y expediente ─────────────────────────────────
    exp_txt = f', Expediente Nro. {expediente}' if expediente else ''
    encabezado = f'Autos caratulados "{caratula}"{exp_txt}.\n\n'
    return encabezado + vistos


# ================================================================
# CORRECCIÓN DE COHERENCIA ENTRE CONSIDERANDOS Y RESUELVO
# ================================================================

_PROMPT_COHERENCIA_RESUELVO_SISTEMA = """Sos un revisor jurídico experto en derecho argentino.
Tu única tarea es reescribir la sección RESUELVO de una sentencia para que sea
estrictamente coherente con las conclusiones alcanzadas en los CONSIDERANDOS.

REGLAS ABSOLUTAS:
1. Leé los CONSIDERANDOS completos e identificá la conclusión sobre CADA agravio:
   si el agravio prospera, se rechaza, se admite parcialmente, etc.
2. Cada conclusión DEBE tener un punto numerado en el RESUELVO.
   Si hay 3 agravios resueltos en considerandos → 3 puntos en RESUELVO.
3. Prestá especial atención a:
   - Agravios sobre la pena de prisión (confirmar / revocar / reducir)
   - Agravios sobre inhabilitación (confirmar / revocar / reducir / suprimir)
   - Agravios sobre calificación legal (confirmar / modificar)
   - Costas (siempre incluir un punto de costas)
4. El RESUELVO no puede decir algo que los considerandos no decidieron.
5. Mantené el lenguaje jurídico formal argentino.
6. No uses corchetes []. Usá los datos concretos del caso.
7. Devolvé SOLO el texto del RESUELVO, empezando con "RESUELVO:" sin nada más antes.
"""


def _corregir_coherencia_resuelvo(considerandos: str, resuelvo: str,
                                   caratula: str, tribunal: str,
                                   instancia: str = '',
                                   agravios: str = '') -> str:
    """
    Para segunda instancia: siempre reescribe el RESUELVO desde los CONSIDERANDOS
    mediante una segunda llamada LLM focalizada (temperatura 0.1, < 1000 tokens).
    Para otras instancias: solo corrige cuando detecta incoherencia por regex.
    Devuelve el RESUELVO corregido (o el original si falla o es coherente).
    """
    import re
    from datetime import date as _date
    _meses = ['enero','febrero','marzo','abril','mayo','junio',
              'julio','agosto','septiembre','octubre','noviembre','diciembre']
    _d = _date.today()
    hoy_str = f"{_d.day} de {_meses[_d.month-1]} de {_d.year}"

    # Todas las instancias de alzada: reescritura siempre activa
    # 'primera' instancia: solo si regex detecta incoherencia (no tiene agravios formales)
    forzar_reescritura = instancia in ('segunda', 'casacion', 'extraordinaria')

    if not forzar_reescritura:
        # Detección rápida de incoherencia por patrones para instancias no-segunda
        patrones_conclusion = [
            (r'corresponde\s+revocar|debe\s+(?:ser\s+)?revocad', 'revocar'),
            (r'el\s+agravio\s+(?:debe\s+)?prospera|hace?\s+lugar\s+al\s+agravio', 'hacer_lugar'),
            (r'el\s+agravio\s+(?:no\s+)?(?:debe\s+)?(?:ser\s+)?rechazad|no\s+ha\s+de\s+prosperar', 'rechazar'),
            (r'corresponde\s+confirmar|confirmarse\s+la\s+sentencia', 'confirmar'),
            (r'corresponde\s+hacer\s+lugar|procede\s+hacer\s+lugar', 'hacer_lugar'),
            (r'se\s+impone\s+revocar|se\s+revoca', 'revocar'),
            (r'cabe\s+revocar|corresponde\s+anular', 'revocar'),
        ]
        conclusiones = set()
        for patron, tipo in patrones_conclusion:
            if re.search(patron, considerandos, re.IGNORECASE):
                conclusiones.add(tipo)

        incoherente = False
        if 'revocar' in conclusiones:
            if not re.search(r'revocar|revocase|anular', resuelvo, re.IGNORECASE):
                incoherente = True
        if 'hacer_lugar' in conclusiones and 'revocar' not in conclusiones:
            if not re.search(r'hacer\s+lugar|hacese\s+lugar|hace\s+lugar', resuelvo, re.IGNORECASE):
                incoherente = True

        if not incoherente:
            return resuelvo  # primera instancia coherente → no gastar llamada LLM

    # Segunda instancia (siempre) o incoherencia detectada en otra instancia:
    # reescribir el RESUELVO con una llamada LLM focalizada
    agravios_txt = (
        f"\nAgravios planteados (exactamente estos, ni más ni menos):\n{agravios}\n"
        if agravios else ''
    )
    prompt_correccion = f"""A continuación están los CONSIDERANDOS completos de una sentencia de {instancia or 'alzada'}:

{considerandos[-3500:]}
{agravios_txt}
El RESUELVO actual (puede ser incorrecto o incompleto):
{resuelvo}

TAREA: Reescribí el RESUELVO de cero siguiendo estas reglas ESTRICTAS:
1. El RESUELVO debe tener exactamente un punto numerado por cada agravio listado arriba,
   NI MÁS NI MENOS — no agregues agravios que no estén en la lista.
2. Cada punto usa la decisión de los considerandos para ese agravio:
   rechazado, admitido o admitido parcialmente.
3. Si los considerandos no analizaron un agravio, resolvelo como "Confirmar la sentencia
   de primera instancia en este punto" (decisión conservadora por defecto).
4. Agregar UN punto final de costas de alzada (último punto).
5. NO usar corchetes [] ni "según consta en autos" — usá los datos concretos del caso.
6. Firmar con el nombre completo del tribunal seguido de la fecha exacta indicada abajo.

Datos del caso:
  Carátula: {caratula}
  Tribunal que resuelve: {tribunal}
  Fecha de la resolución (usala literalmente en la firma): {hoy_str}
"""
    try:
        resuelvo_corregido = _gemini(
            _PROMPT_COHERENCIA_RESUELVO_SISTEMA,
            prompt_correccion,
            max_tokens=900,
            temperature=0.1
        )
        if not resuelvo_corregido or len(resuelvo_corregido.strip()) < 30:
            return resuelvo  # respuesta vacía o demasiado corta → usar original
        # Asegurar que empieza con RESUELVO:
        if 'RESUELVO' not in resuelvo_corregido.upper()[:20]:
            resuelvo_corregido = 'RESUELVO:\n\n' + resuelvo_corregido.strip()
        return resuelvo_corregido.strip()
    except Exception as e:
        print(f"[coherencia_resuelvo] Error en segunda llamada LLM: {e}")
        return resuelvo  # si falla, devolver original sin romper


# ================================================================
# HELPER GEMINI + GROQ CON FALLBACK REGIONAL (para todos los endpoints)
# ================================================================

_GEMINI_MODELS_FALLBACK = [
    'gemini-2.0-flash-lite',
    'gemini-2.0-flash',
    'gemini-2.5-flash-lite',
    'gemini-2.5-flash',
]

# Modelo Groq: LLaMA 3.3 70B — excelente para texto jurídico, sin restricción geográfica
_GROQ_MODEL = 'llama-3.3-70b-versatile'


def _groq(system: str, prompt: str, max_tokens: int = 1500,
          temperature: float = 0.2) -> str:
    """Llama a Groq (LLaMA 3.3 70B) como proveedor alternativo cuando Gemini no está disponible."""
    try:
        from groq import Groq
    except ImportError:
        raise RuntimeError("Módulo 'groq' no instalado.")

    api_key = os.environ.get('GROQ_API_KEY', '')
    if not api_key:
        raise RuntimeError("GROQ_API_KEY no configurada.")

    client = Groq(api_key=api_key)
    completion = client.chat.completions.create(
        model=_GROQ_MODEL,
        messages=[
            {"role": "system", "content": system},
            {"role": "user",   "content": prompt},
        ],
        max_tokens=min(max_tokens, 8000),   # Groq limita a 8k en tier gratuito
        temperature=temperature,
    )
    return completion.choices[0].message.content


def _gemini(system: str, prompt: str, max_tokens: int = 1500,
            temperature: float = 0.2) -> str:
    """
    Llama a Gemini con fallback automático entre modelos.
    Si todos los modelos Gemini fallan (ej. restricción regional en Render),
    intenta automáticamente con Groq (LLaMA 3.3 70B).
    """
    try:
        from google import genai as _genai
        from google.genai import types as _gtypes
    except ImportError:
        # Si ni siquiera está instalado google-genai, ir directo a Groq
        return _groq(system, prompt, max_tokens, temperature)

    api_key = os.environ.get('GEMINI_API_KEY', '')
    if not api_key:
        return _groq(system, prompt, max_tokens, temperature)

    client = _genai.Client(api_key=api_key)
    ultimo_error = None

    for modelo in _GEMINI_MODELS_FALLBACK:
        try:
            resp = client.models.generate_content(
                model=modelo,
                contents=prompt,
                config=_gtypes.GenerateContentConfig(
                    system_instruction=system,
                    temperature=temperature,
                    max_output_tokens=max_tokens,
                )
            )
            return resp.text
        except Exception as e:
            ultimo_error = e
            err = str(e)
            if ('429' in err or 'RESOURCE_EXHAUSTED' in err or 'NOT_FOUND' in err
                    or 'FAILED_PRECONDITION' in err or 'location is not supported' in err.lower()):
                continue   # probar siguiente modelo Gemini
            raise          # error no recuperable (autenticación, red, etc.)

    # Todos los modelos Gemini fallaron — intentar con Groq
    import logging
    logging.warning(f"Todos los modelos Gemini fallaron ({ultimo_error}). Usando Groq como fallback.")
    return _groq(system, prompt, max_tokens, temperature)


def _gemini_json(system: str, prompt: str, max_tokens: int = 1500,
                 temperature: float = 0.1) -> dict:
    """Como _gemini() pero parsea la respuesta como JSON. Limpia markdown si hace falta."""
    texto = _gemini(system, prompt, max_tokens=max_tokens, temperature=temperature)
    texto = texto.strip()
    if texto.startswith("```"):
        texto = texto.split("\n", 1)[1].rsplit("```", 1)[0].strip()
    return json.loads(texto)


_PROMPT_VOTO_SISTEMA = """Sos un juez argentino redactando tu voto en un tribunal colegiado.
La ponencia ya fue redactada por otro miembro del tribunal.
Tu voto debe ser completamente autónomo y estar técnicamente fundado.

TIPOS DE VOTO:
- ADHESIÓN SIMPLE: "Adhiero al voto del Dr./Dra. [ponente]." — sin desarrollo propio.
- CONCURRENCIA: Llegás al mismo resultado pero por distintos fundamentos. Exponés tu razonamiento.
- DISIDENCIA PARCIAL: Acordás en parte con la ponencia pero diferís en algún punto (un rubro, una norma, la cuantía). Explicás qué aceptás y qué rechazás.
- DISIDENCIA TOTAL: Rechazás la ponencia en su totalidad. Exponés tu propio análisis y conclusión.

FORMATO:
El Dr./Dra. [nombre] dijo:
[Desarrollo del voto según el tipo]

Para concurrencias y disidencias: usá considerandos numerados en romano.
Sé técnico, preciso y referenciado en normas reales del derecho argentino.
No inventés precedentes. Si citás Fallos CSJN, usá el formato Fallos: TOMO:PÁGINA."""


_PROMPT_VOTO_ADHESION = "El Dr./Dra. {nombre} dijo:\n\nAdhiero al voto del/la vocal preopinante."


@app.post("/api/generar-voto")
async def generar_voto_vocal(req: VotoRequest):
    """Genera el voto individual de un vocal del tribunal"""

    # Adhesión simple: respuesta inmediata sin llamar a la IA
    if req.vocal.posicion == "adhesion":
        return {"voto": _PROMPT_VOTO_ADHESION.format(nombre=req.vocal.nombre)}

    posicion_label = {
        "concurrencia": "VOTO CONCURRENTE (mismo resultado, distinta fundamentación)",
        "disidencia_parcial": "DISIDENCIA PARCIAL (diferís en algún punto concreto)",
        "disidencia_total": "DISIDENCIA TOTAL (rechazás la ponencia en su totalidad)",
    }.get(req.vocal.posicion, req.vocal.posicion)

    caso = req.caso
    prompt = (
        f"CASO: {caso.get('caratula', '')}\n"
        f"Rama: {caso.get('rama', '')} | Instancia: {caso.get('instancia', '')}\n"
        f"Hechos: {caso.get('hechos_probados', '')[:600]}\n\n"
        f"PONENCIA YA REDACTADA:\n{req.sentencia_ponencia[:4000]}\n\n"
        f"TU POSICIÓN: {posicion_label}\n"
        f"EN QUÉ DIFERÍS: {req.vocal.fundamento or 'Ver posición indicada.'}\n\n"
        f"Redactá tu voto como el Dr./Dra. {req.vocal.nombre}."
    )

    try:
        voto = _gemini(_PROMPT_VOTO_SISTEMA, prompt, max_tokens=2500, temperature=0.3)
        return {"voto": voto.strip()}
    except Exception as e:
        return JSONResponse(status_code=500, content={"error": f"Error al generar voto: {e}"})


# ================================================================
# PONENCIAS COLEGIADAS
# ================================================================

@app.post("/api/ponencia/crear")
async def crear_ponencia(req: PonenciaCreateRequest):
    """
    Guarda la sentencia del vocal preopinante y devuelve un token compartible
    para que los otros vocales del tribunal puedan acceder y votar.
    Funciona para cualquier tribunal colegiado (penal, civil, laboral, TSJ, CSJN).
    """
    _limpiar_ponencias_vencidas()
    token = secrets.token_urlsafe(20)
    _ponencias[token] = {
        "caso":           req.caso,
        "sentencia":      req.sentencia,
        "vocal_ponente":  req.vocal_ponente,
        "created_at":     time.time(),
        "votos":          [],
    }
    return {
        "token":      token,
        "url":        f"/?ponencia={token}",
        "expira_en":  "24 horas",
    }


@app.get("/api/ponencia/{token}")
async def obtener_ponencia(token: str):
    """Devuelve la ponencia + estado de votos para el viewer de los otros vocales."""
    _limpiar_ponencias_vencidas()
    if token not in _ponencias:
        return JSONResponse(status_code=404, content={"error": "Ponencia no encontrada o expirada."})
    p = _ponencias[token]
    # Parsear secciones del considerando para mostrar checkboxes
    secciones = _parsear_considerandos(p["sentencia"].get("considerandos", ""))
    return {
        "caso":             p["caso"],
        "sentencia":        p["sentencia"],
        "vocal_ponente":    p["vocal_ponente"],
        "secciones":        list(secciones.keys()),   # ["I","II","III",...]
        "votos":            p["votos"],
        "vocales_pendientes": _vocales_sin_voto(p),
    }


@app.post("/api/ponencia/{token}/votar")
async def votar_en_ponencia(token: str, req: VotoColegiadoRequest):
    """
    Recibe el voto de un vocal sobre la ponencia:
      - adhesion_total   → texto de adhesión generado en Python (instantáneo)
      - adhesion_parcial → LLM genera sólo los considerandos propios
      - propio           → LLM genera un voto completamente independiente
    """
    _limpiar_ponencias_vencidas()
    if token not in _ponencias:
        return JSONResponse(status_code=404, content={"error": "Ponencia no encontrada o expirada."})

    p          = _ponencias[token]
    caso       = p["caso"]
    sentencia  = p["sentencia"]
    ponente    = p["vocal_ponente"]
    vocal      = req.nombre_vocal
    caratula   = caso.get("caratula", "")

    # ── Adhesión total ──────────────────────────────────────────────────────
    if req.tipo_voto == "adhesion_total":
        voto_texto = (
            f"El/La Vocal Dr./Dra. {vocal} dijo:\n\n"
            f"Adhiero íntegramente a los fundamentos y conclusión expuestos por "
            f"el/la Vocal Preopinante Dr./Dra. {ponente}, votando en idéntico sentido."
        )

    # ── Adhesión parcial ────────────────────────────────────────────────────
    elif req.tipo_voto == "adhesion_parcial":
        secciones        = _parsear_considerandos(sentencia.get("considerandos", ""))
        todos            = list(secciones.keys())
        propios          = req.considerandos_propios or []
        adhiere_a        = [s for s in todos if s not in propios]

        ponencia_resumen = "\n\n".join(
            f"{k}. {v[:600]}" for k, v in secciones.items() if k in propios
        )

        prompt_s = (
            "Sos vocal de un tribunal colegiado. Generá únicamente los considerandos "
            "donde diferís de la ponencia. Citá normas reales del derecho argentino. "
            "No inventes jurisprudencia. Sé técnico y preciso. Usá numeración romana."
        )
        prompt_u = (
            f"CASO: {caratula}\n"
            f"TRIBUNAL: {caso.get('tribunal','')}\n"
            f"RAMA: {caso.get('rama','')}\n\n"
            f"PONENCIA DEL VOCAL PREOPINANTE ({ponente}) — secciones en disputa:\n"
            f"{ponencia_resumen[:3500]}\n\n"
            f"MOTIVO DE LA DISIDENCIA: {req.fundamento_disidencia or 'Ver considerandos propios.'}\n\n"
            f"Redactá los considerandos {', '.join(propios)} como el/la Dr./Dra. {vocal}. "
            f"Al final añadí: 'En cuanto a los restantes considerandos "
            f"({', '.join(adhiere_a)}), adhiero íntegramente al voto del/la "
            f"Vocal Preopinante Dr./Dra. {ponente}.'"
        )
        try:
            texto_propio = _gemini(prompt_s, prompt_u, max_tokens=3000, temperature=0.3)
            voto_texto = f"El/La Vocal Dr./Dra. {vocal} dijo:\n\n{texto_propio.strip()}"
        except Exception as e:
            return JSONResponse(status_code=500, content={"error": f"Error al generar voto: {e}"})

    # ── Voto propio ─────────────────────────────────────────────────────────
    elif req.tipo_voto == "propio":
        # Reutilizamos el mismo motor de generación sobre el mismo caso.
        # El frontend puede también simplemente abrir la pantalla de generación
        # con el caso pre-cargado; aquí ofrecemos la versión API.
        from generador.motor_sentencias import generar, CasoInput
        try:
            caso_input = CasoInput(**caso)
            resultado = generar(caso_input, precedentes=[], normas_contexto=[])
            voto_texto = resultado.get("texto_completo", resultado.get("considerandos", ""))
            voto_texto = f"El/La Vocal Dr./Dra. {vocal} dijo:\n\n{voto_texto}"
        except Exception as e:
            return JSONResponse(status_code=500, content={"error": f"Error al generar voto propio: {e}"})

    else:
        return JSONResponse(status_code=400, content={"error": "tipo_voto debe ser adhesion_total, adhesion_parcial o propio."})

    # Guardar voto en el registro
    _ponencias[token]["votos"].append({
        "nombre":    vocal,
        "tipo":      req.tipo_voto,
        "texto":     voto_texto,
        "timestamp": time.time(),
    })

    return {
        "voto":           voto_texto,
        "tipo_voto":      req.tipo_voto,
        "nombre_vocal":   vocal,
        "total_votos":    len(_ponencias[token]["votos"]),
        "pendientes":     _vocales_sin_voto(_ponencias[token]),
    }


@app.get("/api/ponencia/{token}/componer")
async def componer_sentencia_final(token: str):
    """
    Una vez que todos los vocales votaron, compone la sentencia
    final con VISTOS + votos ordenados + dispositivo por mayoría.
    """
    _limpiar_ponencias_vencidas()
    if token not in _ponencias:
        return JSONResponse(status_code=404, content={"error": "Ponencia no encontrada."})

    p         = _ponencias[token]
    sentencia = p["sentencia"]
    votos     = p["votos"]
    ponente   = p["vocal_ponente"]

    # Construir texto completo
    partes = []
    partes.append(sentencia.get("vistos", ""))
    partes.append("\nY CONSIDERANDO:\n")

    # Voto preopinante
    partes.append(
        f"\nEl/La Vocal Preopinante Dr./Dra. {ponente} dijo:\n\n"
        + sentencia.get("considerandos", "").replace("CONSIDERANDO:\n\n", "").strip()
    )

    # Votos restantes
    for v in votos:
        partes.append(f"\n\n{'─'*60}\n\n{v['texto']}")

    # Dispositivo
    partes.append(f"\n\n{'═'*60}\n\nPOR ELLO, EL TRIBUNAL RESUELVE:\n\n")
    partes.append(sentencia.get("resuelvo", "").replace("RESUELVO:\n\n", "").strip())

    texto_final = "\n".join(partes)

    return {
        "texto_final":    texto_final,
        "vocal_ponente":  ponente,
        "votos_emitidos": len(votos) + 1,   # +1 preopinante
        "pendientes":     _vocales_sin_voto(p),
    }


@app.post("/api/exportar-pdf")
async def exportar_pdf(req: ExportarPDFRequest):
    """Genera un PDF de la sentencia y lo devuelve como base64"""
    try:
        import io, base64
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import cm
            from reportlab.lib import colors
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
            from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER
            REPORTLAB_OK = True
        except ImportError:
            REPORTLAB_OK = False

        if not REPORTLAB_OK:
            return JSONResponse(status_code=503, content={"error": "reportlab no instalado."})

        buf = io.BytesIO()
        doc = SimpleDocTemplate(
            buf, pagesize=A4,
            rightMargin=2.5*cm, leftMargin=3*cm,
            topMargin=3*cm, bottomMargin=2.5*cm
        )

        styles = getSampleStyleSheet()
        estilo_normal = ParagraphStyle(
            'Normal_JA',
            parent=styles['Normal'],
            fontName='Times-Roman',
            fontSize=12,
            leading=22,
            alignment=TA_JUSTIFY,
            spaceAfter=6,
        )
        estilo_titulo = ParagraphStyle(
            'Titulo_JA',
            parent=styles['Normal'],
            fontName='Times-Bold',
            fontSize=13,
            leading=18,
            alignment=TA_CENTER,
            spaceAfter=12,
            spaceBefore=16,
            textColor=colors.HexColor('#1a365d'),
        )
        estilo_seccion = ParagraphStyle(
            'Seccion_JA',
            parent=styles['Normal'],
            fontName='Times-Bold',
            fontSize=12,
            leading=18,
            spaceAfter=6,
            spaceBefore=14,
            textColor=colors.HexColor('#1a365d'),
        )

        story = []

        # Encabezado
        if req.caratula:
            story.append(Paragraph(req.caratula.upper(), estilo_titulo))
        if req.tribunal:
            story.append(Paragraph(req.tribunal, ParagraphStyle('sub', parent=estilo_normal, alignment=TA_CENTER, fontSize=11)))
        if req.instancia:
            story.append(Paragraph(f"Instancia: {req.instancia}", ParagraphStyle('sub2', parent=estilo_normal, alignment=TA_CENTER, fontSize=10, textColor=colors.grey)))
        story.append(HRFlowable(width="100%", thickness=1.5, color=colors.HexColor('#c6a84b'), spaceAfter=16))

        # Cuerpo
        SECCIONES = {'VISTOS', 'Y CONSIDERANDO', 'CONSIDERANDO', 'RESUELVO', 'SE RESUELVE'}
        for linea in req.texto.split('\n'):
            t = linea.strip()
            if not t:
                story.append(Spacer(1, 6))
                continue
            # Detectar sección principal
            t_upper = t.rstrip(':').upper()
            if t_upper in SECCIONES or any(t_upper.startswith(s) for s in SECCIONES):
                story.append(Paragraph(t, estilo_seccion))
            else:
                # Escapar caracteres especiales para ReportLab
                t_safe = t.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                story.append(Paragraph(t_safe, estilo_normal))

        doc.build(story)
        pdf_bytes = buf.getvalue()
        return {"pdf_base64": base64.b64encode(pdf_bytes).decode('utf-8')}

    except Exception as e:
        return JSONResponse(status_code=500, content={"error": f"Error generando PDF: {e}"})


# Autenticación simple basada en env vars
# Formato esperado en env: TRIBUNALES_AUTH="tribunal_id:hash,tribunal_id2:hash2"
# El hash es sha256 de la contraseña
def _verificar_credenciales(tribunal: str, password: str) -> bool:
    """Verifica credenciales de tribunal contra la variable de entorno TRIBUNALES_AUTH"""
    auth_config = os.environ.get('TRIBUNALES_AUTH', '')
    if not auth_config:
        # Si no hay config, modo demo: cualquier tribunal/password pasa
        return True
    pwd_hash = hashlib.sha256(password.encode()).hexdigest()
    for entry in auth_config.split(','):
        parts = entry.strip().split(':')
        if len(parts) == 2 and parts[0].strip() == tribunal and hmac.compare_digest(parts[1].strip(), pwd_hash):
            return True
    return False


@app.post("/api/login-tribunal")
async def login_tribunal(req: UsuarioLogin):
    """[LEGACY] Autenticación simple de tribunal - renombrado para no conflictuar con /api/login"""
    if _verificar_credenciales(req.tribunal, req.password):
        # Token simple: HMAC del tribunal + timestamp del día (válido 24hs)
        from datetime import date
        secret = os.environ.get('SESSION_SECRET', 'justicia-argentina-secret')
        payload = f"{req.tribunal}:{date.today().isoformat()}"
        token = hmac.new(secret.encode(), payload.encode(), hashlib.sha256).hexdigest()  # type: ignore
        return {"ok": True, "token": token, "tribunal": req.tribunal}
    return JSONResponse(status_code=401, content={"ok": False, "error": "Credenciales incorrectas"})


@app.post("/api/verificar-token")
async def verificar_token(req: dict):
    """Verifica si un token de sesión es válido"""
    from datetime import date
    token = req.get("token", "")
    tribunal = req.get("tribunal", "")
    secret = os.environ.get('SESSION_SECRET', 'justicia-argentina-secret')
    payload = f"{tribunal}:{date.today().isoformat()}"
    expected = hmac.new(secret.encode(), payload.encode(), hashlib.sha256).hexdigest()
    valid = hmac.compare_digest(token, expected) if token and tribunal else False
    return {"valid": valid, "tribunal": tribunal if valid else ""}


_PROMPT_COHERENCIA_SISTEMA = """Sos un revisor jurídico que analiza la coherencia interna de proyectos de sentencia argentinos.

Verificá:
1. ¿El RESUELVO responde TODAS las cuestiones planteadas en el CONSIDERANDO? (congruencia)
2. ¿Hay contradicciones entre los argumentos del CONSIDERANDO y la decisión del RESUELVO?
3. ¿Hay argumentos en el CONSIDERANDO que llevarían a una conclusión distinta a la del RESUELVO?
4. ¿Las costas son coherentes con el resultado?
5. ¿Hay puntos del RESUELVO que no tienen sustento en el CONSIDERANDO?

Respondé EXCLUSIVAMENTE con un objeto JSON válido, sin markdown ni texto adicional:
{
  "coherente": true|false,
  "score": 0.0-1.0,
  "observaciones": [
    {"tipo": "contradiccion|omision|incongruencia|advertencia", "descripcion": "...", "gravedad": "alta|media|baja"}
  ]
}
Si no hay observaciones, devolvé array vacío."""


@app.post("/api/coherencia")
async def verificar_coherencia(req: CoherenciaRequest):
    """Verifica la coherencia interna de la sentencia generada"""
    cuestiones_txt = "\n".join(f"- {c}" for c in req.cuestiones_a_resolver) if req.cuestiones_a_resolver else "No especificadas."
    prompt = (
        f"Carátula: {req.caratula}\n"
        f"Cuestiones a resolver:\n{cuestiones_txt}\n\n"
        f"PROYECTO DE SENTENCIA:\n{req.sentencia[:6000]}"
    )
    try:
        return _gemini_json(_PROMPT_COHERENCIA_SISTEMA, prompt, max_tokens=800)
    except json.JSONDecodeError as e:
        return JSONResponse(status_code=500, content={"error": f"Respuesta IA no es JSON: {e}"})
    except Exception as e:
        return JSONResponse(status_code=500, content={"error": f"Error: {e}"})


_PROMPT_ANALISIS_SISTEMA = """Sos un asistente jurídico experto en derecho argentino. Analizás casos judiciales desde la perspectiva del tribunal.

Los hechos ya están fijados por el juez. Tu análisis se enfoca exclusivamente en el plano jurídico.

Respondé EXCLUSIVAMENTE con un objeto JSON válido, sin markdown, sin texto adicional antes ni después.

Formato requerido:
{
  "mapa_controversia": {
    "nucleo_juridico": "Una oración que resume el nudo normativo central del caso",
    "posicion_actor": "Qué norma/interpretación invoca el actor y por qué",
    "posicion_demandado": "Qué norma/interpretación invoca el demandado y por qué",
    "normas_en_disputa": ["Art. X Ley Y — descripción breve", "..."],
    "tension_interpretativa": "Cuál es la tensión entre las dos posiciones jurídicas"
  },
  "cuestiones_oficio": [
    {"tipo": "prescripción|caducidad|legitimación|nulidad|inconstitucionalidad|competencia|otra", "descripcion": "...", "urgencia": "alta|media|baja", "norma": "Art. X Ley Y"}
  ],
  "plazos_procesales": [
    {"nombre": "Nombre del plazo", "dias": N, "norma": "Art. X CPCC...", "observacion": "..."}
  ],
  "citas_partes": ["fallo o norma citada por alguna parte", "..."]
}

Si no hay cuestiones de oficio detectadas, devolvé array vacío. Si no hay plazos relevantes, devolvé array vacío."""


@app.post("/api/analizar")
async def analizar_caso(req: CasoRequest):
    """Análisis jurídico auxiliar: mapa de controversia, cuestiones de oficio, plazos, citas"""
    partes_txt = ""
    for p in req.partes:
        partes_txt += (
            f"\n{p.rol.upper()} — {p.nombre}:\n"
            f"  Pretensión: {p.pretension}\n"
            f"  Fundamentos jurídicos: {p.fundamentos_juridicos}\n"
            f"  Argumentos: {p.argumentos}\n"
            f"  Jurisprudencia citada: {p.jurisprudencia_citada}\n"
        )
    prompt = (
        f"CASO: {req.caratula}\n"
        f"Rama: {req.rama} | Tipo proceso: {req.tipo_proceso} | "
        f"Jurisdicción: {req.jurisdiccion} | Tribunal: {req.tribunal}\n\n"
        f"CUESTIONES A RESOLVER:\n" + "\n".join(f"- {c}" for c in req.cuestiones_a_resolver) + "\n\n"
        f"POSICIONES DE LAS PARTES:{partes_txt}"
    )
    try:
        analisis = _gemini_json(_PROMPT_ANALISIS_SISTEMA, prompt, max_tokens=1500)
    except json.JSONDecodeError as e:
        return JSONResponse(status_code=500, content={"error": f"Respuesta IA no es JSON válido: {e}"})
    except Exception as e:
        return JSONResponse(status_code=500, content={"error": f"Error al analizar: {e}"})

    # Feature 4: verificar citas de partes contra el índice
    citas = analisis.get("citas_partes", [])
    verificacion_citas = []
    if citas and INDEXADOR_DISPONIBLE:
        for cita in citas[:6]:  # máximo 6 para no demorar
            try:
                resultados = indexador.buscar_precedentes(query=cita, rama=req.rama, n_results=1)
                if resultados and resultados[0].get("distancia", 1) < 0.5:
                    r = resultados[0]
                    verificacion_citas.append({
                        "cita": cita,
                        "encontrada": True,
                        "caratula": r["metadata"].get("caratula", ""),
                        "tribunal": r["metadata"].get("tribunal_tipo", ""),
                        "fecha": r["metadata"].get("fecha", ""),
                    })
                else:
                    verificacion_citas.append({"cita": cita, "encontrada": False})
            except Exception:
                verificacion_citas.append({"cita": cita, "encontrada": None})

    analisis["verificacion_citas_partes"] = verificacion_citas
    return analisis


_PROMPT_CHAT_SISTEMA = """Sos un asistente jurídico especializado que ayuda a magistrados argentinos a revisar proyectos de sentencia.

Tu rol:
- Evaluar si las observaciones del magistrado son jurídicamente correctas y están bien fundadas
- Señalar si hay normas o jurisprudencia relevante no considerada
- Identificar inconsistencias en el razonamiento jurídico
- Responder con precisión técnica: doctrina, jurisprudencia CSJN, CCyCN, CPCC, etc.
- Ser directo sobre si una observación está o NO bien orientada, con fundamento

Cuando el magistrado señale una omisión:
- Si está bien fundada: confirmalo y desarrollá el argumento
- Si está parcialmente fundada: aclará qué parte es correcta y qué no
- Si NO está bien fundada: decilo claramente con fundamento jurídico

Respondé en español, de manera concisa pero técnicamente precisa. Máximo 350 palabras por respuesta."""


@app.post("/api/chat")
async def chat_sentencia(req: ChatRequest):
    """Chat interactivo para revisar y discutir el proyecto de sentencia"""
    historial_txt = ""
    for msg in req.historial:
        label = "MAGISTRADO" if msg.rol == "user" else "SISTEMA"
        historial_txt += f"\n\n{label}: {msg.contenido}"
    caso = req.caso
    prompt = (
        f"PROYECTO DE SENTENCIA ACTUAL:\n{req.sentencia_actual[:3500]}\n\n"
        f"DATOS DEL CASO:\n"
        f"Carátula: {caso.get('caratula', '')}\n"
        f"Rama: {caso.get('rama', '')}\n"
        f"Hechos: {caso.get('hechos_probados', '')[:600]}\n"
        f"\nCONVERSACIÓN PREVIA:{historial_txt}\n\n"
        f"MAGISTRADO: {req.mensaje}"
    )
    try:
        respuesta = _gemini(_PROMPT_CHAT_SISTEMA, prompt, max_tokens=700, temperature=0.3)
        return {"respuesta": respuesta.strip()}
    except Exception as e:
        return JSONResponse(status_code=500, content={"error": f"Error al consultar IA: {e}"})


_PROMPT_REGENERAR_SISTEMA = """Sos un juez argentino experto redactando sentencias judiciales.
Reescribí el proyecto de sentencia incorporando todas las correcciones y observaciones validadas en la conversación.
Mantené la estructura formal argentina: VISTOS, CONSIDERANDO (con análisis detallado), RESUELVO.
La sentencia nueva debe ser completa y autosuficiente. No uses frases como "como se señaló" o "según lo anterior"."""


@app.post("/api/regenerar")
async def regenerar_con_chat(req: RegenerarRequest):
    """Regenera la sentencia incorporando el feedback del chat"""
    historial_txt = ""
    for msg in req.historial:
        label = "MAGISTRADO" if msg.rol == "user" else "SISTEMA"
        historial_txt += f"\n\n{label}: {msg.contenido}"

    prompt = (
        f"PROYECTO DE SENTENCIA ORIGINAL:\n{req.sentencia_actual}\n\n"
        f"CONVERSACIÓN DE REVISIÓN CON EL MAGISTRADO:{historial_txt}\n\n"
        f"Reescribí la sentencia completa incorporando todas las correcciones y observaciones válidas de la conversación."
    )

    try:
        texto = _gemini(_PROMPT_REGENERAR_SISTEMA, prompt, max_tokens=4000, temperature=0.3)
        return {"texto_completo": texto.strip()}
    except Exception as e:
        return JSONResponse(status_code=500, content={"error": f"Error al regenerar sentencia: {e}"})


_PROMPT_VALORACION_PRUEBA = """Sos un experto en razonamiento probatorio judicial aplicado al derecho argentino.
Tu función es analizar si el razonamiento inductivo del tribunal —el salto inferencial de la prueba al hecho—
es lógicamente válido según los estándares de la sana crítica racional.

DISTINCIÓN FUNDAMENTAL (imprescindible):
- NO determinás qué ocurrió (quaestio facti — intocable en casación)
- SÍ evaluás si el MÉTODO DE VALORACIÓN fue lógicamente correcto
- SÍ evaluás si la prueba producida soporta inductivamente la conclusión fáctica

ESTÁNDARES DE EVALUACIÓN:
1. SANA CRÍTICA RACIONAL: lógica formal, máximas de la experiencia, conocimiento científico
   (art. 193 CPP Córdoba / art. 241 CPPN y equivalentes provinciales)
2. VALORACIÓN INTEGRAL: no puede omitirse prueba decisiva que contradiga la conclusión
3. COHERENCIA INTERNA: el razonamiento no puede ser autocontradictorio
4. FUNDAMENTACIÓN: la sentencia debe explicitar los pasos del razonamiento
5. NON SEQUITUR: la conclusión debe seguirse con alta probabilidad de las premisas probadas

VICIOS CASABLES (entidad para abrir la casación):
- Omisión de prueba decisiva y conducente (el tribunal ignoró prueba que contraría el hecho)
- Razonamiento contrario a las reglas de la lógica (error formal: non sequitur, contradictio)
- Violación de máximas de la experiencia (contrario al id quod plerumque accidit verificable)
- Contradicción interna entre la valoración probatoria y la conclusión
- Valoración selectiva: se eligió la parte favorable de una prueba ignorando la parte contraria
- Otorgar certeza donde solo hay probabilidad (salto lógico injustificado)

VICIOS NO CASABLES (dentro del margen de valoración discrecional legítima):
- El tribunal eligió una entre dos interpretaciones igualmente razonables de la prueba
- La credibilidad de un testigo (salvo irrazonabilidad manifiesta o contradicción con otras pruebas)
- La prueba es débil pero el tribunal la integró coherentemente con otras pruebas

Respondé EXCLUSIVAMENTE con JSON válido, sin markdown:
{
  "reconstruccion": {
    "cadena_inferencial": "descripción de la cadena lógica del tribunal: prueba X → inferencia Y → hecho H",
    "premisas": ["P1: ...", "P2: ..."],
    "conclusion_factica": "el hecho que el tribunal tuvo por acreditado"
  },
  "analisis_logico": {
    "validez": "válido|cuestionable|inválido",
    "score_razonamiento": 0.0,
    "problemas": [
      {
        "tipo": "non_sequitur|omision_prueba|contradiccion_interna|violacion_maximas|valoracion_selectiva|certeza_infundada|otro",
        "descripcion": "descripción precisa y técnica del problema",
        "gravedad": "alta|media|baja",
        "prueba_involucrada": "qué elemento probatorio específico está en cuestión"
      }
    ]
  },
  "evaluacion_agravio": {
    "tiene_razon_impugnante": true,
    "casable": true,
    "fundamento": "fundamento técnico de la conclusión",
    "estandar_aplicado": "sana crítica racional — art. 193 CPP Córdoba"
  },
  "sintesis": "Párrafo de síntesis del análisis (2-3 párrafos) con lenguaje de casación"
}
Si el razonamiento del tribunal es correcto y el agravio no tiene razón, también decilo con fundamento."""


@app.post("/api/analizar-valoracion-prueba")
async def analizar_valoracion_prueba(req: ValoracionPruebaRequest):
    """
    Analiza el razonamiento probatorio inductivo del tribunal:
    ¿La prueba producida soporta lógicamente la conclusión fáctica?
    ¿Tiene razón el impugnante cuando dice que la valoración fue arbitraria?
    """
    # Cuando el frontend envía la sentencia completa como fuente única,
    # construimos el prompt indicando al modelo que extraiga los componentes
    sentencia_completa = req.prueba_transcripta  # el frontend envía sentencia_recurrida aquí
    tiene_desglose = bool(req.hecho_probado and len(req.hecho_probado) > 20)

    if tiene_desglose:
        # Versión con campos separados (ingreso manual detallado)
        bloque_fuente = f"""══ HECHO QUE EL TRIBUNAL TUVO POR PROBADO ══
{req.hecho_probado}

══ PRUEBA PRODUCIDA Y VALORADA ══
{req.prueba_transcripta[:4000]}

══ RAZONAMIENTO PROBATORIO DE LA SENTENCIA ══
{req.razonamiento_probatorio[:3000]}"""
    else:
        # Versión con sentencia completa — el modelo extrae los componentes
        bloque_fuente = f"""══ SENTENCIA RECURRIDA (texto completo o síntesis) ══
A partir de esta sentencia, identificá: (1) los hechos que tuvo por probados,
(2) la prueba en que se basó, (3) el razonamiento valorativo empleado.
Luego procedé con el análisis lógico.

{sentencia_completa[:6000]}"""

    prompt = f"""CASO: {req.caratula}
Tribunal recurrido: {req.tribunal_recurrido}
Tipo de proceso: {req.tipo_proceso} | Jurisdicción: {req.jurisdiccion}

{bloque_fuente}

══ AGRAVIOS DEL IMPUGNANTE SOBRE LA VALORACIÓN ══
{req.agravios_valoracion[:2000] if req.agravios_valoracion else "(no especificados)"}

Analizá si el razonamiento inductivo del tribunal es lógicamente válido y si el impugnante tiene razón."""

    try:
        return _gemini_json(_PROMPT_VALORACION_PRUEBA, prompt, max_tokens=2000, temperature=0.1)
    except json.JSONDecodeError as e:
        return JSONResponse(status_code=500, content={"error": f"Respuesta IA no es JSON: {e}"})
    except Exception as e:
        return JSONResponse(status_code=500, content={"error": f"Error: {e}"})


def _ocr_pdf_con_gemini(contenido_pdf: bytes) -> tuple[str, int]:
    """
    OCR de PDF escaneado usando Gemini Vision.
    Convierte cada página a imagen con pymupdf y manda al modelo multimodal.
    Devuelve (texto_extraído, cantidad_páginas).
    """
    import fitz  # pymupdf
    import base64

    api_key = os.environ.get('GEMINI_API_KEY', '')
    if not api_key:
        raise RuntimeError("GEMINI_API_KEY no configurada para OCR.")

    doc = fitz.open(stream=contenido_pdf, filetype="pdf")
    paginas = len(doc)
    textos = []

    # Modelos con capacidad de visión, en orden de preferencia
    modelos_vision = ['gemini-2.0-flash', 'gemini-2.5-flash', 'gemini-2.0-flash-lite']

    from google import genai as _genai
    from google.genai import types as _gtypes

    client = _genai.Client(api_key=api_key)

    # Procesar de a máximo 20 páginas para no exceder límites
    MAX_PAGINAS = 20
    if paginas > MAX_PAGINAS:
        import logging
        logging.warning(f"OCR: PDF tiene {paginas} páginas, procesando solo las primeras {MAX_PAGINAS}.")

    for num_pag in range(min(paginas, MAX_PAGINAS)):
        page = doc[num_pag]
        # Renderizar a imagen PNG a 150 DPI (balance calidad/tamaño)
        mat = fitz.Matrix(150/72, 150/72)
        pix = page.get_pixmap(matrix=mat, alpha=False)
        img_bytes = pix.tobytes("png")
        img_b64 = base64.b64encode(img_bytes).decode()

        prompt_ocr = (
            "Transcribí exactamente el texto de esta imagen de una sentencia judicial argentina. "
            "Mantené párrafos, numeración y estructura. No agregues comentarios ni explicaciones. "
            "Solo el texto tal como aparece en la imagen."
        )

        ultimo_error = None
        for modelo in modelos_vision:
            try:
                resp = client.models.generate_content(
                    model=modelo,
                    contents=[
                        _gtypes.Part.from_bytes(data=img_bytes, mime_type="image/png"),
                        prompt_ocr,
                    ],
                    config=_gtypes.GenerateContentConfig(
                        max_output_tokens=4096,
                        temperature=0.0,
                    )
                )
                textos.append(resp.text or "")
                break
            except Exception as e:
                ultimo_error = e
                err = str(e)
                if ('FAILED_PRECONDITION' in err or '429' in err or
                        'NOT_FOUND' in err or 'location is not supported' in err.lower()):
                    continue
                raise
        else:
            raise RuntimeError(f"OCR falló en página {num_pag + 1}: {ultimo_error}")

    doc.close()
    return "\n\n".join(textos), paginas


@app.post("/api/extraer-texto")
async def extraer_texto(archivo: UploadFile = File(...)):
    """
    Extrae el texto de un archivo PDF, Word (.docx) o TXT subido por el usuario.
    Para PDFs escaneados (sin capa de texto), aplica OCR con Gemini Vision.
    Devuelve { "texto": str, "paginas": int, "nombre": str, "ocr": bool }
    """
    nombre = archivo.filename or ""
    ext = nombre.rsplit(".", 1)[-1].lower() if "." in nombre else ""
    contenido = await archivo.read()

    if ext == "pdf":
        try:
            import pdfplumber
        except ImportError:
            return JSONResponse(status_code=500, content={"error": "pdfplumber no instalado."})
        try:
            texto_paginas = []
            paginas = 0
            with pdfplumber.open(io.BytesIO(contenido)) as pdf:
                paginas = len(pdf.pages)
                for page in pdf.pages:
                    t = page.extract_text()
                    if t:
                        texto_paginas.append(t)

            texto = "\n\n".join(texto_paginas).strip()
            chars_utiles = len(texto.replace('\n', '').replace(' ', ''))

            # Si el texto extraído es demasiado escaso → PDF escaneado → OCR
            # Umbral: menos de 100 caracteres útiles por página en promedio
            if chars_utiles < paginas * 100:
                import logging
                logging.info(f"PDF '{nombre}': texto escaso ({chars_utiles} chars), aplicando OCR con Gemini Vision.")
                try:
                    texto_ocr, paginas = _ocr_pdf_con_gemini(contenido)
                    aviso = ""
                    if paginas > 20:
                        aviso = f" (solo primeras 20 de {paginas} páginas — PDF muy extenso)"
                    return {
                        "texto": texto_ocr,
                        "paginas": min(paginas, 20),
                        "nombre": nombre,
                        "ocr": True,
                        "aviso": aviso,
                    }
                except Exception as e_ocr:
                    # Si el OCR también falla, devolver lo que había (puede ser vacío)
                    return {
                        "texto": texto,
                        "paginas": paginas,
                        "nombre": nombre,
                        "ocr": False,
                        "aviso": f"OCR no disponible: {e_ocr}. El PDF puede ser una imagen sin capa de texto.",
                    }

            return {"texto": texto, "paginas": paginas, "nombre": nombre, "ocr": False}

        except Exception as e:
            return JSONResponse(status_code=500, content={"error": f"Error al leer PDF: {e}"})

    elif ext in ("docx", "doc"):
        try:
            from docx import Document
        except ImportError:
            return JSONResponse(status_code=500, content={"error": "python-docx no instalado."})
        try:
            doc = Document(io.BytesIO(contenido))
            parrafos = [p.text for p in doc.paragraphs if p.text.strip()]
            texto = "\n".join(parrafos)
            return {"texto": texto, "paginas": len(doc.paragraphs), "nombre": nombre, "ocr": False}
        except Exception as e:
            return JSONResponse(status_code=500, content={"error": f"Error al leer Word: {e}"})

    elif ext == "txt":
        try:
            texto = contenido.decode("utf-8", errors="replace")
            return {"texto": texto, "paginas": 1, "nombre": nombre, "ocr": False}
        except Exception as e:
            return JSONResponse(status_code=500, content={"error": f"Error al leer txt: {e}"})

    else:
        return JSONResponse(
            status_code=400,
            content={"error": f"Formato no soportado: .{ext}. Use PDF, DOCX o TXT."}
        )


@app.post("/api/buscar-normativa-municipal")
async def buscar_normativa_municipal(req: dict):
    """
    Búsqueda on-demand de normativa municipal vía SAIJ API.
    Usa la API de SAIJ para buscar ordenanzas, decretos y resoluciones municipales.

    Body: { "municipio": str, "query": str, "n_results": int }
    """
    import sys
    sys.path.insert(0, "C:/Users/Juan/AppData/Roaming/Python/Python314/site-packages")
    try:
        from saij_mcp.client import search as saij_search
    except ImportError:
        return JSONResponse(
            status_code=503,
            content={"error": "saij-mcp no instalado en este entorno."}
        )

    municipio = req.get("municipio", "")
    query_text = req.get("query", "")
    n_results = min(req.get("n_results", 5), 25)

    query = f"{municipio} {query_text}".strip()
    if not query:
        return JSONResponse(status_code=400, content={"error": "Debe especificar municipio o query."})

    try:
        resultado = saij_search(query, doc_type="legislacion", field="titulo", limit=n_results)
        normas = []
        for doc in resultado.get("results", []):
            normas.append({
                "tipo": doc.get("tipo_norma", ""),
                "numero": doc.get("numero_norma", ""),
                "titulo": doc.get("titulo", ""),
                "jurisdiccion": doc.get("jurisdiccion", ""),
                "url": doc.get("url", ""),
            })
        return {
            "municipio": municipio,
            "query": query_text,
            "total_api": resultado.get("total", 0),
            "resultados": normas,
            "nota": "Resultados obtenidos en tiempo real desde SAIJ. Verificar vigencia."
        }
    except Exception as e:
        return JSONResponse(status_code=502, content={"error": f"Error consultando SAIJ: {e}"})


@app.get("/api/estadisticas")
async def estadisticas():
    """Estadísticas del sistema"""
    from indexador.qdrant_store import get_qdrant_client
    stats = {
        "indexador_jurisprudencia": INDEXADOR_DISPONIBLE,
        "indexador_legislacion_nacional": INFOLEG_DISPONIBLE,
        "indexador_legislacion_provincial": PROVINCIAL_DISPONIBLE,
        "colecciones": {}
    }
    # Contar todos los puntos en todas las colecciones Qdrant
    try:
        client = get_qdrant_client()
        if client:
            for col in client.get_collections().collections:
                try:
                    cnt = client.count(collection_name=col.name).count
                    stats["colecciones"][col.name] = cnt
                except Exception:
                    pass
    except Exception as e:
        stats["error_colecciones"] = str(e)
    return stats


# ================================================================
# FRONTEND
# ================================================================

@app.get("/login", response_class=HTMLResponse)
async def login_page(request: Request):
    """Página de login con código de invitación."""
    if _esta_autenticado(request):
        return RedirectResponse(url='/', status_code=302)
    return HTMLResponse(content=_LOGIN_HTML.format(error_html=''))


@app.post("/api/login", response_class=RedirectResponse)
async def login(request: Request):
    """Valida el código de invitación y emite cookie de sesión."""
    # Leer formulario directamente sin validación de Pydantic
    form_data = await request.form()
    codigo = str(form_data.get('codigo', '')).strip().upper()

    # Por ahora: permitir acceso sin restricción (cualquier código o sin código)
    # Más adelante: if not AUTH_HABILITADO or codigo in INVITE_CODES:
    if True:
        # Emitir cookie de sesión
        cookie_val = _crear_cookie(codigo or 'TESTING')
        resp = RedirectResponse(url='/', status_code=302)
        resp.set_cookie(
            key=SESSION_COOKIE,
            value=cookie_val,
            max_age=SESSION_DAYS * 86400,
            httponly=True,
            samesite='lax',
            secure=os.environ.get('RENDER', '') != '',  # HTTPS en Render, HTTP local
        )
        # Registrar acceso
        _accesos[codigo or 'GUEST'].append({
            'ip': request.client.host if request.client else 'unknown',
            'timestamp': time.strftime('%Y-%m-%d %H:%M:%S'),
        })
        import logging
        logging.info(f"[AUTH] Acceso concedido: código={codigo or 'GUEST'} ip={request.client.host if request.client else '?'}")
        return resp


@app.get("/logout")
async def logout():
    """Cierra la sesión."""
    resp = RedirectResponse(url='/login', status_code=302)
    resp.delete_cookie(SESSION_COOKIE)
    return resp


@app.get("/api/admin/accesos")
async def ver_accesos(request: Request):
    """Lista accesos registrados. Solo accesible con SESSION_SECRET como token."""
    token = request.headers.get('X-Admin-Token', '')
    if not hmac.compare_digest(token, SESSION_SECRET):
        return JSONResponse(status_code=403, content={"error": "No autorizado"})
    return {"accesos": dict(_accesos), "codigos_activos": list(INVITE_CODES)}


@app.get("/", response_class=HTMLResponse)
async def frontend(request: Request):
    """Sirve la interfaz web — requiere autenticación si AUTH_HABILITADO."""
    if not _esta_autenticado(request):
        return RedirectResponse(url='/login', status_code=302)
    frontend_path = os.path.join(ROOT_DIR, "src", "frontend", "index.html")
    if os.path.exists(frontend_path):
        with open(frontend_path, 'r', encoding='utf-8') as f:
            return HTMLResponse(content=f.read())
    return HTMLResponse(content="<h1>Frontend no encontrado</h1>", status_code=404)


# ================================================================
# MAIN
# ================================================================

if __name__ == '__main__':
    import uvicorn
    port = int(os.environ.get('PORT', 8000))
    print("=" * 70)
    print("JUSTICIA ARGENTINA - Sistema de Resolución Asistida")
    print(f"Servidor: http://localhost:{port}")
    print(f"Indexador: {'Disponible' if INDEXADOR_DISPONIBLE else 'No disponible'}")
    print(f"API Key: {'Configurada' if os.environ.get('ANTHROPIC_API_KEY') else 'No configurada (modo demo)'}")
    print("=" * 70)
    uvicorn.run(app, host="0.0.0.0", port=port)
